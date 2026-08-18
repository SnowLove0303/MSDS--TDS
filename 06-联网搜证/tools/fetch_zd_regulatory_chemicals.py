# -*- coding: utf-8 -*-
from playwright.sync_api import sync_playwright
import docx, re
from pathlib import Path

CHROME = r'E:\MorenAnzhuangLujing\Chrome\Chrome\Application\chrome.exe'
ROOT_OUT = Path(r"F:\正式项目与模块化内容\冠志\MSDS\Word 覆写模块\数据库与推断引擎\法规匹配库\法规原文归档\目录公告")

TARGETS = [
    {
        "name": "重点监管的危险化学品名录（两批完整版）",
        "url": "https://www.so.com/link?m=eVUjapyTtOuDze%2FMiCbjc1TDU2WvCsHWCKj2t9Fg3ciWK02z%2Fwnel0dKjsJXC1PLQ2JlKAcq2dqhPWPPmgHX3vM8Kypn6AsSFtT%2BrpIksn2%2FQMAC6Mo6cukcBx76NIRDKSnby7GHJIwR51EqoLsPivF9skMqrTPVlmeTGDBp1RqboLFz4In2h0fYxNwP4QeFpoi7aU1nb968tv3pAzV1UqdSiS24Uy4nFB4G5KoJqYqGn57ShMuROO0g4t%2FVFJphCeV4f2OHggh%2BlJbGlzaU9ZodHjzRv%2FOVHbEVb%2BT7gs5Bsm5ooV6Ha8bUSmo9MFmv1g62NnhJBu2o7ZaA8"
    }
]

def main():
    with sync_playwright() as p:
        b = p.chromium.launch(headless=True, executable_path=CHROME, args=['--no-sandbox'])
        context = b.new_context()
        page = context.new_page()
        
        # 检索首批和第二批重点监管名录
        page.goto("https://www.so.com/s?q=重点监管危险化学品名录+两批完整名单+CAS", timeout=30000)
        page.wait_for_timeout(2000)
        
        links = page.evaluate('''() => Array.from(document.querySelectorAll('h3 a, .res-title a')).map(a => [(a.innerText || '').trim(), a.href]);''')
        print("Search result links:")
        for l in links[:5]:
            print("  ", l)
            
        if links:
            page.goto(links[0][1], timeout=30000)
            page.wait_for_timeout(2000)
            text = page.evaluate('() => document.body.innerText')
            print("Article text len:", len(text))
            
            docx_p = ROOT_OUT / "重点监管危险化学品名录（第一批和第二批完整版）.docx"
            txt_p = ROOT_OUT / "重点监管危险化学品名录（第一批和第二批完整版）.txt"
            
            txt_p.write_text(text, encoding="utf-8")
            doc = docx.Document()
            doc.add_heading("重点监管危险化学品名录（第一批和第二批完整版）", level=0)
            for line in text.splitlines():
                ls = line.strip()
                if ls:
                    doc.add_paragraph(ls)
            doc.save(str(docx_p))
            print(f"✅ 成功保存: {docx_p.name}")
            
        b.close()

if __name__ == '__main__':
    main()
