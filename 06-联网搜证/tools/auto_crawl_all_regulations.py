# -*- coding: utf-8 -*-
"""全自动化搜索、提取、校验并归档所有缺失的法律法规与行业标准."""
import urllib.parse, re, time
from pathlib import Path
from playwright.sync_api import sync_playwright
import docx

CHROME = r'E:\MorenAnzhuangLujing\Chrome\Chrome\Application\chrome.exe'
ROOT_DIR = Path(r'F:\正式项目与模块化内容\冠志\MSDS\Word 覆写模块\数据库与推断引擎\法规匹配库')
LAW_DIR = ROOT_DIR / "法规原文归档"
STD_DIR = ROOT_DIR / "标准原文归档"

TARGETS = [
    {"name": "中华人民共和国安全生产法", "type": "law", "sub": "法律", "query": "中华人民共和国安全生产法 2021 全文"},
    {"name": "中华人民共和国职业病防治法", "type": "law", "sub": "法律", "query": "中华人民共和国职业病防治法 2018最新 全文"},
    {"name": "中华人民共和国消防法", "type": "law", "sub": "法律", "query": "中华人民共和国消防法 2021修正 全文"},
    {"name": "中华人民共和国固体废物污染环境防治法", "type": "law", "sub": "法律", "query": "中华人民共和国固体废物污染环境防治法 2020修订 全文"},
    {"name": "易制毒化学品管理条例", "type": "law", "sub": "行政法规", "query": "易制毒化学品管理条例 国务院令第445号 全文"},
    {"name": "使用有毒物品作业场所劳动保护条例", "type": "law", "sub": "行政法规", "query": "使用有毒物品作业场所劳动保护条例 国务院令第352号 全文"},
    {"name": "易制爆危险化学品名录", "type": "law", "sub": "目录公告", "query": "易制爆危险化学品名录 2017年版 公安部 全文"},
    {"name": "重点监管的危险化学品名录", "type": "law", "sub": "目录公告", "query": "首批和第二批重点监管的危险化学品名录 完整版"},
    {"name": "GB 15258-2009 化学品安全标签编写规定", "type": "std", "sub": "SDS框架", "query": "GB 15258-2009 化学品安全标签编写规定 全文"},
    {"name": "GB 13690-2009 化学品分类和危险性公示 通则", "type": "std", "sub": "SDS框架", "query": "GB 13690-2009 化学品分类和危险性公示 通则 全文"},
    {"name": "GB 15603-2022 危险化学品仓库储存通则", "type": "std", "sub": "储存消防", "query": "GB 15603-2022 危险化学品仓库储存通则 全文"},
    {"name": "GB 12158-2006 防止静电事故通用导则", "type": "std", "sub": "储存消防", "query": "GB 12158-2006 防止静电事故通用导则 全文"},
    {"name": "GB 50016-2014(2018年版) 建筑设计防火规范", "type": "std", "sub": "储存消防", "query": "GB 50016-2014(2018年版) 建筑设计防火规范 仓库火灾危险性 全文"},
    {"name": "GB 50140-2005 建筑灭火器配置设计规范", "type": "std", "sub": "储存消防", "query": "GB 50140-2005 建筑灭火器配置设计规范 全文"},
    {"name": "WS 444-2014 化学品中毒医疗救治技术导则", "type": "std", "sub": "急救与职业卫生", "query": "WS 444-2014 化学品中毒医疗救治技术导则 全文"},
    {"name": "GBZ 71-2013 职业性急性化学物中毒诊断标准(总则)", "type": "std", "sub": "急救与职业卫生", "query": "GBZ 71-2013 职业性急性化学物中毒诊断标准(总则) 全文"},
    {"name": "GBZ 2.2-2007 工作场所有害因素职业接触限值 第2部分：物理因素", "type": "std", "sub": "急救与职业卫生", "query": "GBZ 2.2-2007 物理因素 全文"},
    {"name": "GB 2890-2009 呼吸防护 自吸过滤式防毒面具", "type": "std", "sub": "职业卫生", "query": "GB 2890-2009 呼吸防护 自吸过滤式防毒面具 全文"},
    {"name": "GB 2626-2019 呼吸防护 自吸过滤式防颗粒物呼吸器", "type": "std", "sub": "职业卫生", "query": "GB 2626-2019 呼吸防护 自吸过滤式防颗粒物呼吸器 全文"},
    {"name": "GB-T 18664-2002 呼吸防护用品的选择、使用与维护", "type": "std", "sub": "职业卫生", "query": "GB/T 18664-2002 呼吸防护用品的选择、使用与维护 全文"},
    {"name": "GB 50483-2019 化工建设项目环境保护设计标准", "type": "std", "sub": "环保与危废", "query": "GB 50483-2019 化工建设项目环境保护设计标准 全文"},
    {"name": "GB 18597-2023 危险废物贮存污染控制标准", "type": "std", "sub": "环保与危废", "query": "GB 18597-2023 危险废物贮存污染控制标准 全文"},
    {"name": "GB 18484-2020 危险废物焚烧污染控制标准", "type": "std", "sub": "环保与危废", "query": "GB 18484-2020 危险废物焚烧污染控制标准 全文"},
    {"name": "JT-T 617-2019 危险货物道路运输规则", "type": "std", "sub": "运输", "query": "JT/T 617-2019 危险货物道路运输规则 全文"},
    {"name": "GB-T 21844-2008 化合物 闪点测定 快速平衡闭杯法", "type": "std", "sub": "理化与毒理测试", "query": "GB/T 21844-2008 闪点测定 快速平衡闭杯法 全文"},
    {"name": "GB-T 21853-2008 化学品 沸点测定", "type": "std", "sub": "理化与毒理测试", "query": "GB/T 21853-2008 沸点测定 全文"},
    {"name": "GB-T 21845-2008 化学品 水溶解度测定", "type": "std", "sub": "理化与毒理测试", "query": "GB/T 21845-2008 水溶解度测定 全文"},
    {"name": "GB-T 21578-2008 危险品 反应性试验方法", "type": "std", "sub": "理化与毒理测试", "query": "GB/T 21578-2008 危险品 反应性试验方法 全文"},
    {"name": "GB-T 21807-2008 化学品 鱼类急性毒性试验", "type": "std", "sub": "生态毒理测试", "query": "GB/T 21807-2008 鱼类急性毒性试验 全文"},
    {"name": "GB-T 21809-2008 化学品 溞类急性活动抑制试验", "type": "std", "sub": "生态毒理测试", "query": "GB/T 21809-2008 溞类急性活动抑制试验 全文"},
    {"name": "GB-T 21805-2008 化学品 淡水藻生长抑制试验", "type": "std", "sub": "生态毒理测试", "query": "GB/T 21805-2008 淡水藻生长抑制试验 全文"}
]

def save_docx_and_txt(out_dir, name, content):
    out_dir.mkdir(parents=True, exist_ok=True)
    clean_name = re.sub(r'[\\/:*?"<>|]', "_", name)
    txt_file = out_dir / f"{clean_name}.txt"
    txt_file.write_text(content, encoding="utf-8")
    
    docx_file = out_dir / f"{clean_name}.docx"
    doc = docx.Document()
    doc.add_heading(name, level=0)
    for p in content.splitlines():
        p_str = p.strip()
        if p_str:
            doc.add_paragraph(p_str)
    doc.save(str(docx_file))
    print(f"  ✅ [归档成功] {docx_file.name} ({len(content)} 字符)")

def fetch_target(context, item):
    name = item["name"]
    stype = item["type"]
    sub = item["sub"]
    query = item["query"]
    out_dir = (LAW_DIR if stype == "law" else STD_DIR) / sub
    
    clean_name = re.sub(r'[\\/:*?"<>|]', "_", name)
    existing = list(out_dir.glob(f"{clean_name[:12]}*.docx"))
    if existing and existing[0].stat().st_size > 10000:
        print(f"⏭ [已完整归档] {name}")
        return True

    print(f"\n🚀 [获取中] {name} ({sub})...")
    page = context.new_page()
    try:
        search_url = f"https://www.baidu.com/s?wd={urllib.parse.quote(query)}"
        page.goto(search_url, timeout=25000)
        page.wait_for_timeout(2000)
        
        results = page.evaluate('''() => {
            const rows = document.querySelectorAll('.result, .c-container');
            return Array.from(rows).map(r => {
                const h3 = r.querySelector('h3');
                const a = r.querySelector('a');
                return {
                    title: h3 ? h3.innerText.trim() : '',
                    href: a ? a.href : ''
                };
            }).filter(x => x.title && x.href);
        }''')
        
        best_link = None
        for r in results:
            t = r["title"]
            # 优先选择包含全文、修正、条例的详情页
            if "百度" in t and "文库" in t: continue
            if "安全生产法" in t or "防治法" in t or "消防法" in t or "条例" in t or "名录" in t or "GB" in t or "标准" in t or "规范" in t:
                best_link = r["href"]
                print(f"  命中条目: {t[:40]}")
                break
                
        if not best_link and results:
            best_link = results[0]["href"]

        if best_link:
            page.goto(best_link, timeout=30000)
            page.wait_for_timeout(2000)
            
            # 提取主要文本
            text = page.evaluate('''() => {
                const el = document.querySelector('article, main, #article, .article-content, #content, .content, .main-content, #view_content, .law-content, .detail-content, .text, body');
                return el ? el.innerText : document.body.innerText;
            }''')
            
            # 过滤清洗
            start = text.find('目　录')
            if start < 0: start = text.find('第一章')
            if start < 0: start = text.find('第一条')
            if start < 0: start = text.find('1 范围')
            if start < 0: start = text.find('前言')
            if start >= 0:
                text = text[start:]
                
            if len(text) > 800:
                save_docx_and_txt(out_dir, name, text)
                return True
            else:
                print(f"  ⚠️ 页面抓取文本较短 ({len(text)} 字符)")
        else:
            print(f"  ❌ 未检索到合适链接: {name}")
    except Exception as e:
        print(f"  ❌ 抓取异常: {e}")
    finally:
        page.close()
    return False

def main():
    with sync_playwright() as p:
        b = p.chromium.launch(headless=True, executable_path=CHROME, args=['--no-sandbox'])
        context = b.new_context(user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/126.0.0.0 Safari/537.36")
        success = 0
        for item in TARGETS:
            if fetch_target(context, item):
                success += 1
            time.sleep(1)
        b.close()
        print(f"\n=============================\n🎉 任务完成! 成功归档/校验: {success} / {len(TARGETS)}\n=============================")

if __name__ == '__main__':
    main()
