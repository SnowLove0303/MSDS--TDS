# -*- coding: utf-8 -*-
"""一键批量抓取并清洗所有缺失的国家法律、行政法规、部委名录全文并生成 docx 与 txt."""
import sys, os, time, re
from pathlib import Path
from playwright.sync_api import sync_playwright
import docx

CHROME = r'E:\MorenAnzhuangLujing\Chrome\Chrome\Application\chrome.exe'
ROOT_DIR = Path(r'F:\正式项目与模块化内容\冠志\MSDS\Word 覆写模块\数据库与推断引擎\法规匹配库\法规原文归档')

TARGET_LAWS = [
    {
        "name": "中华人民共和国固体废物污染环境防治法",
        "sub": "法律",
        "url": "https://www.mee.gov.cn/ywgz/fgbz/fl/202004/t20200430_777580.shtml",
        "type": "html"
    },
    {
        "name": "中华人民共和国安全生产法",
        "sub": "法律",
        "url": "https://flk.npc.gov.cn/detail?id=ff8081817a66b816017a7956b7db0ad4&fileId=&type=&title=%E4%B8%AD%E5%8D%8E%E4%BA%BA%E6%B0%91%E5%85%B1%E5%92%8C%E5%9B%BD%E5%AE%89%E5%85%A8%E7%94%9F%E4%BA%A7%E6%B3%95",
        "type": "flk_ofd"
    },
    {
        "name": "中华人民共和国消防法",
        "sub": "法律",
        "url": "https://flk.npc.gov.cn/detail?id=ff8081817ab22e0c017abd909312060a&fileId=&type=&title=%E4%B8%AD%E5%8D%8E%E4%BA%BA%E6%B0%91%E5%85%B1%E5%92%8C%E5%9B%BD%E6%B6%88%E9%98%B2%E6%B3%95",
        "type": "flk_ofd"
    },
    {
        "name": "中华人民共和国职业病防治法",
        "sub": "法律",
        "url": "https://flk.npc.gov.cn/detail?id=ff8080816f135f46016f1a8e1b6f0e4b&fileId=&type=&title=%E4%B8%AD%E5%8D%8E%E4%BA%BA%E6%B0%91%E5%85%B1%E5%92%8C%E5%9B%BD%E8%81%8C%E4%B8%9A%E7%97%85%E9%98%B2%E6%B2%BB%E6%B3%95",
        "type": "flk_ofd"
    },
    {
        "name": "易制毒化学品管理条例",
        "sub": "行政法规",
        "url": "https://flk.npc.gov.cn/detail?id=ff8080816f3cbb3c016f40edff6f0c05&fileId=&type=&title=%E6%98%93%E5%88%B6%E6%AF%92%E5%8C%96%E5%AD%A6%E5%93%81%E7%AE%A1%E7%90%86%E6%9D%A1%E4%BE%8B",
        "type": "flk_ofd"
    },
    {
        "name": "使用有毒物品作业场所劳动保护条例",
        "sub": "行政法规",
        "url": "https://flk.npc.gov.cn/detail?id=ff8080816f3cbb3c016f40954b080517&fileId=&type=&title=%E4%BD%BF%E7%94%A8%E6%9C%89%E6%AF%92%E7%89%A9%E5%93%81%E4%BD%9C%E4%B8%9A%E5%9C%BA%E6%89%80%E5%8A%B3%E5%8A%A8%E4%BF%9D%E6%8A%A4%E6%9D%A1%E4%BE%8B",
        "type": "flk_ofd"
    }
]

def save_doc_and_txt(name, text, out_dir):
    out_dir.mkdir(parents=True, exist_ok=True)
    # 保存 txt
    txt_path = out_dir / f"{name}.txt"
    txt_path.write_text(text, encoding="utf-8")
    
    # 保存 docx
    docx_path = out_dir / f"{name}.docx"
    doc = docx.Document()
    doc.add_heading(name, level=0)
    for p in text.splitlines():
        p_clean = p.strip()
        if p_clean:
            doc.add_paragraph(p_clean)
    doc.save(str(docx_path))
    print(f"  ✅ 成功保存: {docx_path.name} ({len(text)} 字)")

def extract_flk_ofd(page, url):
    page.goto(url, timeout=40000)
    page.wait_for_timeout(3500)
    
    frames = [f for f in page.frames if 'flkofd' in f.url]
    if not frames:
        return None
    f = frames[0]
    # 模拟滚动到底部触发全量 OFD 渲染
    for _ in range(50):
        f.evaluate('() => window.scrollBy(0, 1500)')
        time.sleep(0.15)
        
    text = f.evaluate('''() => {
        return document.body.innerText;
    }''')
    return text

def extract_html(page, url):
    page.goto(url, timeout=30000)
    page.wait_for_timeout(2000)
    text = page.evaluate('''() => {
        const el = document.querySelector('.content_body, #main_content, .article-content, #content, .content, main, body');
        return el ? el.innerText : document.body.innerText;
    }''')
    # 截取正文
    start = text.find('目　录')
    if start < 0: start = text.find('第一章')
    if start < 0: start = text.find('第一条')
    if start >= 0:
        return text[start:]
    return text

def main():
    with sync_playwright() as p:
        b = p.chromium.launch(headless=True, executable_path=CHROME, args=['--no-sandbox'])
        context = b.new_context(user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36")
        page = context.new_page()
        
        for item in TARGET_LAWS:
            name = item["name"]
            sub = item["sub"]
            url = item["url"]
            ltype = item["type"]
            out_dir = ROOT_DIR / sub
            
            print(f"\n[抓取] {name} ({sub})...")
            try:
                if ltype == "flk_ofd":
                    content = extract_flk_ofd(page, url)
                else:
                    content = extract_html(page, url)
                    
                if content and len(content) > 300:
                    save_doc_and_txt(name, content, out_dir)
                else:
                    print(f"  ❌ 抓取内容过短: {len(content) if content else 0} 字")
            except Exception as e:
                print(f"  ❌ 抓取异常: {e}")
            time.sleep(1)
        b.close()

if __name__ == '__main__':
    main()
