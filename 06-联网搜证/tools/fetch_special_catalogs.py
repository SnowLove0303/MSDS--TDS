# -*- coding: utf-8 -*-
from playwright.sync_api import sync_playwright
import docx, time
from pathlib import Path

CHROME = r'E:\MorenAnzhuangLujing\Chrome\Chrome\Application\chrome.exe'
ROOT_OUT = Path(r"F:\正式项目与模块化内容\冠志\MSDS\Word 覆写模块\数据库与推断引擎\法规匹配库\法规原文归档\目录公告")

TARGETS = [
    {
        "name": "易制爆危险化学品名录（2017年版）",
        "url": "https://www.mps.gov.cn/n2254314/n6409334/c6475685/content.html"
    },
    {
        "name": "首批重点监管的危险化学品名录",
        "url": "https://www.mem.gov.cn/gk/gwgg/aqsc/wxhxp/201106/t20110621_237599.shtml"
    },
    {
        "name": "第二批重点监管的危险化学品名录",
        "url": "https://www.mem.gov.cn/gk/gwgg/aqsc/wxhxp/201302/t20130206_237672.shtml"
    }
]

def main():
    with sync_playwright() as p:
        b = p.chromium.launch(headless=True, executable_path=CHROME, args=['--no-sandbox'])
        context = b.new_context()
        page = context.new_page()
        
        for it in TARGETS:
            name = it["name"]
            url = it["url"]
            print(f"\n[Crawling Notice] {name} -> {url}")
            try:
                page.goto(url, timeout=30000)
                page.wait_for_timeout(2500)
                
                text = page.evaluate('''() => {
                    const main = document.querySelector('.article-content, #main_content, #content, .content, body');
                    return main ? main.innerText : document.body.innerText;
                }''')
                
                ROOT_OUT.mkdir(parents=True, exist_ok=True)
                docx_path = ROOT_OUT / f"{name}.docx"
                txt_path = ROOT_OUT / f"{name}.txt"
                
                txt_path.write_text(text, encoding="utf-8")
                doc = docx.Document()
                doc.add_heading(name, level=0)
                for line in text.splitlines():
                    ls = line.strip()
                    if ls:
                        doc.add_paragraph(ls)
                doc.save(str(docx_path))
                print(f"  ✅ 成功归档: {docx_path.name} ({len(text)} 字符)")
            except Exception as e:
                print(f"  ❌ 抓取异常: {e}")
        b.close()

if __name__ == '__main__':
    main()
