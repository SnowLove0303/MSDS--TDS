# -*- coding: utf-8 -*-
"""全量获取 MSDS 推断引擎法规标准库中缺失的法律、行政法规、部委名录及国家/行业标准."""
import sys, os, time, re, json
from pathlib import Path
from playwright.sync_api import sync_playwright
import docx
import requests

CHROME = r'E:\MorenAnzhuangLujing\Chrome\Chrome\Application\chrome.exe'
ROOT_DIR = Path(r'F:\正式项目与模块化内容\冠志\MSDS\Word 覆写模块\数据库与推断引擎\法规匹配库')
LAW_ROOT = ROOT_DIR / "法规原文归档"
STD_ROOT = ROOT_DIR / "标准原文归档"

# 需要获取的法规与标准清单
ITEMS = [
    # 1. 核心法律
    {"name": "中华人民共和国安全生产法", "sub": "法律", "type": "law", "query": "中华人民共和国安全生产法 2021全文 site:mem.gov.cn OR site:gov.cn"},
    {"name": "中华人民共和国职业病防治法", "sub": "法律", "type": "law", "query": "中华人民共和国职业病防治法 2018全文 site:gov.cn OR site:nhc.gov.cn"},
    {"name": "中华人民共和国消防法", "sub": "法律", "type": "law", "query": "中华人民共和国消防法 2021全文 site:gov.cn OR site:mem.gov.cn"},
    {"name": "中华人民共和国固体废物污染环境防治法", "sub": "法律", "type": "law", "query": "中华人民共和国固体废物污染环境防治法 2020全文 site:mee.gov.cn OR site:gov.cn"},
    
    # 2. 行政法规与目录
    {"name": "易制毒化学品管理条例", "sub": "行政法规", "type": "law", "query": "易制毒化学品管理条例 国务院令第445号 site:gov.cn"},
    {"name": "使用有毒物品作业场所劳动保护条例", "sub": "行政法规", "type": "law", "query": "使用有毒物品作业场所劳动保护条例 国务院令第352号 site:gov.cn"},
    {"name": "易制爆危险化学品名录", "sub": "目录公告", "type": "law", "query": "易制爆危险化学品名录 2017年版 公安部 site:gov.cn OR site:mps.gov.cn"},
    {"name": "重点监管的危险化学品名录", "sub": "目录公告", "type": "law", "query": "首批和第二批重点监管的危险化学品名录 site:mem.gov.cn OR site:gov.cn"},

    # 3. 关键国家标准与行业标准
    {"name": "GB 15258-2009 化学品安全标签编写规定", "sub": "SDS框架", "type": "std", "query": "GB 15258-2009 化学品安全标签编写规定 pdf download OR 全文"},
    {"name": "GB 13690-2009 化学品分类和危险性公示 通则", "sub": "SDS框架", "type": "std", "query": "GB 13690-2009 化学品分类和危险性公示 通则 pdf"},
    {"name": "GB 15603-2022 危险化学品仓库储存通则", "sub": "储存消防", "type": "std", "query": "GB 15603-2022 危险化学品仓库储存通则 pdf"},
    {"name": "GB 12158-2006 防止静电事故通用导则", "sub": "储存消防", "type": "std", "query": "GB 12158-2006 防止静电事故通用导则 pdf"},
    {"name": "GB 50016-2014(2018年版) 建筑设计防火规范", "sub": "储存消防", "type": "std", "query": "GB 50016-2014(2018年版) 建筑设计防火规范 pdf"},
    {"name": "GB 50140-2005 建筑灭火器配置设计规范", "sub": "储存消防", "type": "std", "query": "GB 50140-2005 建筑灭火器配置设计规范 pdf"},
    {"name": "WS 444-2014 化学品中毒医疗救治技术导则", "sub": "急救与职业卫生", "type": "std", "query": "WS 444-2014 化学品中毒医疗救治技术导则 pdf"},
    {"name": "GBZ 71-2013 职业性急性化学物中毒诊断标准(总则)", "sub": "急救与职业卫生", "type": "std", "query": "GBZ 71-2013 职业性急性化学物中毒诊断标准 pdf"},
    {"name": "GBZ 2.2-2007 工作场所有害因素职业接触限值 第2部分：物理因素", "sub": "急救与职业卫生", "type": "std", "query": "GBZ 2.2-2007 物理因素 pdf"},
    {"name": "GB 2890-2009 呼吸防护 自吸过滤式防毒面具", "sub": "职业卫生", "type": "std", "query": "GB 2890-2009 呼吸防护 自吸过滤式防毒面具 pdf"},
    {"name": "GB 2626-2019 呼吸防护 自吸过滤式防颗粒物呼吸器", "sub": "职业卫生", "type": "std", "query": "GB 2626-2019 呼吸防护 自吸过滤式防颗粒物呼吸器 pdf"},
    {"name": "GB-T 18664-2002 呼吸防护用品的选择、使用与维护", "sub": "职业卫生", "type": "std", "query": "GB/T 18664-2002 呼吸防护用品的选择、使用与维护 pdf"},
    {"name": "GB 50483-2019 化工建设项目环境保护设计标准", "sub": "环保与危废", "type": "std", "query": "GB 50483-2019 化工建设项目环境保护设计标准 pdf"},
    {"name": "GB 18597-2023 危险废物贮存污染控制标准", "sub": "环保与危废", "type": "std", "query": "GB 18597-2023 危险废物贮存污染控制标准 pdf site:mee.gov.cn OR pdf"},
    {"name": "GB 18484-2020 危险废物焚烧污染控制标准", "sub": "环保与危废", "type": "std", "query": "GB 18484-2020 危险废物焚烧污染控制标准 pdf site:mee.gov.cn OR pdf"},
    {"name": "JT-T 617-2019 危险货物道路运输规则", "sub": "运输", "type": "std", "query": "JT/T 617-2019 危险货物道路运输规则 pdf"},
    {"name": "GB-T 21844-2008 化合物 闪点测定 快速平衡闭杯法", "sub": "理化与毒理测试", "type": "std", "query": "GB/T 21844-2008 闪点测定 快速平衡闭杯法 pdf"},
    {"name": "GB-T 21853-2008 化学品 沸点测定", "sub": "理化与毒理测试", "type": "std", "query": "GB/T 21853-2008 沸点测定 pdf"},
    {"name": "GB-T 21845-2008 化学品 水溶解度测定", "sub": "理化与毒理测试", "type": "std", "query": "GB/T 21845-2008 水溶解度测定 pdf"},
    {"name": "GB-T 21578-2008 危险品 反应性试验方法", "sub": "理化与毒理测试", "type": "std", "query": "GB/T 21578-2008 危险品 反应性试验方法 pdf"},
    {"name": "GB-T 21807-2008 化学品 鱼类急性毒性试验", "sub": "生态毒理测试", "type": "std", "query": "GB/T 21807-2008 鱼类急性毒性试验 pdf"},
    {"name": "GB-T 21809-2008 化学品 溞类急性活动抑制试验", "sub": "生态毒理测试", "type": "std", "query": "GB/T 21809-2008 溞类急性活动抑制试验 pdf"},
    {"name": "GB-T 21805-2008 化学品 淡水藻生长抑制试验", "sub": "生态毒理测试", "type": "std", "query": "GB/T 21805-2008 淡水藻生长抑制试验 pdf"},
]

def save_as_docx(title, content, out_docx_path):
    doc = docx.Document()
    doc.add_heading(title, level=0)
    for line in content.splitlines():
        l = line.strip()
        if l:
            doc.add_paragraph(l)
    doc.save(str(out_docx_path))

def fetch_item(context, item):
    name = item["name"]
    stype = item["type"]
    sub = item["sub"]
    query = item["query"]
    
    out_dir = (LAW_ROOT if stype == "law" else STD_ROOT) / sub
    out_dir.mkdir(parents=True, exist_ok=True)
    
    clean_fname = re.sub(r'[\\/:*?"<>|]', "_", name)
    existing_files = list(out_dir.glob(f"{clean_fname[:15]}*"))
    if existing_files:
        print(f"⏭ [已存在] {name} -> {existing_files[0].name}")
        return True

    print(f"\n🚀 [正在获取] {name} ({sub})...")
    page = context.new_page()
    try:
        search_url = f"https://www.bing.com/search?q={query}"
        page.goto(search_url, timeout=25000)
        page.wait_for_timeout(2000)
        
        # 寻找结果链接
        links = page.evaluate('''() => {
            return Array.from(document.querySelectorAll('a')).map(a => ({
                text: a.innerText.trim(),
                href: a.href
            })).filter(x => x.text.length > 2 && x.href.startsWith('http'));
        }''')
        
        target_href = None
        for l in links:
            t = l["text"]
            h = l["href"]
            if "bing.com" in h or "microsoft.com" in h or "javascript" in h:
                continue
            # 优先匹配标题中包含关键词的
            if any(k in t for k in [name[:6], "全文", "标准", "法规", "下载", "PDF"]):
                target_href = h
                print(f"  命中链接: {t[:40]} -> {h[:80]}")
                break
                
        if not target_href and len(links) > 0:
            for l in links:
                if "bing.com" not in l["href"]:
                    target_href = l["href"]
                    break

        if target_href:
            # 访问目标页
            try:
                page.goto(target_href, timeout=30000, wait_until="load")
                page.wait_for_timeout(2000)
            except Exception as e:
                print(f"  页面跳转提示: {e}")
                
            final_url = page.url
            print(f"  当前落地页: {final_url}")
            
            # 检查是否有直接附件下载（.pdf / .docx）
            pdf_url = page.evaluate('''() => {
                const as = Array.from(document.querySelectorAll('a'));
                for (const a of as) {
                    const href = a.href || '';
                    if (href.endsWith('.pdf') || href.endsWith('.docx') || href.endsWith('.doc')) {
                        return href;
                    }
                }
                return null;
            }''')
            
            downloaded = False
            if pdf_url and pdf_url.startswith("http"):
                print(f"  发现直接文件下载: {pdf_url}")
                try:
                    r = requests.get(pdf_url, headers={"User-Agent": "Mozilla/5.0"}, timeout=30)
                    if len(r.content) > 1000:
                        ext = ".pdf" if pdf_url.endswith(".pdf") else ".docx"
                        dest = out_dir / f"{clean_fname}{ext}"
                        dest.write_bytes(r.content)
                        print(f"  ✅ [附件下载成功] {dest.name} ({len(r.content)} 字节)")
                        downloaded = True
                except Exception as ex:
                    print(f"  附件下载异常: {ex}")
                    
            if not downloaded:
                # 抓取页面全文正文
                body_text = page.evaluate('''() => {
                    const el = document.querySelector('article, main, #content, .content, .article-content, .detail-content, .law-content, body');
                    return el ? el.innerText : document.body.innerText;
                }''')
                
                if len(body_text) > 200:
                    docx_file = out_dir / f"{clean_fname}.docx"
                    txt_file = out_dir / f"{clean_fname}.txt"
                    save_as_docx(name, body_text, docx_file)
                    txt_file.write_text(body_text, encoding="utf-8")
                    print(f"  ✅ [正文解析保存成功] {docx_file.name} ({len(body_text)} 字)")
                    downloaded = True
                else:
                    print(f"  ❌ 抓取内容过短 ({len(body_text)} 字)")
        else:
            print(f"  ❌ 未找到有效链接: {name}")
    except Exception as e:
        print(f"  ❌ 处理发生异常: {e}")
    finally:
        page.close()

def main():
    with sync_playwright() as p:
        b = p.chromium.launch(headless=True, executable_path=CHROME, args=['--no-sandbox'])
        context = b.new_context(user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/126.0.0.0 Safari/537.36")
        for item in ITEMS:
            fetch_item(context, item)
            time.sleep(1)
        b.close()

if __name__ == '__main__':
    main()
