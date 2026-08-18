# -*- coding: utf-8 -*-
from playwright.sync_api import sync_playwright
import docx, re
from pathlib import Path

CHROME = r'E:\MorenAnzhuangLujing\Chrome\Chrome\Application\chrome.exe'
ROOT_OUT = Path(r"F:\正式项目与模块化内容\冠志\MSDS\Word 覆写模块\数据库与推断引擎\法规匹配库\法规原文归档\目录公告")

TARGETS = [
    {
        "name": "易制爆危险化学品名录（2017年版）",
        "url": "https://www.so.com/link?m=eVdLevyhQq7qZBk1yEXL%2BrJHhaF8Apf6nWaipTdUXkH6ukZKsBzJGYRCISez9YdC8yBvraz%2BLVwC0WTgVg%2BNnJkKMrHaq88tIJ53Bj2IA%2BXyo6JIuxlTtCpzC5wMLSikSiAfzoY49EEc3%2Bt1crUJZ0hGjmP3XWLHcaGkuUYd8m%2BH%2FDY2k2lC6Z0eXOxuP%2BjdQRZ4Ert%2F99ACZLalctzlJMGIhasIGmEvhQvYSffFgRJYRvEAMog7GfOhPGm2kknjYbUF%2BD2A7gXL89lkMeiJO4s4TcWTfd%2BXnt7uM9jQRaiqq4bahmA7CKc0Lcvc%3D"
    }
]

def main():
    with sync_playwright() as p:
        b = p.chromium.launch(headless=True, executable_path=CHROME, args=['--no-sandbox'])
        context = b.new_context()
        page = context.new_page()
        
        for it in TARGETS:
            name = it["name"]
            url = it["url"]
            print(f"\n[Crawling] {name} -> {url}")
            page.goto(url, timeout=30000)
            page.wait_for_timeout(2000)
            
            text = page.evaluate('''() => {
                const el = document.querySelector('article, #content, .content, .main_content, .article-content, body');
                return el ? el.innerText : '';
            }''')
            print(f"  Extracted text len: {len(text)}")
            if len(text) > 400:
                txt_path = ROOT_OUT / f"{name}.txt"
                docx_path = ROOT_OUT / f"{name}.docx"
                txt_path.write_text(text, encoding="utf-8")
                
                doc = docx.Document()
                doc.add_heading(name, level=0)
                for line in text.splitlines():
                    ls = line.strip()
                    if ls:
                        doc.add_paragraph(ls)
                doc.save(str(docx_path))
                print(f"  ✅ 成功保存: {docx_path.name}")
        b.close()

if __name__ == '__main__':
    main()
