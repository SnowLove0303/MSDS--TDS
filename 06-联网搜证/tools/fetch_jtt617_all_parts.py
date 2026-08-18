# -*- coding: utf-8 -*-
from playwright.sync_api import sync_playwright
import docx, time, requests
from pathlib import Path

CHROME = r'E:\MorenAnzhuangLujing\Chrome\Chrome\Application\chrome.exe'
ROOT_OUT = Path(r"F:\正式项目与模块化内容\冠志\MSDS\Word 覆写模块\数据库与推断引擎\法规匹配库\标准原文归档\运输")

with sync_playwright() as p:
    b = p.chromium.launch(headless=True, executable_path=CHROME, args=['--no-sandbox'])
    context = b.new_context()
    page = context.new_page()
    page.goto("https://hbba.sacinfo.org.cn/stdList?key=JT/T%20617", timeout=30000)
    page.wait_for_timeout(2000)
    
    links = page.evaluate('''() => {
        return Array.from(document.querySelectorAll('a')).map(a => ({
            text: (a.innerText || '').trim(),
            href: a.href
        })).filter(x => x.text.includes('JT/T 617') || x.text.includes('危险货物道路运输规则'));
    }''')
    print(f"Found {len(links)} JT/T 617 parts:")
    
    for l in links:
        part_name = l["text"]
        part_url = l["href"]
        print(f"\n[Crawling] {part_name} -> {part_url}")
        
        page.goto(part_url, timeout=30000)
        page.wait_for_timeout(2000)
        
        detail_text = page.evaluate('''() => {
            const main = document.querySelector('.main, .content, #content, .container, body');
            return main ? main.innerText : document.body.innerText;
        }''')
        
        ROOT_OUT.mkdir(parents=True, exist_ok=True)
        docx_path = ROOT_OUT / f"{part_name}.docx"
        txt_path = ROOT_OUT / f"{part_name}.txt"
        
        txt_path.write_text(detail_text, encoding="utf-8")
        doc = docx.Document()
        doc.add_heading(part_name, level=0)
        for line in detail_text.splitlines():
            ls = line.strip()
            if ls:
                doc.add_paragraph(ls)
        doc.save(str(docx_path))
        print(f"  ✅ 成功归档备案信息: {docx_path.name}")
        time.sleep(1)
        
    b.close()
