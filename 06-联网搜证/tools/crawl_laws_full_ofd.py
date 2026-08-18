# -*- coding: utf-8 -*-
from playwright.sync_api import sync_playwright
import docx, time
from pathlib import Path

CHROME = r'E:\MorenAnzhuangLujing\Chrome\Chrome\Application\chrome.exe'

TARGETS = [
    {
        "name": "中华人民共和国安全生产法",
        "sub": "法律",
        "url": "https://flk.npc.gov.cn/detail?id=ff8081817a66b816017a7956b7db0ad4&fileId=&type=&title=%E4%B8%AD%E5%8D%8E%E4%BA%BA%E6%B0%91%E5%85%B1%E5%92%8C%E5%9B%BD%E5%AE%89%E5%85%A8%E7%94%9F%E4%BA%A7%E6%B3%95"
    },
    {
        "name": "中华人民共和国消防法",
        "sub": "法律",
        "url": "https://flk.npc.gov.cn/detail?id=ff8081817ab22e0c017abd909312060a&fileId=&type=&title=%E4%B8%AD%E5%8D%8E%E4%BA%BA%E6%B0%91%E5%85%B1%E5%92%8C%E5%9B%BD%E6%B6%88%E9%98%B2%E6%B3%95"
    },
    {
        "name": "中华人民共和国职业病防治法",
        "sub": "法律",
        "url": "https://flk.npc.gov.cn/detail?id=ff8080816f135f46016f1a8e1b6f0e4b&fileId=&type=&title=%E4%B8%AD%E5%8D%8E%E4%BA%BA%E6%B0%91%E5%85%B1%E5%92%8C%E5%9B%BD%E8%81%8C%E4%B8%9A%E7%97%85%E9%98%B2%E6%B2%BB%E6%B3%95"
    },
    {
        "name": "易制毒化学品管理条例",
        "sub": "行政法规",
        "url": "https://flk.npc.gov.cn/detail?id=ff8080816f3cbb3c016f40edff6f0c05&fileId=&type=&title=%E6%98%93%E5%88%B6%E6%AF%92%E5%8C%96%E5%AD%A6%E5%93%81%E7%AE%A1%E7%90%86%E6%9D%A1%E4%BE%8B"
    },
    {
        "name": "使用有毒物品作业场所劳动保护条例",
        "sub": "行政法规",
        "url": "https://flk.npc.gov.cn/detail?id=ff8080816f3cbb3c016f40954b080517&fileId=&type=&title=%E4%BD%BF%E7%94%A8%E6%9C%89%E6%AF%92%E7%89%A9%E5%93%81%E4%BD%9C%E4%B8%9A%E5%9C%BA%E6%89%80%E5%8A%B3%E5%8A%A8%E4%BF%9D%E6%8A%A4%E6%9D%A1%E4%BE%8B"
    }
]

ROOT_OUT = Path(r"F:\正式项目与模块化内容\冠志\MSDS\Word 覆写模块\数据库与推断引擎\法规匹配库\法规原文归档")

def crawl_law(page, item):
    name = item["name"]
    sub = item["sub"]
    url = item["url"]
    out_dir = ROOT_OUT / sub
    out_dir.mkdir(parents=True, exist_ok=True)
    
    print(f"\n[Crawl] {name} ({sub})...")
    page.goto(url, timeout=30000)
    page.wait_for_timeout(3500)
    
    frames = [f for f in page.frames if 'flkofd' in f.url]
    if not frames:
        print("  ❌ 未找到 OFD frame")
        return
        
    f = frames[0]
    
    # 模拟在 frame 内部连续向下翻页与滚动，确保每一页 text-layer 都渲染
    for p_idx in range(60):
        f.evaluate('() => { window.scrollBy(0, 1000); const el = document.querySelector("#viewer, body"); if(el) el.scrollTop += 1000; }')
        time.sleep(0.1)
        
    # 获取全部文本
    text_lines = f.evaluate('''() => {
        const divs = document.querySelectorAll('div[class*="text"], span[class*="text"], span, div');
        const set = [];
        divs.forEach(d => {
            const t = d.innerText ? d.innerText.trim() : '';
            if (t && t.length > 0 && !set.includes(t)) {
                // 如果只是一两个字或纯数字页码过滤
            }
        });
        return document.body.innerText;
    }''')
    
    print(f"  Extracted text length: {len(text_lines)} characters")
    if len(text_lines) > 500:
        txt_path = out_dir / f"{name}.txt"
        txt_path.write_text(text_lines, encoding="utf-8")
        
        docx_path = out_dir / f"{name}.docx"
        doc = docx.Document()
        doc.add_heading(name, level=0)
        for line in text_lines.splitlines():
            l_s = line.strip()
            if l_s:
                doc.add_paragraph(l_s)
        doc.save(str(docx_path))
        print(f"  ✅ 成功保存: {docx_path.name} ({docx_path.stat().st_size} bytes)")
    else:
        print(f"  ❌ 抓取内容过少")

def main():
    with sync_playwright() as p:
        b = p.chromium.launch(headless=True, executable_path=CHROME, args=['--no-sandbox'])
        context = b.new_context(user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36")
        page = context.new_page()
        for it in TARGETS:
            crawl_law(page, it)
            time.sleep(1)
        b.close()

if __name__ == '__main__':
    main()
