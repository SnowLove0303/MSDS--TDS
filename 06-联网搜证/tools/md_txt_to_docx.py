
# -*- coding: utf-8 -*-
"""md/txt -> docx 批量转换，中文公文排版."""
import sys, re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# GHS 编号 -> 标准象形图 PNG 文件名（用于象形图库表格图片嵌入）
GHS_PIC = {
    'GHS01': 'GHS01_爆炸物.png', 'GHS02': 'GHS02_易燃.png',
    'GHS03': 'GHS03_氧化.png', 'GHS04': 'GHS04_加压气体.png',
    'GHS05': 'GHS05_腐蚀.png', 'GHS06': 'GHS06_急性毒性.png',
    'GHS07': 'GHS07_感叹号(警示).png', 'GHS08': 'GHS08_健康危害.png',
    'GHS09': 'GHS09_环境危害.png',
}
_EMOJI_RE = re.compile(r'[\U0001F000-\U0001FAFF☀-➿⬀-⯿️]')

def set_font(run, name_cn='宋体', size=12, bold=False, color=None):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name_cn)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')

def add_runs(p, text, name_cn, size, bold, color=None):
    parts = re.split(r'(\*\*.*?\*\*|`[^`]*`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2]); set_font(r, name_cn, size, True, color)
        elif part.startswith('`') and part.endswith('`'):
            r = p.add_run(part[1:-1]); set_font(r, '宋体', size-1, False, color or (0x33,0x33,0x33))
        else:
            r = p.add_run(part); set_font(r, name_cn, size, bold, color)

def add_para(doc, text, style='body'):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if style == 'h1':
        pf.space_before = Pt(18); pf.space_after = Pt(12)
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_runs(p, text, '黑体', 16, True)
    elif style == 'h2':
        pf.space_before = Pt(12); pf.space_after = Pt(8)
        add_runs(p, text, '黑体', 14, True)
    elif style == 'h3':
        pf.space_before = Pt(10); pf.space_after = Pt(6)
        add_runs(p, text, '黑体', 12, True)
    elif style == 'h4':
        pf.space_before = Pt(8); pf.space_after = Pt(4)
        add_runs(p, text, '黑体', 12, True)
    elif style == 'quote':
        pf.left_indent = Cm(0.74); pf.line_spacing = 1.4
        pf.space_before = Pt(3); pf.space_after = Pt(3)
        add_runs(p, text, '楷体', 10.5, False)
    elif style == 'code':
        pf.left_indent = Cm(0.74); pf.line_spacing = 1.15
        pf.space_before = Pt(3); pf.space_after = Pt(3)
        add_runs(p, text, '宋体', 9, False, (0x33,0x33,0x33))
    else:
        pf.first_line_indent = Pt(24)
        pf.line_spacing = 1.5
        add_runs(p, text, '宋体', 12, False)
    return p

def add_table(doc, rows, pictogram_dir=None):
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = [(rows[0][k].strip().strip('*') if k < len(rows[0]) else '')
               for k in range(n_cols)]
    for i, row in enumerate(rows):
        for j in range(n_cols):
            cell = table.cell(i, j)
            txt = row[j].strip() if j < len(row) else ''
            p = cell.paragraphs[0]
            if i == 0:
                add_runs(p, txt, '黑体', 10, True)
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), 'D9E2F3')
                cell._tc.get_or_add_tcPr().append(shd)
            elif pictogram_dir and '象形图' in headers[j]:
                # 象形图列：GHS 编号/emoji 描述 -> 真实 PNG
                m = re.search(r'(GHS\d{2})', txt)
                if m and m.group(1) in GHS_PIC:
                    keep = None if _EMOJI_RE.search(txt) else m.group(1)
                    _cell_pictogram(cell, pictogram_dir, m.group(1), keep=keep)
                elif _EMOJI_RE.search(txt):
                    # 表0：emoji 描述不含编号，从行内其他单元格定位 GHS 编号
                    gid = None
                    for cell_txt in row:
                        mm = re.search(r'(GHS\d{2})', cell_txt)
                        if mm and mm.group(1) in GHS_PIC:
                            gid = mm.group(1); break
                    if gid:
                        _cell_pictogram(cell, pictogram_dir, gid, keep=None)
                    else:
                        add_runs(p, txt, '宋体', 10, False)
                else:
                    add_runs(p, txt, '宋体', 10, False)
            else:
                add_runs(p, txt, '宋体', 10, False)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
    return table

def _cell_pictogram(cell, pictogram_dir, gid, keep):
    """在表格单元格内插入 GHS 象形图 PNG，可选在下方保留编号文字."""
    png = Path(pictogram_dir) / GHS_PIC.get(gid, '')
    p = cell.paragraphs[0]
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    if png.exists():
        run = p.add_run()
        run.add_picture(str(png), width=Cm(1.5))
    if keep:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(keep)
        set_font(r2, '宋体', 8, False, (0x40, 0x40, 0x40))
        p2.paragraph_format.space_after = Pt(2)
        p2.paragraph_format.space_before = Pt(1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def setup(doc):
    for sec in doc.sections:
        sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(2.6); sec.right_margin = Cm(2.6)

def md_to_docx(md_path, docx_path, pictogram_dir=None):
    lines = Path(md_path).read_text(encoding='utf-8').splitlines()
    doc = Document(); setup(doc)
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1; continue
        if line.strip().startswith('```'):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                add_para(doc, lines[i].rstrip(), 'code'); i += 1
            i += 1; continue
        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            add_para(doc, m.group(2).strip(), f'h{min(len(m.group(1)),4)}')
            i += 1; continue
        if line.strip().startswith('|') and i+1 < len(lines) and re.match(r'^\s*\|[\s:\-|]+\|\s*$', lines[i+1]):
            header = [c.strip() for c in line.strip().strip('|').split('|')]
            i += 2
            rows = [header]
            while i < len(lines) and lines[i].strip().startswith('|'):
                rows.append([c.strip() for c in lines[i].strip().strip('|').split('|')])
                i += 1
            add_table(doc, rows, pictogram_dir)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue
        if line.strip().startswith('>'):
            add_para(doc, line.strip().lstrip('>').strip(), 'quote')
            i += 1; continue
        m = re.match(r'^\s*[-*+]\s+(.*)', line)
        if m:
            add_para(doc, '• ' + m.group(1).strip(), 'body'); i += 1; continue
        m = re.match(r'^\s*\d+[.、]\s+(.*)', line)
        if m:
            add_para(doc, m.group(0).strip(), 'body'); i += 1; continue
        add_para(doc, line.strip(), 'body'); i += 1
    doc.save(docx_path)
    return len(lines)

def txt_to_docx(txt_path, docx_path):
    lines = Path(txt_path).read_text(encoding='utf-8').splitlines()
    doc = Document(); setup(doc)
    for line in lines:
        s = line.strip()
        if not s:
            continue
        if re.match(r'^第[一二三四五六七八九十百]+章', s) and len(s) < 30:
            add_para(doc, s, 'h3')
        else:
            add_para(doc, s, 'body')
    doc.save(docx_path)
    return len(lines)

if __name__ == '__main__':
    # 象形图库文档转换时传入标准象形图目录，表格中自动嵌入真实 PNG
    pic_dir = Path(__file__).resolve().parent.parent / '标准原文归档' / '象形图库' / '标准象形图'
    for fp in sys.argv[1:]:
        p = Path(fp)
        if not p.exists():
            print(f'MISS {fp}'); continue
        docx_p = p.with_suffix('.docx')
        try:
            if p.suffix.lower() == '.md':
                n = md_to_docx(p, docx_p, pictogram_dir=pic_dir if '象形图' in p.name else None)
            else:
                n = txt_to_docx(p, docx_p)
            print(f'OK {p.name} -> {docx_p.name} ({n} lines)')
        except Exception as e:
            print(f'ERR {p.name}: {e}')
