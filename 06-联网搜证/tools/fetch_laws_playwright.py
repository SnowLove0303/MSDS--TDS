# -*- coding: utf-8 -*-
"""从国家法律法规数据库 (flk.npc.gov.cn) 或 中国政府网 (gov.cn) 抓取核心法律法规 Word/PDF."""
import sys, json, time, re
from pathlib import Path
from playwright.sync_api import sync_playwright
import docx

CHROME = r'E:\MorenAnzhuangLujing\Chrome\Chrome\Application\chrome.exe'
OUT_DIR = Path(r'F:\正式项目与模块化内容\冠志\MSDS\Word 覆写模块\数据库与推断引擎\法规匹配库\法规原文归档')

LAWS_TO_FETCH = [
    {"name": "中华人民共和国安全生产法", "type": "法律", "year": "2021"},
    {"name": "中华人民共和国职业病防治法", "type": "法律", "year": "2018"},
    {"name": "中华人民共和国消防法", "type": "法律", "year": "2021"},
    {"name": "中华人民共和国固体废物污染环境防治法", "type": "法律", "year": "2020"},
    {"name": "易制毒化学品管理条例", "type": "行政法规", "year": ""},
    {"name": "使用有毒物品作业场所劳动保护条例", "type": "行政法规", "year": ""},
    {"name": "易制爆危险化学品名录", "type": "目录公告", "year": "2017"},
    {"name": "重点监管的危险化学品名录", "type": "目录公告", "year": ""},
]

def search_and_download_law(page, law_info):
    name = law_info["name"]
    ltype = law_info["type"]
    target_dir = OUT_DIR / ltype
    target_dir.mkdir(parents=True, exist_ok=True)
    
    print(f"\n[检索] {name} ({ltype})...")
    # 优先搜索 flk.npc.gov.cn
    search_url = f"https://flk.npc.gov.cn/search.html?search={name}"
    try:
        page.goto(search_url, timeout=30000, wait_until="networkidle")
        page.wait_for_timeout(2000)
    except Exception as e:
        print(f"  flk.npc.gov.cn 访问超时: {e}")
        # fallback to baidu
        baidu_url = f"https://www.baidu.com/s?wd={name}+site%3Agov.cn+OR+site%3Anpc.gov.cn"
        page.goto(baidu_url, timeout=30000, wait_until="networkidle")
        page.wait_for_timeout(1500)

    # 尝试找到匹配链接
    links = page.evaluate('''() => {
        return Array.from(document.querySelectorAll('a')).map(a => ({
            text: a.innerText.trim(),
            href: a.href
        })).filter(x => x.text.length > 2 && x.href.startsWith('http'));
    }''')
    
    target_link = None
    for link in links:
        if name in link["text"] or (name[:6] in link["text"] and "法" in link["text"]):
            target_link = link["href"]
            print(f"  找到详情链接: {link['text'][:40]} -> {target_link}")
            break
            
    if not target_link:
        # try bing
        bing_url = f"https://cn.bing.com/search?q={name}+最新+全文+site%3Agov.cn+OR+site%3Anpc.gov.cn"
        page.goto(bing_url, timeout=30000, wait_until="networkidle")
        page.wait_for_timeout(1500)
        links = page.evaluate('''() => {
            return Array.from(document.querySelectorAll('a')).map(a => ({
                text: a.innerText.trim(),
                href: a.href
            })).filter(x => x.text.length > 2 && x.href.startsWith('http'));
        }''')
        for link in links:
            if name in link["text"]:
                target_link = link["href"]
                print(f"  Bing 找到详情链接: {link['text'][:40]} -> {target_link}")
                break

    if target_link:
        page.goto(target_link, timeout=30000, wait_until="networkidle")
        page.wait_for_timeout(2000)
        
        # 尝试寻找 word / pdf 下载按钮
        dl_link = page.evaluate('''() => {
            const btns = Array.from(document.querySelectorAll('a, button'));
            for (const b of btns) {
                const t = (b.innerText || '') + (b.getAttribute('title') || '') + (b.getAttribute('href') || '');
                if (t.includes('Word') || t.includes('DOCX') || t.includes('docx') || t.includes('.docx') || t.includes('PDF') || t.includes('pdf') || t.includes('下载')) {
                    if (b.href && b.href.startsWith('http')) return b.href;
                }
            }
            return null;
        }''')
        
        # 提取页面正文
        content_text = page.evaluate('''() => {
            const article = document.querySelector('#content, .content, .article-content, #view_content, .detail-content, .law-content, article, main, .main-content');
            if (article) return article.innerText;
            return document.body.innerText;
        }''')
        
        # 保存为 docx 和 txt
        docx_path = target_dir / f"{name}.docx"
        doc = docx.Document()
        doc.add_heading(name, level=0)
        for line in content_text.splitlines():
            line_s = line.strip()
            if line_s:
                doc.add_paragraph(line_s)
        doc.save(str(docx_path))
        
        txt_path = target_dir / f"{name}.txt"
        txt_path.write_text(content_text, encoding="utf-8")
        print(f"  ✅ 成功保存: {docx_path.name} ({len(content_text)} 字)")
        return True
    else:
        print(f"  ❌ 未找到 {name} 链接")
        return False

def main():
    with sync_playwright() as p:
        b = p.chromium.launch(headless=True, executable_path=CHROME, args=['--no-sandbox'])
        page = b.new_page()
        for item in LAWS_TO_FETCH:
            search_and_download_law(page, item)
            time.sleep(1)
        b.close()

if __name__ == '__main__':
    main()
