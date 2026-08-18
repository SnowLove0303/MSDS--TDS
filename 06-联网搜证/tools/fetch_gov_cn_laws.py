# -*- coding: utf-8 -*-
"""针对性从中国政府网 (gov.cn) 抓取官方颁布的标准全文 HTML 并转化为 docx / txt."""
import sys, os, time, re
from pathlib import Path
from playwright.sync_api import sync_playwright
import docx

CHROME = r'E:\MorenAnzhuangLujing\Chrome\Chrome\Application\chrome.exe'
ROOT_DIR = Path(r'F:\正式项目与模块化内容\冠志\MSDS\Word 覆写模块\数据库与推断引擎\法规匹配库\法规原文归档')

GOV_LAWS = [
    {
        "name": "中华人民共和国职业病防治法",
        "sub": "法律",
        "url": "https://www.gov.cn/banshi/2005-08/21/content_25089.htm"
    },
    {
        "name": "中华人民共和国消防法",
        "sub": "法律",
        "url": "https://www.gov.cn/flfg/2008-10/29/content_1134314.htm"
    },
    {
        "name": "中华人民共和国安全生产法",
        "sub": "法律",
        "url": "https://www.gov.cn/flfg/2014-09/01/content_2743841.htm"
    },
    {
        "name": "易制毒化学品管理条例",
        "sub": "行政法规",
        "url": "https://www.gov.cn/zwgk/2005-08/30/content_27608.htm"
    },
    {
        "name": "使用有毒物品作业场所劳动保护条例",
        "sub": "行政法规",
        "url": "https://www.gov.cn/gongbao/content/2002/content_61556.htm"
    },
    {
        "name": "易制爆危险化学品名录",
        "sub": "目录公告",
        "url": "https://www.mps.gov.cn/n2254314/n6409334/c6475685/content.html"
    }
]

def save_doc_and_txt(name, text, out_dir):
    out_dir.mkdir(parents=True, exist_ok=True)
    txt_path = out_dir / f"{name}.txt"
    txt_path.write_text(text, encoding="utf-8")
    
    docx_path = out_dir / f"{name}.docx"
    doc = docx.Document()
    doc.add_heading(name, level=0)
    for p in text.splitlines():
        p_clean = p.strip()
        if p_clean:
            doc.add_paragraph(p_clean)
    doc.save(str(docx_path))
    print(f"  ✅ 成功保存: {docx_path.name} ({len(text)} 字)")

def main():
    with sync_playwright() as p:
        b = p.chromium.launch(headless=True, executable_path=CHROME, args=['--no-sandbox'])
        context = b.new_context(user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36")
        page = context.new_page()
        
        for item in GOV_LAWS:
            name = item["name"]
            sub = item["sub"]
            url = item["url"]
            out_dir = ROOT_DIR / sub
            
            print(f"\n[Gov.cn 抓取] {name} ({sub})...")
            try:
                page.goto(url, timeout=30000)
                page.wait_for_timeout(2000)
                
                text = page.evaluate('''() => {
                    const el = document.querySelector('#UCAP-CONTENT, #main_content, .article-content, #content, .content, td.p1, body');
                    return el ? el.innerText : document.body.innerText;
                }''')
                
                # 过滤前导杂质
                start = text.find('目　录')
                if start < 0: start = text.find('第一章')
                if start < 0: start = text.find('第一条')
                if start >= 0:
                    text = text[start:]
                    
                if len(text) > 400:
                    save_doc_and_txt(name, text, out_dir)
                else:
                    print(f"  ❌ 抓取内容过短: {len(text)} 字")
            except Exception as e:
                print(f"  ❌ 抓取异常: {e}")
            time.sleep(1)
        b.close()

if __name__ == '__main__':
    main()
