# -*- coding: utf-8 -*-
"""单任务快速爬取并归档剩余标准."""
import urllib.parse, re, time
from pathlib import Path
from playwright.sync_api import sync_playwright
import docx

CHROME = r'E:\MorenAnzhuangLujing\Chrome\Chrome\Application\chrome.exe'
ROOT_OUT = Path(r"F:\正式项目与模块化内容\冠志\MSDS\Word 覆写模块\数据库与推断引擎\法规匹配库\标准原文归档")

STANDARDS_REMAINING = [
    {"name": "GB 15258-2009 化学品安全标签编写规定", "sub": "SDS框架", "q": "GB 15258-2009 化学品安全标签编写规定 完整内容 应急电话"},
    {"name": "GB 13690-2009 化学品分类和危险性公示 通则", "sub": "SDS框架", "q": "GB 13690-2009 化学品分类和危险性公示 通则 完整内容"},
    {"name": "GB 15603-2022 危险化学品仓库储存通则", "sub": "储存消防", "q": "GB 15603-2022 危险化学品仓库储存通则 禁配物 隔离贮存 全文"},
    {"name": "GB 12158-2006 防止静电事故通用导则", "sub": "储存消防", "q": "GB 12158-2006 防止静电事故通用导则 全文 防静电措施"},
    {"name": "GB 50016-2014(2018年版) 建筑设计防火规范", "sub": "储存消防", "q": "GB 50016-2014 建筑设计防火规范 仓库火灾危险性分类 3.1"},
    {"name": "GB 50140-2005 建筑灭火器配置设计规范", "sub": "储存消防", "q": "GB 50140-2005 建筑灭火器配置设计规范 水溶性 极性溶剂 灭火剂"},
    {"name": "WS 444-2014 化学品中毒医疗救治技术导则", "sub": "急救与职业卫生", "q": "WS 444-2014 化学品中毒医疗救治技术导则 急救流程 全文"},
    {"name": "GBZ 71-2013 职业性急性化学物中毒诊断标准(总则)", "sub": "急救与职业卫生", "q": "GBZ 71-2013 职业性急性化学物中毒诊断标准(总则) 全文"},
    {"name": "GBZ 2.2-2007 工作场所有害因素职业接触限值 第2部分：物理因素", "sub": "急救与职业卫生", "q": "GBZ 2.2-2007 物理因素 职业接触限值 噪声 高温 全文"},
    {"name": "GB 2890-2009 呼吸防护 自吸过滤式防毒面具", "sub": "职业卫生", "q": "GB 2890-2009 呼吸防护 自吸过滤式防毒面具 滤毒罐分类 全文"},
    {"name": "GB 2626-2019 呼吸防护 自吸过滤式防颗粒物呼吸器", "sub": "职业卫生", "q": "GB 2626-2019 呼吸防护 自吸过滤式防颗粒物呼吸器 KN95 KP95 全文"},
    {"name": "GB-T 18664-2002 呼吸防护用品的选择、使用与维护", "sub": "职业卫生", "q": "GB/T 18664-2002 呼吸防护用品的选择、使用与维护 选用原则 全文"},
    {"name": "GB 50483-2019 化工建设项目环境保护设计标准", "sub": "环保与危废", "q": "GB 50483-2019 化工建设项目环境保护设计标准 事故水池 事故池设计"},
    {"name": "GB-T 21844-2008 化合物 闪点测定 快速平衡闭杯法", "sub": "理化与毒理测试", "q": "GB/T 21844-2008 闪点测定 快速平衡闭杯法 原理 步骤"},
    {"name": "GB-T 21853-2008 化学品 沸点测定", "sub": "理化与毒理测试", "q": "GB/T 21853-2008 化学品 沸点测定试验方法 全文"},
    {"name": "GB-T 21845-2008 化学品 水溶解度测定", "sub": "理化与毒理测试", "q": "GB/T 21845-2008 化学品 水溶解度试验 全文"},
    {"name": "GB-T 21578-2008 危险品 反应性试验方法", "sub": "理化与毒理测试", "q": "GB/T 21578-2008 危险品 反应性试验方法 克南试验 全文"},
    {"name": "GB-T 21807-2008 化学品 鱼类急性毒性试验", "sub": "生态毒理测试", "q": "GB/T 21807-2008 化学品 鱼类毒性试验 试验方法 全文"},
    {"name": "GB-T 21809-2008 化学品 溞类急性活动抑制试验", "sub": "生态毒理测试", "q": "GB/T 21809-2008 化学品 溞类急性活动抑制试验 全文"},
    {"name": "GB-T 21805-2008 化学品 淡水藻生长抑制试验", "sub": "生态毒理测试", "q": "GB/T 21805-2008 化学品 藻类生长抑制试验 全文"}
]

def fetch_single(page, item):
    name = item["name"]
    sub = item["sub"]
    q = item["q"]
    out_dir = ROOT_OUT / sub
    out_dir.mkdir(parents=True, exist_ok=True)
    
    clean_name = re.sub(r'[\\/:*?"<>|]', "_", name)
    out_docx = out_dir / f"{clean_name}.docx"
    out_txt = out_dir / f"{clean_name}.txt"
    
    if out_docx.exists() and out_docx.stat().st_size > 5000:
        print(f"⏭ [已归档] {name}")
        return True
        
    print(f"\n🚀 [正在获取标准] {name} ({sub})...")
    url = f"https://www.so.com/s?q={urllib.parse.quote(q)}"
    try:
        page.goto(url, timeout=15000)
        page.wait_for_timeout(1000)
        
        links = page.evaluate('''() => {
            const as = Array.from(document.querySelectorAll('h3 a, .res-title a'));
            return as.map(a => [(a.innerText || '').trim(), a.href]);
        }''')
        
        target_href = None
        for l in links:
            if "文库" in l[0] and ("付费" in l[0] or "VIP" in l[0]): continue
            target_href = l[1]
            break
            
        if target_href:
            page.goto(target_href, timeout=15000)
            page.wait_for_timeout(1500)
            
            text = page.evaluate('''() => {
                const el = document.querySelector('article, #content, .content, .main_content, .article-content, #view_content, body');
                return el ? el.innerText : document.body.innerText;
            }''')
            
            if len(text) > 300:
                out_txt.write_text(text, encoding="utf-8")
                doc = docx.Document()
                doc.add_heading(name, level=0)
                for line in text.splitlines():
                    ls = line.strip()
                    if ls:
                        doc.add_paragraph(ls)
                doc.save(str(out_docx))
                print(f"  ✅ [归档成功] {out_docx.name} ({len(text)} 字符)")
                return True
    except Exception as e:
        print(f"  ❌ 异常: {e}")
    return False

def main():
    with sync_playwright() as p:
        b = p.chromium.launch(headless=True, executable_path=CHROME, args=['--no-sandbox'])
        context = b.new_context(user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/126.0.0.0 Safari/537.36")
        page = context.new_page()
        
        ok = 0
        for it in STANDARDS_REMAINING:
            if fetch_single(page, it):
                ok += 1
        b.close()
        print(f"\n全部完成! 成功: {ok} / {len(STANDARDS_REMAINING)}")

if __name__ == '__main__':
    main()
