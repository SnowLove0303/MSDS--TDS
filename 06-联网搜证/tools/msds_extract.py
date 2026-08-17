# -*- coding: utf-8 -*-
"""批量提取 MSDS docx 的产品画像: 产品名/中文名/分类/危险概述/成分表.
用法: python msds_extract.py <MSDS目录> <输出json>"""
import sys, json, re, glob
from pathlib import Path
from docx import Document

def cell_text(c):
    return c.text.strip().replace('\n', ' ').replace('\r', '')

def extract(path):
    doc = Document(path)
    meta = {'file': Path(path).name, 'tables': []}
    # 段落中也可能有标题/文本
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    meta['paras'] = paras[:10]
    for tb in doc.tables:
        rows = []
        for r in tb.rows:
            cells = [cell_text(c) for c in r.cells]
            # 去重相邻合并单元格
            dedup = []
            for c in cells:
                if not dedup or dedup[-1] != c:
                    dedup.append(c)
            rows.append(dedup)
        meta['tables'].append(rows)
    return meta

def find_key(meta, keys):
    for rows in meta['tables']:
        for row in rows:
            joined = '|'.join(row)
            for k in keys:
                if k in joined:
                    # 取该行后续单元格
                    vals = [c for c in row[1:] if c and not c.startswith('|')]
                    if vals: return vals[0]
                    # 下一行
    return ''

def main():
    root = Path(sys.argv[1])
    out = Path(sys.argv[2])
    results = []
    for f in sorted(root.rglob('*.docx')):
        try:
            m = extract(f)
            # 从表0(物料及供方)提取产品名/中文名
            name = chn = cat = ''
            for rows in m['tables']:
                for row in rows:
                    j = '|'.join(row)
                    if '产品名称' in j: name = row[-1]
                    elif '中文名称' in j: chn = row[-1]
                    elif '化学品分类' in j: cat = row[-1]
            # 危险概述
            hazard = ''
            for rows in m['tables']:
                for row in rows:
                    j = '|'.join(row)
                    if 'GHS分类' in j:
                        hazard = row[-1]
                        break
                if hazard: break
            # 成分表
            comps = []
            for rows in m['tables']:
                for i, row in enumerate(rows):
                    j = '|'.join(row)
                    if 'CAS编号' in j:
                        for rr in rows[i+1:]:
                            if not rr or not rr[0]: continue
                            if any(k in '|'.join(rr) for k in ['成分/组成','SDS','化学品','产品类型']): break
                            # 尝试解析 名称 + CAS + %
                            cj = '|'.join(rr)
                            cas = re.search(r'\b\d{2,7}-\d{2}-\d\b', cj)
                            if cas:
                                comps.append({'row': rr, 'cas': cas.group(0), 'raw': cj})
                        break
            results.append({'file': Path(f).name, 'brand': '冠志' if '冠志' in str(f) else '国彩',
                            'name': name, 'chn': chn, 'cat': cat, 'hazard': hazard, 'comps': comps})
        except Exception as e:
            results.append({'file': Path(f).name, 'error': str(e)})
    out.write_text(json.dumps(results, ensure_ascii=False, indent=1), encoding='utf-8')
    print(f'提取 {len(results)} 份 → {out}')
    # 统计
    no_comps = [r for r in results if not r.get('comps')]
    print('无成分条目:', len(no_comps))
    for r in no_comps[:5]:
        print('  -', r.get('file'), r.get('error',''))

if __name__ == '__main__':
    main()
