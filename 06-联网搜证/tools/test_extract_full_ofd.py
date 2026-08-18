# -*- coding: utf-8 -*-
from playwright.sync_api import sync_playwright
import docx
from pathlib import Path

CHROME = r'E:\MorenAnzhuangLujing\Chrome\Chrome\Application\chrome.exe'
url = "https://flk.npc.gov.cn/detail?id=ff8081817a66b816017a7956b7db0ad4&fileId=&type=&title=%E4%B8%AD%E5%8D%8E%E4%BA%BA%E6%B0%91%E5%85%B1%E5%92%8C%E5%9B%BD%E5%AE%89%E5%85%A8%E7%94%9F%E4%BA%A7%E6%B3%95"

with sync_playwright() as p:
    b = p.chromium.launch(headless=True, executable_path=CHROME, args=['--no-sandbox'])
    context = b.new_context(user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/126.0.0.0 Safari/537.36")
    page = context.new_page()
    page.goto(url, timeout=30000)
    page.wait_for_timeout(3000)
    
    f = next(f for f in page.frames if 'flkofd' in f.url)
    print("Frame loaded. Scrolling and loading all pages...")
    
    # 模拟在 OFD reader 内部滚动到底部加载全部页
    for i in range(45):
        f.evaluate('() => window.scrollBy(0, 1200)')
        page.wait_for_timeout(300)
        
    full_text = f.evaluate('''() => {
        const textElements = document.querySelectorAll('.text, div[class*="text"], span');
        const lines = [];
        let currentText = '';
        textElements.forEach(el => {
            const t = el.innerText.trim();
            if (t) lines.push(t);
        });
        return document.body.innerText;
    }''')
    
    print(f"Full Text Extracted: {len(full_text)} characters.")
    print("Sample:\n", full_text[:400])
    
    # 保存为 docx
    doc = docx.Document()
    doc.add_heading("中华人民共和国安全生产法（2021年最新修正版）", level=0)
    for p_line in full_text.splitlines():
        p_clean = p_line.strip()
        if p_clean:
            doc.add_paragraph(p_clean)
    out_p = Path("中华人民共和国安全生产法.docx")
    doc.save(str(out_p))
    print(f"Saved {out_p.name} ({out_p.stat().st_size} bytes)")
    b.close()
