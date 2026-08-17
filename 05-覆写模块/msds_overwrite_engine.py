#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
MSDS 覆写引擎（增删查改） —— 以模板为准（模板驱动）
====================================================
依赖：结构读取 的 core.extract.read_msds（只读检索，用于"查"与"闭环校验"）

覆写原则（以模板 PEA-4139 为准构建通路）：
  - 模板的 序号、标签、字段结构、总体格式 是覆写通路的骨架/契约；
  - 默认模板驱动：模板字段行全部保留（按模板原顺序为输出主序），写入项只做”值源”，
    按归一化标签匹配模板行并覆写值格；标签格继承模板原文（序号/标签/空格/双 run 格式 100% 保留）；
  - 改：写入项标签匹配到模板行 → 仅覆写值格（标签格不动，模板格式完全继承）；
  - 增：写入项有、模板没有 → 克隆同 section 参考行（deepcopy，格式继承），
        用写入项的 seq+label 重建标签格后插入到相邻模板行之间；
  - 删：仅当写入项该行显式 “delete”: true → 删除对应模板行
        （模板驱动不再”默认删除模板多出行”；模板行默认就是骨架）；
  - 排版空格围栏：新增行只替换文字/符号，保留参考行标签格与值格的排版空格；
  - 校验：覆写后调用 read_msds 读回输出文档，与写入项逐字段对比。

写入项 schema (JSON):
{
  "keep_structure": "all",      // 默认 all=模板驱动；也可 [2] 仅指定节模板驱动
  "sections": {
    "1": [ {"label": "中文名称", "value": "..."}, ... ],
    "3": { "产品类型": "混合物",
           "components": [{"name":"...","cas":"...","conc":"..."}, ...] },
    "9": [ {"label": "外观", "value": "乳白色液体"}, ... ]
  }
}
写入项 label 仅作匹配键（归一化+字段映射对齐模板）；值格覆写用 value。
显式删除模板行：{"label": "手套材质行", "delete": true}

用法:
  python msds_overwrite_engine.py --template <模板.docx> --write-items <写入项.json> --out <输出.docx>
"""

import argparse
import copy
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.document import Document as _Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P

# ---- 依赖：结构读取（MSDS 检索工具），仅用于“查”和“校验” ----
STRUCTURE_READ_DIR = Path(r"F:\正式项目与模块化内容\Word 覆写模块\结构读取")
if str(STRUCTURE_READ_DIR) not in sys.path:
    sys.path.insert(0, str(STRUCTURE_READ_DIR))
try:
    from core.extract import read_msds, build_hierarchy
    from core.structure import (normalize_component_name,
                                normalize_component_cas,
                                normalize_component_conc)
except Exception as e:  # pragma: no cover
    read_msds = None
    print(f"[WARN] 无法导入结构读取工具（read_msds）：{e}")

def _norm_comp(value, fn):
    """用结构读取的归一化函数处理成分值；失败时降级为去空白。"""
    try:
        return fn(value)
    except Exception:
        if value is None:
            return ""
        return re.sub(r"\s+", "", str(value))

# ---- 保留清单：固定/公司信息字段，永不参与删除 ----
# 键为 section 号；值为该 section 中受保护的归一化标签集合。
DEFAULT_PRESERVE = {
    1: {"产品名称", "供应商信息", "供应商名称", "供应商地址", "电话", "传真"},
}

# ---- note 行覆写键（单格通栏正文行的定位通道）----
# 推导方案/写入项用“语义键”匹配单格正文 note 行（其文本不呈“标签：”形态，
# 无法用 label 归一化匹配）。mode='startswith' 表示行文本以 pat 开头即命中。
# 纯值驱动节：模板字段省略即保留模板原值，不自动填"无数据/无"。
# S11 毒性资料按 GB/T 17519 属"纯值驱动"（方案给了值才写，省略保留，
# 不自行判断数据来源/缺数据语义）。
_PURE_VALUE_SECTIONS = {11}

NOTE_ROW_KEYS = {
    8: {
        "手套材质_氟化橡胶FKM": ("startswith", "氟化橡胶"),
        "手套材质_丁基橡胶IIR": ("startswith", "丁基橡胶"),
        "手套材质_丁腈橡胶NBR": ("startswith", "丁腈橡胶"),
    },
    12: {
        "生态毒性_延续": ("startswith", "其他"),
    },
    15: {
        "法规入口": ("startswith", "符合下列法规要求"),
        "其它的规定": ("startswith", "其它的规定"),
    },
    16: {
        "免责声明": ("startswith", "就我们所掌握的知识信息"),
    },
}

# ---- 参考行选择策略（新增行克隆哪个模板行的格式/排版）----
# 取值：
#   None / "auto"         自动：与目标 seq 相邻的模板字段行（用它的格式/排版）
#   "last"                最后一个字段行（历史默认行为）
#   "label:中文名称"       按标签指定（如 "label:中文名称"，跨字段也能用）
# 也可在 overwrite(ref_pick=...) 或写入项 JSON 的 sections.<N>.ref 里覆盖。
DEFAULT_REF_PICK = "auto"

def log(msg, level="INFO"):
    print(f"[{level}] {msg}")

# ---- 字段别名：Excel/推断输出与模板写法不一致时归一化（如 阀值/阈值）----
FIELD_ALIASES = {
    "嗅觉阀值": "嗅觉阈值",
}

# ---- 字段映射规则层（以模板为主）----
# 把源文件/检索输出里的字段写法，映射到模板的标准字段写法（打标签/建通道）。
# 键=归一化后的源标签；值=归一化后的模板标签。分节维护，避免跨节误匹配。
# 归一化规则与 norm_label 一致（去空白 + FIELD_ALIASES）。
SECTION_FIELD_MAP = {
    1: {},
    2: {
        # 源 S2 单格文本拆出的写法 → 模板 S2 标准字段
        "GHS危险性类别": "GHS危险性类别",        # 模板 2.1
        "物质或混合物的分类": "GHS危险性类别",    # 源 2.1 子标题 → 模板 2.1
        "物质或混合物分类": "GHS危险性类别",
        "标签要素": "GHS标签要素",               # 模板 2.2
        # 源文件单格文本把"标签要素"下的内容写成"GHS-象形图 必须列在标签上的有害成份…"，
        # 其语义是标签要素（必须列在标签上的有害成分），归并到模板 2.2 GHS标签要素。
        "GHS-象形图": "GHS标签要素",
        "GHS象形图": "GHS象形图",                # 模板 2.2 下（真正的象形图描述）
        "危害性说明": "危险性说明",               # 模板 2.4
        "防范说明": "防范说明",                   # 模板 2.5
        "其他危险": "其他危险",                   # 模板 2.6
    },
    3: {},
    9: {
        "嗅觉阀值": "嗅觉阈值",
        # 模板字段名比源文件更完整/规范 → 以模板为主做通道
        "pH值": "pH值（1%水溶液）",               # 模板 9.3
        "相对密度": "密度",                       # 模板 9.12（相对密度=密度，模板用"密度"）
        "辛醇/水分配系数的对数值": "辛醇/水分配系数对数值",   # 模板 9.15（源多"的"）
    },
}

# ---- 字段映射外部配置（隐患2：映射通道不锚定单个模板）----
# 内置 SECTION_FIELD_MAP 作为兜底默认；启动时自动加载同目录 field_maps.json 合并覆盖。
# 后续引入中英文/冠志国彩等新模板时，只需新增一份 field_maps.<模板名>.json，无需改代码。
FIELD_MAP_FILE = Path(__file__).parent / "field_maps.json"

def _load_field_map(path=None):
    """加载外部字段映射配置并与内置合并。
    外部配置 schema: {"2": {"GHS-象形图": "GHS标签要素"}, ...}，节号可为字符串。
    返回合并后的完整映射 dict（不修改内置 SECTION_FIELD_MAP）。
    """
    merged = copy.deepcopy(SECTION_FIELD_MAP)
    p = Path(path) if path else (FIELD_MAP_FILE if FIELD_MAP_FILE.exists() else None)
    if p is not None:
        try:
            ext = json.loads(p.read_text(encoding="utf-8"))
            n_rules = 0
            for sec, mapping in ext.items():
                try:
                    sec_i = int(sec)
                except (TypeError, ValueError):
                    continue                    # 非数字键（如 _说明 注释）跳过
                if not isinstance(mapping, dict):
                    continue
                merged.setdefault(sec_i, {}).update(mapping)
                n_rules += len(mapping)
            log(f"字段映射配置已加载: {p} ({n_rules} 条规则)")
        except Exception as e:
            log(f"字段映射配置加载失败（用内置默认）: {e}", "WARN")
    return merged

FIELD_MAP = _load_field_map()

# ---- 空值覆写策略（隐患4：区分“缺字段”与“空值”）----
# 用户语义：写入项含该字段但值为空 → 锚定并覆成空值（清空模板）。
# 但模板可能预填了合理默认值，被清空需要可审计，因此默认 warn（覆写+告警）。
#   overwrite  空值覆写模板（清空，不留排版空格残留）
#   preserve   空值跳过（保留模板原值，仅“缺字段”才不动模板）
#   warn       同 overwrite，且模板原值非空时记录 WARN
EMPTY_POLICY_DEFAULT = "warn"

# ---- keep_structure 节必需字段清单（隐患3：保留模板结构不掩盖缺字段）----
# 这些是模板标准字段（映射目标标签），keep_structure 节若写入项缺失它们，
# 该行会保留模板默认值（如信号词"警告"），容易被误当真实数据 → 覆写时告警。
REQUIRED_FIELDS = {
    2: {"GHS危险性类别", "信号词", "危险性说明", "防范说明"},
}

def map_field_label(sec_no, label):
    """按节把源标签映射为模板标准标签；无映射则返回原标签。"""
    nl = norm_label(label)
    m = FIELD_MAP.get(sec_no, {})
    target = m.get(nl)
    if target:
        return target
    # 别名表（跨节）
    return FIELD_ALIASES.get(nl, nl)

def seq_key(seq):
    """seq 字符串转版本号元组用于大小比较：'9.19'->(9,19)，'9.2'->(9,2)。"""
    if not seq:
        return None
    parts = str(seq).split(".")
    nums = []
    for p in parts:
        if p.isdigit():
            nums.append(int(p))
        else:
            break
    return tuple(nums)

# ------------------------- 通用工具 -------------------------
def norm_field(text):
    """从字段行标签文本提取 (seq, label)。
    模板示例：'9.1  外观：' -> ('9.1','外观')；'中文名称：' -> ('','中文名称')。
    """
    if not text:
        return "", ""
    s = str(text).replace("　", " ").replace("\xa0", " ").strip()
    m = re.match(r"^(\d+(?:\.\d+)*)\s*(.*)$", s)
    if m and m.group(2):
        seq, rest = m.group(1), m.group(2).strip()
    else:
        seq, rest = "", s
    # 去尾部冒号/标点/空白
    rest = re.sub(r"[：:，,。；;]+$", "", rest).strip()
    rest = rest.replace(" ", "")
    return seq, norm_label(rest)

def norm_label(s):
    """归一化标签（供匹配）：去空白 + 别名映射。"""
    if not s:
        return ""
    s = str(s).replace(" ", "").replace("　", "").replace("\xa0", "")
    return FIELD_ALIASES.get(s, s)

def _rpr_of_run(r):
    return r._r.rPr if r is not None else None

# ------------------------- 格式捕获-重放 -------------------------
# 模板值格/标签格可能存在：多段落、多 run 混排（中文宋体+数字 Times）、
# 超链接、字段域（页码/日期）、图片（S2 象形图/S14 标签）。
# 覆写原则（以模板格式为准）：捕获目标单元格每个段落的 pPr + 代表 run 的 rPr，
# 新值按 \n 拆段逐段重放；多余段落删除、不足克隆末段格式。

def _tc_paragraphs(cell):
    """返回目标单元格的 w:p 列表（无段落则补一个空段落）。"""
    if isinstance(cell, _Cell):
        tc = cell._tc
    else:
        tc = cell
    paras = tc.findall(qn("w:p"))
    if not paras:
        p = tc.makeelement(qn("w:p"), {})
        tc.append(p)
        paras = [p]
    return paras

def _capture_para(p):
    """捕获段落格式：pPr（保留不动）+ run/hyperlink 的 rPr（代表格式）。
    隐患5加固：优先取普通 run 的 rPr；只有段落全部是超链接时才取超链接 rPr。
    （超链接样式常带下划线/蓝色，覆写降级为纯文本后不应继承超链接外观。）
    返回 {'rPr': deepcopy 或 None}。pPr 不单独拷贝，因为段落元素本身保留。
    """
    hyper_rPr = None
    for child in p.iterchildren():
        if child.tag == qn("w:r"):
            rp = child.find(qn("w:rPr"))
            if rp is not None:
                return {"rPr": copy.deepcopy(rp)}   # 首个普通 run 即返回
        elif child.tag == qn("w:hyperlink"):
            for r in child.findall(qn("w:r")):
                rp = r.find(qn("w:rPr"))
                if rp is not None and hyper_rPr is None:
                    hyper_rPr = copy.deepcopy(rp)
    return {"rPr": hyper_rPr}

def _clear_para_content(p):
    """清空段落内全部 run/超链接/域（保留 pPr 与段落属性）。"""
    for child in list(p):
        if child.tag in (qn("w:r"), qn("w:hyperlink"), qn("w:fldSimple")):
            p.remove(child)

def _replay_para(p, text, fmt):
    """按捕获格式 fmt 重建单段文本。text 含 \n 用 <w:br/> 软换行（不产生新段落）。
    超链接降级为纯文本（格式仍继承 rPr）；图片/域随旧内容一并清除。
    """
    _clear_para_content(p)
    new_r = p.makeelement(qn("w:r"), {})
    if fmt["rPr"] is not None:
        new_r.append(copy.deepcopy(fmt["rPr"]))
    lines = str(text).split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            br = p.makeelement(qn("w:br"), {})
            new_r.append(br)
        t = p.makeelement(qn("w:t"), {})
        t.text = line
        if line and (line[0].isspace() or line[-1].isspace()):
            t.set(qn("xml:space"), "preserve")
        new_r.append(t)
    p.append(new_r)
    return new_r

def set_paragraph_text(p_elm, text, xml_space=True):
    """单段格式捕获-重放（保留 pPr + 代表 run rPr）。"""
    fmt = _capture_para(p_elm)
    _replay_para(p_elm, text, fmt)

def set_cell_text(cell, text):
    """格式捕获-重放整格：捕获每个段落格式，新值按 \\n 拆段逐段重放。
    段落数不足 → 克隆末段格式；段落数超出 → 删除多余段落（保留首个）。
    """
    paras = _tc_paragraphs(cell)
    fmts = [_capture_para(p) for p in paras]
    lines = [str(line) for line in str(text).split("\n")]
    # 段落数不足：克隆末段（deepcopy 后清内容，保留其 pPr/rPr 载体）。
    # 注意 anchor 必须随插入推进，否则重复 addnext(同一锚点) 会让新段落逆序。
    if fmts and len(paras) < len(lines):
        src_fmt = _capture_para(paras[-1])
        anchor = paras[-1]
        for _ in range(len(lines) - len(paras)):
            new_p = copy.deepcopy(anchor)
            _clear_para_content(new_p)
            anchor.addnext(new_p)
            anchor = new_p
            paras.append(new_p)
            fmts.append(src_fmt)
    # 段落数超出：删除多余段落
    if len(paras) > len(lines) and lines:
        for p in paras[len(lines):]:
            p.getparent().remove(p)
        paras = paras[:len(lines)]
        fmts = fmts[:len(lines)]
    # 逐段重放
    for p, line, fmt in zip(paras, lines, fmts):
        _replay_para(p, line, fmt)


def _label_runs(p):
    """返回段落内普通 run 列表（跳过超链接/域，仅文本 run）。"""
    return [r for r in p.iterchildren() if r.tag == qn("w:r")]


def set_label_text(cell, text):
    """标签格格式捕获-重放（保留模板'序号run + 中文run'的多run结构）。

    模板标签格标准格式（PEA-4139 惯例）：
      前导序号 run → 西文（Times New Roman / None）10.5pt 加粗；
      中文标签 run  → 默认（宋体/Arial）12pt 加粗。
    重建时按模板 run 结构重放：前导序号部分继承模板首个含文本 run 的 rPr，
    中文标签部分继承模板最后一个含文本 run 的 rPr。

    注意：模板个别标签格存在尾部对齐空格 run（如 S9 '9.2 嗅觉阈值' 为
    [序号 TNR10.5][中文 Arial12][尾空格 TNR10.5] 三 run）。此类纯空格 run
    不做格式样本，否则中文标签会被错误应用西文小字号。

    模板只有单 run 或文本无序号时 → 退化为单 run 重放（保留模板 rPr）。
    """
    paras = _tc_paragraphs(cell)
    p = paras[0]
    runs = _label_runs(p)
    # 仅取“含非空白文本”的 run 作为格式样本（跳过纯尾部空格 run）
    text_rprs = []
    for r in runs:
        t = ''.join(x.text or '' for x in r.findall(qn("w:t")))
        if not t.strip():
            continue
        rp = r.find(qn("w:rPr"))
        text_rprs.append(copy.deepcopy(rp) if rp is not None else None)

    def _mk_run(rpr, txt):
        r = p.makeelement(qn("w:r"), {})
        if rpr is not None:
            r.append(copy.deepcopy(rpr))
        t = p.makeelement(qn("w:t"), {})
        t.text = txt
        if txt and (txt[0].isspace() or txt[-1].isspace()):
            t.set(qn("xml:space"), "preserve")
        r.append(t)
        p.append(r)

    _clear_para_content(p)
    text = str(text)
    # 分离前导序号（数字+可选点+空白）与中文标签主体
    m = re.match(r"^(\s*\d+(?:\.\d+)*\s*)(.*)$", text, re.S)
    if m and m.group(2).strip() and len(text_rprs) >= 2:
        _mk_run(text_rprs[0], m.group(1))
        _mk_run(text_rprs[-1], m.group(2))
    else:
        rpr = text_rprs[-1] if text_rprs else None
        _mk_run(rpr, text)

# ------------------------- 排版空格围栏 -------------------------
# 模板中很多字段行用“空格”做排版对齐（如 '9.2  嗅觉阈值：      ' 的多空格/尾空格）。
# 围栏原则：覆写时优先只替换“文字/符号”，保留模板的排版空格结构。

def _cell_w_text(tr, ci):
    """取 tr 第 ci 个单元格的首段纯文本（拼合 w:t）。"""
    tcs = tr.findall(qn("w:tc"))
    if ci >= len(tcs):
        return ""
    p = tcs[ci].find(qn("w:p"))
    if p is None:
        return ""
    return "".join(t.text or "" for t in p.iter(qn("w:t")))

def replace_label_keep_spacing(ref_text, seq, label, keep_ref_seq=True):
    """标签格排版护栏：保留模板的序号后空格与尾部对齐空格，只替换序号与文字。
    结构：前导空白 | 序号 | 序号后空白 | 文字 | 标点 | 尾部空白
    - keep_ref_seq=True（默认）：写入项给 seq 用写入项，否则沿用参考行/模板行序号
      （子标题行如 '8.1 暴露控制' 不因写入项无 seq 而丢序号；模板 '9.1  外观：' 同理）；
    - keep_ref_seq=False（新增行）：写入项无 seq 时**不引入**参考行序号，
      避免新增字段误沿用相邻字段序号（如 S11 新字段全部变成 '11.10'）。
    - 冒号跟随参考行风格（模板 '9.1  外观：' 带冒号、'8.1 暴露控制' 无冒号都保留）。
    示例：ref='9.2  嗅觉阈值：      ', seq='9.23', label='折射率'
        -> '9.23  折射率：      '
    """
    ref_text = ref_text or ""
    m = re.match(r"^(\s*)(\d+(?:\.\d+)*)?(\s*)(.*?)(\s*)$", ref_text, re.S)
    lead, ref_seq, seq_space, _body, tail = m.groups()
    body = (_body or "").strip()
    has_colon = bool(re.search(r"[：:]\s*$", body))
    # 序号：写入项给了就用写入项；否则按 keep_ref_seq 决定是否沿用参考行
    if keep_ref_seq:
        new_seq = (seq or "").strip() or (ref_seq or "")
    else:
        new_seq = (seq or "").strip()
    if new_seq:
        # 有序号：继承参考行的序号后空格排版（含无空格，如 '1.1产品名称：'）
        space = seq_space
    else:
        # 无序号：模板若有空格则继承（'中文名称： '），否则不加空格
        space = seq_space or ""
    core = f"{new_seq}{space}{label}" + ("：" if has_colon else "")
    return f"{lead}{core}{tail}"

def replace_value_keep_spacing(ref_text, new_text):
    """值格排版护栏：保留模板前导/尾随排版空格，只替换内容本体（文字/符号）。
    示例：ref='  无数据  ', new='7-9' -> '  7-9  '
    """
    ref_text = ref_text or ""
    m = re.match(r"^(\s*)(.*?)(\s*)$", ref_text, re.S)
    lead, _core, tail = m.groups()
    return f"{lead}{new_text}{tail}"

def cell_paragraph_xml(cell):
    if isinstance(cell, _Cell):
        return cell.paragraphs[0]._p
    return cell

def iter_block_items(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("unknown parent")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def section_tables(doc):
    """按 body 顺序识别 16 节表格：{sec_no: Table}。"""
    result = {}
    for block in iter_block_items(doc):
        if not isinstance(block, Table):
            continue
        first = block.rows[0].cells[0].text.strip()
        m = re.match(r"^(\d+)\.", first)
        if m:
            result[int(m.group(1))] = block
    return result

def section_sub_labels(template_path, sections=None):
    """用检索工具（build_hierarchy）确认模板父子级，返回每节“二级子标题”标签集。
    子标题行（如 S8 的 8.1 暴露控制、S3 的 3.2 成分）是模板结构骨架：
    不参与字段删除/重建，仅当写入项含该标签时覆写其标签格。
    返回 {sec_no: {norm_label, ...}}。
    """
    if read_msds is None or build_hierarchy is None:
        return {}
    try:
        r = read_msds(str(template_path))
        nodes = build_hierarchy(r)
    except Exception as e:
        log(f"读取模板父子级失败（跳过子标题保护）: {e}", "WARN")
        return {}
    result = {}
    for node in nodes:
        if sections is not None and node.number not in sections:
            continue
        subs = {norm_label(bt.title) for bt in node.big_titles
                if bt.kind == "sub" and bt.title}
        if subs:
            result[node.number] = subs
    return result

# ------------------------- 物理行映射 -------------------------
def row_kind(tr):
    """表格行分类：
      - 'field'：标签|值 两列结构行（首格有文本、≥2 独立 tc），可匹配/可删；
      - 'note' ：单格整行（通栏说明/跨列子标题）或首格空 → 结构行，不参与匹配/删除。
    """
    tcs = tr.findall(qn("w:tc"))
    if len(tcs) < 2:
        return "note"
    if not _cell_w_text(tr, 0).strip():
        return "note"
    return "field"

def _is_label_note_row(tr):
    """判断单格行是否为“标签形态”：有前导序号，或以冒号结尾。
    用于区分 note 标签行（'符合下列法规要求：'）与纯值行（法规内容行），
    纯值行可被上级标签行的值区重建删除，标签行则保留。
    """
    txt = _cell_w_text(tr, 0).strip()
    if re.match(r"^\d+(?:\.\d+)*\s*[^\d\s]", txt):
        return True
    if re.search(r"[：:]\s*$", txt):
        return True
    return False

def _delete_note_value_rows(tr, tbl):
    """删除 tr 之后连续的“纯值单格行”（值区），直到遇到 field 行或标签形态行。
    用于 S15 法规：'符合下列法规要求：' 标签行后的法规单格行是值区，
    覆写该字段的多行值后应删除模板原法规行，避免新旧法规重复。
    """
    rows = tbl.rows
    delete = []
    started = False
    for row in rows:
        rtr = row._tr
        if not started:
            if rtr is tr:
                started = True
            continue
        tcs = rtr.findall(qn("w:tc"))
        if len(tcs) < 2 and not _is_label_note_row(rtr):
            delete.append(rtr)
        else:
            break
    if delete:
        for r in delete:
            r.getparent().remove(r)
    return len(delete)

def row_map(tbl):
    """返回 {归一化标签: (row_idx, seq)}，跳过首行标题与结构行（note）。"""
    mapping = {}
    for ri in range(1, len(tbl.rows)):
        if row_kind(tbl.rows[ri]._tr) == "note":
            continue
        seq, label = norm_field(tbl.rows[ri].cells[0].text)
        mapping[norm_label(label)] = (ri, seq)
    return mapping

def find_component_rows(tbl):
    """定位 S3 成分表：返回 (产品类型行idx, 表头行idx, 成分起始idx)。"""
    prod_row = None
    comp_header = None
    for ri in range(1, len(tbl.rows)):
        seq, label = norm_field(tbl.rows[ri].cells[0].text)
        if label == "产品类型" and prod_row is None:
            prod_row = ri
        if label == "化学品名称" and comp_header is None:
            comp_header = ri
    comp_start = comp_header + 1 if comp_header is not None else None
    return prod_row, comp_header, comp_start

# ------------------------- 执行：改 -------------------------
def _cell_text_of(cell):
    """取单元格纯文本（兼容 _Cell 与 w:tc 元素）。"""
    if isinstance(cell, _Cell):
        return cell.text
    paras = cell.findall(qn("w:p"))
    return "\n".join("".join(t.text or "" for t in p.iter(qn("w:t"))) for p in paras)

def op_update_cell(cell, value, tag, keep_spacing=True, empty_policy="warn"):
    """覆写值格文本，格式继承。
    keep_spacing=True 时保留模板排版空格，只替换文字/符号。
    empty_policy: overwrite/preserve/warn（见 EMPTY_POLICY_DEFAULT）。
    隐患4：空值按策略处理——preserve 跳过；overwrite/warn 清空且不留排版空格残留；
    warn 在模板原值非空时记录告警（模板预填默认值被清空需可审计）。
    返回 (旧值, 新值-不含排版空格)。"""
    old = _cell_text_of(cell)
    val_str = "" if value is None else str(value)
    if not val_str.strip():
        if empty_policy == "preserve":
            return old, val_str
        set_cell_text(cell, "")
        if empty_policy == "warn" and old.strip():
            log(f"空值覆写（模板原值被清空）: {tag} {old!r} -> ''", "WARN")
        return old, val_str
    new_text = replace_value_keep_spacing(old, val_str) if keep_spacing else val_str
    set_cell_text(cell, new_text)
    return old, val_str

# ------------------------- 执行：删 -------------------------
def op_delete_rows(trs):
    """物理删除表格行（先移除靠后的行避免索引位移）。"""
    for tr in trs:
        tr.getparent().remove(tr)

# ------------------------- 执行：增 -------------------------
def find_insert_target(tbl, seq, rows_seq):
    """在 section 表格内定位新字段行的插入目标 tr。
    rows_seq: [(seq, tr)] 当前字段行的有序列表。
    返回 (anchor_tr, mode) mode in {'before','after'}。
    """
    key = seq_key(seq)
    if key is not None:
        for s, tr in rows_seq:
            sk = seq_key(s)
            if sk is not None and sk >= key:
                return tr, "before"
    # 无序号或全小于：追加到最后一个字段行之后
    if rows_seq:
        return rows_seq[-1][1], "after"
    return tbl.rows[0]._tr, "after"

def op_insert_row(tbl, ref_tr, new_label_text, new_value_text, tag):
    """克隆参考行 tr，替换标签/值文本，插到 ref_tr 后。返回新 tr。"""
    new_tr = copy.deepcopy(ref_tr)
    tcs = new_tr.findall(qn("w:tc"))
    if not tcs:
        raise RuntimeError(f"{tag}: 参考行无单元格，无法克隆")
    # 标签格（第 1 个 tc）取首个段落
    p_label = tcs[0].find(qn("w:p"))
    if p_label is None:
        raise RuntimeError(f"{tag}: 标签格无段落")
    set_paragraph_text(p_label, new_label_text)
    # 值格（第 2 个 tc，若存在）
    if len(tcs) >= 2:
        p_value = tcs[1].find(qn("w:p"))
        if p_value is not None:
            set_paragraph_text(p_value, new_value_text)
    ref_tr.addnext(new_tr)
    return new_tr

def pick_ref_row(tbl, rows_seq, seq, ref_pick):
    """按策略选择参考行 tr（新增行克隆其格式/排版）。
    rows_seq: [(seq, tr)] 当前字段行有序序列。
    ref_pick: None/'auto' | 'last' | 'label:标签' | 'row:行号'
    返回 tr 或 None。
    """
    if ref_pick is None:
        ref_pick = DEFAULT_REF_PICK
    pick = str(ref_pick or "auto")
    if pick.startswith("label:"):
        target = norm_label(pick[len("label:"):].strip())
        for _s, tr in rows_seq:
            _seq, _label = norm_field(_cell_w_text(tr, 0))
            if norm_label(_label) == target:
                return tr
        # 标签未找到：回退到最后一个
        log(f"参考行标签 {target!r} 未找到，回退到最后一个字段行", "WARN")
        return rows_seq[-1][1] if rows_seq else None
    if pick.startswith("row:"):
        try:
            ridx = int(pick[len("row:"):].strip())
            if 0 <= ridx < len(tbl.rows):
                return tbl.rows[ridx]._tr
        except ValueError:
            pass
        log(f"参考行行号 {pick!r} 无效，回退到最后一个字段行", "WARN")
        return rows_seq[-1][1] if rows_seq else None
    if pick == "last":
        return rows_seq[-1][1] if rows_seq else None
    # auto：与目标 seq 相邻的模板字段行
    key = seq_key(seq)
    if key is not None:
        best = None
        best_after = None
        for _s, tr in rows_seq:
            sk = seq_key(_s)
            if sk is None:
                continue
            if sk == key:
                return tr
            if sk > key and (best_after is None or sk < best_after[0]):
                best_after = (sk, tr)
            elif sk < key and (best is None or sk > best[0]):
                best = (sk, tr)
        if best_after is not None:
            return best_after[1]     # 优先取比目标大的最近一个（更接近目标）
        if best is not None:
            return best[1]
    return rows_seq[-1][1] if rows_seq else None

# ------------------------- 新增行边框衔接修正 -------------------------
# 模板字段行边框规律：首字段行 top=single（接节标题 bottom）/ 内部字段行 top=bottom=dotted /
# 末字段行 bottom=single（表格底边）。新增行克隆参考行时可能带末行边框（tdbs）或
# note 行边框（top=nil），插到中间后与上下行边框不连续（表现为“没继承表格内外框线”）。
# 修正策略：把连续新增行视为一个“块”，块首 top ← 块前模板行 bottom、
#           块末 bottom ← 块后模板行 top、块内其余边 = 该节内部字段行分隔样式（dotted）。
_BORDER_RANK = ["top", "left", "bottom", "right", "insideH", "insideV", "start", "end"]

def _tc_borders(tc):
    """取（不存在则建）单元格 tcBorders 元素。"""
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = tc.makeelement(qn("w:tcPr"), {})
        tc.insert(0, tcPr)
    b = tcPr.find(qn("w:tcBorders"))
    if b is None:
        b = tcPr.makeelement(qn("w:tcBorders"), {})
        tcPr.append(b)
    return b

def _border_of(tc, edge):
    """返回 tc 指定边边框元素（可能 None）。"""
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        return None
    b = tcPr.find(qn("w:tcBorders"))
    if b is None:
        return None
    return b.find(qn("w:" + edge))

def _has_real_border(tr, edge):
    """tr 是否任一单元格在指定边有实际边框（val 非 nil/none）。"""
    for tc in tr.findall(qn("w:tc")):
        b = _border_of(tc, edge)
        if b is not None and b.get(qn("w:val")) not in (None, "", "nil"):
            return True
    return False

def _tc_at(tr, ci):
    """返回 tr 第 ci 列 tc；越界时取最后一个（合并通栏行兜底）。"""
    tcs = tr.findall(qn("w:tc"))
    if ci < len(tcs):
        return tcs[ci]
    return tcs[-1] if tcs else None

def _make_border_elem(tc, edge, val):
    """构造 w:{edge} val={val} sz=4 color=auto space=0 边框元素。"""
    b = tc.makeelement(qn("w:" + edge), {})
    b.set(qn("w:val"), val)
    b.set(qn("w:color"), "auto")
    b.set(qn("w:sz"), "4")
    b.set(qn("w:space"), "0")
    return b

def _border_rank(el):
    local = el.tag.rsplit("}", 1)[-1]
    try:
        return _BORDER_RANK.index(local)
    except ValueError:
        return 99

def _set_border_adjacent(tc, edge, src_tc, src_edge, fallback_val):
    """把 src_tc 的 src_edge 边框（deepcopy，保留 val/sz/color）应用到 tc 的 edge 边框。
    src_tc 为 None 或 src 无该边框元素时，用 fallback_val 构造新边框。"""
    dst = _tc_borders(tc)
    for e in dst.findall(qn("w:" + edge)):
        dst.remove(e)
    new_b = None
    if src_tc is not None:
        sb = _border_of(src_tc, src_edge)
        if sb is not None:
            # 源边框为 nil/空 = “无线”，继承会造断线（如 S16 新增块末行接
            # 免责声明 top=nil）→ 视同无源，用 fallback 补齐
            _v = sb.get(qn("w:val"))
            if _v not in (None, "", "nil"):
                new_b = copy.deepcopy(sb)
                new_b.tag = qn("w:" + edge)
    if new_b is None:
        new_b = _make_border_elem(tc, edge, fallback_val)
    rank = _border_rank(new_b)
    inserted = False
    for child in dst:
        if _border_rank(child) > rank:
            child.addprevious(new_b)
            inserted = True
            break
    if not inserted:
        dst.append(new_b)

def _prev_template_row(ordered, tr, new_tr_ids):
    """返回有序行列表中 tr 之前（不含 tr）最近的模板（非新增）行；无则 None。"""
    for i, t in enumerate(ordered):
        if t is tr:
            for j in range(i - 1, -1, -1):
                if id(ordered[j]) not in new_tr_ids:
                    return ordered[j]
            return None
    return None

def _next_template_row(ordered, tr, new_tr_ids):
    """返回有序行列表中 tr 之后（不含 tr）最近的模板（非新增）行；无则 None。"""
    for i, t in enumerate(ordered):
        if t is tr:
            for j in range(i + 1, len(ordered)):
                if id(ordered[j]) not in new_tr_ids:
                    return ordered[j]
            return None
    return None

def _row_bottom_to_dotted(up):
    """原模板末行被新增行“顶到中间”后，其 bottom 从 single（表格底边）改为 dotted（行间分隔）。
    仅改显式为 single 的边；nil/None（靠表格级底边）保持不动。"""
    if up is None:
        return
    for tc in up.findall(qn("w:tc")):
        b = _border_of(tc, "bottom")
        if b is not None and b.get(qn("w:val")) == "single":
            b.set(qn("w:val"), "dotted")

def fix_first_row_border(tbl, ordered, tag=""):
    """紧邻节标题的首个数据行：其 top 与标题 bottom 不一致时修正。
    用于 prep 删除引导行（如 S11/S12 的“该产品无可用的…”）后，首字段行
    top 残留为中间行样式（dotted）而与标题底边（single）断线的情况。
    仅当首行 top 边框存在且值不同时改动；无 top 边框（S1/S16 功能行）不动。"""
    if not ordered:
        return
    title_tcs = tbl.rows[0]._tr.findall(qn("w:tc"))
    if not title_tcs:
        return
    title_bottom = _border_of(title_tcs[0], "bottom")
    if title_bottom is None:
        return
    want = title_bottom.get(qn("w:val"))
    if not want:
        return
    first = ordered[0]
    changed = False
    for tc in first.findall(qn("w:tc")):
        b = _border_of(tc, "top")
        if b is not None and b.get(qn("w:val")) != want:
            b.set(qn("w:val"), want)
            changed = True
    if changed:
        log(f"{tag}: 首行 top 边框修正为 {want}（衔接节标题底边）")

def fix_new_row_borders(tbl, ordered, new_tr_ids, tag=""):
    """对新增行块重设 top/bottom 边框，使与相邻模板行衔接（见模块注释）。
    处理顺序：先把被顶到中间的模板末行 bottom 还原为 dotted，再重设新增行块边框。"""
    blocks = []
    cur = []
    for tr in ordered:
        if id(tr) in new_tr_ids:
            cur.append(tr)
        elif cur:
            blocks.append(cur)
            cur = []
    if cur:
        blocks.append(cur)
    for blk in blocks:
        up = _prev_template_row(ordered, blk[0], new_tr_ids)
        down = _next_template_row(ordered, blk[-1], new_tr_ids)
        if down is None:
            # 本块成为表格新末行：原模板末行（块前）被顶到中间，bottom 还原为行间分隔
            _row_bottom_to_dotted(up)
        n = len(blk)
        for i, tr in enumerate(blk):
            is_first = (i == 0)
            is_last = (i == n - 1)
            for ci, tc in enumerate(tr.findall(qn("w:tc"))):
                if is_first:
                    src = _tc_at(up, ci) if up is not None else None
                    _set_border_adjacent(tc, "top", src, "bottom", fallback_val="single")
                else:
                    _set_border_adjacent(tc, "top", None, None, fallback_val="dotted")
                if is_last:
                    _src_tr, _src_edge = down, "top"
                    if down is not None and not _has_real_border(down, "top"):
                        # down 行无实际顶线（顶线由节标题/上一行底边提供，如 S16 免责声明
                        # 紧跟节标题）→ 末行 bottom 继承其上方模板行底边，避免断线
                        _prev = _prev_template_row(ordered, down, new_tr_ids)
                        if _prev is not None:
                            _src_tr, _src_edge = _prev, "bottom"
                    src = _tc_at(_src_tr, ci) if _src_tr is not None else None
                    _set_border_adjacent(tc, "bottom", src, _src_edge, fallback_val="single")
                else:
                    _set_border_adjacent(tc, "bottom", None, None, fallback_val="dotted")

# ------------------------- 单节同步（改/删/增，以写入项为准重建） -------------------------
def sync_section(tbl, sec_no, items, preserve_labels, ref_pick=None, order="write",
                 struct_labels=None, keep_structure=False, empty_policy="warn",
                 missing_policy="preserve", missing_text="无数据"):
    """对单个普通 section 执行改/删/增 —— 重建该节字段行为“写入项顺序”。
    items: [{'seq','label','value'}]，顺序 = 真实文件字段顺序（写入项顺序即输出顺序）。
    preserve_labels: 该 section 受保护标签集合（模板中有、写入项没有时保留原位）。
    struct_labels: 模板结构行标签（子标题等，来自检索工具父子级），不参与删除/重建。
    keep_structure: True=以模板为主（模板字段行全部保留，写入项映射到的覆写值，
                    写入项没有的模板行保留模板原值）；False=以写入项为准（模板多出删除）。
    order: "write"（默认）= 严格按写入项顺序重建字段行；
           "auto" = 与 "write" 等价（旧逻辑按 seq 排序导致顺序错乱，已弃用）。
    empty_policy: 空值覆写策略（见 op_update_cell）。
    missing_policy: 模板字段行存在、写入项未给该字段时的值格处理：
        "preserve"（默认）= 保留模板原值；
        "no_data"/"none" = 值格非空（旧产品残留）→ 覆写为 missing_text（"无数据"/"无"），
        清除残留；值格已空则不动；受保护字段与纯值驱动节（S11）不受影响。
    missing_text: missing_policy 生效时的填充文本（默认"无数据"）。
    返回该节操作日志。
    """
    logs = []
    struct_labels = struct_labels or set()

    # 1) 分类收集模板行（跳过首行标题）：
    #    - field 行：参与匹配/删除/重建
    #    - note 行（单格通栏/首格空）：结构行，保留原位；写入项含该标签则覆写整格
    tmpl_rows = []              # [(ri, tr, seq, nl)]
    note_rows = []              # [(ri, tr)]
    mapping = {}                # nl -> (tr, seq)
    orig_idx = {}               # id(tr) -> 原行号
    for ri in range(1, len(tbl.rows)):
        tr = tbl.rows[ri]._tr
        orig_idx[id(tr)] = ri
        if row_kind(tr) == "note":
            note_rows.append((ri, tr))
            continue
        seq, label = norm_field(tbl.rows[ri].cells[0].text)
        nl = norm_label(label)
        tmpl_rows.append((ri, tr, seq, nl))
        if nl not in mapping:
            mapping[nl] = (tr, seq)

    # 写入项标签先映射到模板标准标签（以模板为主建通道）
    mapped_items = []
    for it in items:
        m = dict(it)
        m["label"] = map_field_label(sec_no, it["label"])
        m["src_label"] = it["label"]           # 保留源标签（日志用）
        mapped_items.append(m)

    write_labels = {norm_label(it["label"]) for it in mapped_items}
    keep_labels = preserve_labels | struct_labels
    ref_seq_list = [(seq, tr) for (_ri, tr, seq, _nl) in tmpl_rows]
    consumed = set()            # 已被结构行（note）覆写消费的 item id

    # 2) 结构行（note）覆写：写入项含该标签 → 整格重放（标签格式继承），否则保留原位。
    #    值非空时写入同格（标签\n值）；并删除该标签行后连续的值区单格行
    #    （如 S15 法规列表：'符合下列法规要求：' 后的法规单格行）。
    #    note 行匹配两通道：① 行文本归一化标签；② NOTE_ROW_KEYS 语义键
    #    （单格正文行如免责声明，文本不呈“标签：”形态时用 key 前缀定位）。
    _note_keys = NOTE_ROW_KEYS.get(sec_no, {})
    for (ri, tr) in note_rows:
        tcs = tr.findall(qn("w:tc"))
        if not tcs:
            continue
        _s, _l = norm_field(_cell_w_text(tr, 0))
        nl = norm_label(_l)
        # 匹配文本：标签格优先；标签格空（如 S12 12.1 延续行）时用值格文本，
        # 使 NOTE_ROW_KEYS 的 startswith 语义键（如 '生态毒性_延续'→'其他'）能命中
        row_text = _cell_w_text(tr, 0)
        if not row_text.strip() and len(tcs) > 1:
            row_text = _cell_w_text(tr, 1)
        for it in mapped_items:
            il = norm_label(it["label"])
            hit = bool(nl) and il == nl
            _key_hit = None
            if not hit:
                for _key, (mode, pat) in _note_keys.items():
                    if il == _key and mode == "startswith" and row_text.startswith(pat):
                        hit = True
                        _key_hit = _key
                    if il == _key:
                        break
            if not hit:
                continue
            val = it.get("value")
            cont_cols = _key_hit is not None and len(tcs) > 1
            if cont_cols:
                # 语义键命中的延续/注释两列行（如 S12 cont、S15 法规入口）：
                # 标签格保留模板原文（通常为空），只操作值格
                if val is not None and str(val).strip():
                    set_cell_text(tcs[1], str(val))
                    _delete_note_value_rows(tr, tbl)
                else:
                    set_cell_text(tcs[1], "")
                logs.append(("update", sec_no, ri, it["src_label"], "", val))
                consumed.add(id(it))
                break
            label_text = replace_label_keep_spacing(row_text,
                                                    it.get("seq", ""), it["label"])
            if val is not None and str(val).strip():
                full = label_text + "\n" + str(val)
                set_cell_text(tcs[0], full)
                _delete_note_value_rows(tr, tbl)
            else:
                set_cell_text(tcs[0], label_text)
            logs.append(("update", sec_no, ri, it["src_label"], "", it["value"]))
            consumed.add(id(it))
            break

    # 3) 按写入项顺序构建最终行序列（模板驱动通路）：
    #    模板已有该标签 → 复用原行（骨架保留）；否则 → 克隆参考行（新增）。
    #    显式 "delete": true → 删除该模板行（模板驱动下唯一的删行途径）。
    #    模板驱动（keep_structure=True）：复用行标签格【完全继承模板原文】，
    #    仅覆写值格 —— 模板的序号/标签/空格/双 run 格式 100% 保留；
    #    新增行才用写入项的 seq+label 重建标签格。
    final = []                  # [(tr, is_new, it)]
    reused_ids = set()
    del_trs = []
    for it in mapped_items:
        if id(it) in consumed:
            continue
        nl = norm_label(it["label"])
        got = mapping.get(nl)
        want_delete = bool(it.get("delete"))
        if got is not None and want_delete:
            del_trs.append(got[0])
            logs.append(("delete", sec_no, -1, it["src_label"], ""))
            continue
        if got is not None and id(got[0]) not in reused_ids:
            tr, _seq = got
            reused_ids.add(id(tr))
            is_new = False
        else:
            ref_tr = pick_ref_row(tbl, ref_seq_list, it.get("seq", ""), ref_pick)
            if ref_tr is None:
                # 无字段行节（如 S16 note 区）：参考最后一行（正文 note 样式），
                # 而非首行节标题，避免新增行带标题格式。
                ref_tr = (tbl.rows[-1]._tr if len(tbl.rows) > 1 else tbl.rows[0]._tr)
            tr = copy.deepcopy(ref_tr)
            is_new = True
        tcs = tr.findall(qn("w:tc"))
        if keep_structure and not is_new:
            # 模板驱动：复用行标签格继承模板原文，只覆写值格
            pass
        else:
            # 新增行 / 写入项驱动模式：用写入项的 seq+label 重建标签格
            #  - 改行：模板原行有序号 → 用写入项 seq 或沿用模板序号；
            #         模板原行无序号（如 S7 '安全操作防范：'）→ 不引入写入项序号（S7 修复）。
            #  - 增行：写入项给 seq 用写入项；无 seq 不沿用参考行序号（S11 修复）。
            it_seq = (it.get("seq") or "").strip()
            if is_new:
                label_seq, keep_ref = it_seq, False
            else:
                label_seq, keep_ref = (it_seq if _seq else ""), True
            new_label_text = replace_label_keep_spacing(_cell_w_text(tr, 0),
                                                        label_seq, it["label"],
                                                        keep_ref_seq=keep_ref)
            val = it.get("value")
            if tcs:
                if is_new and len(tcs) == 1 and val not in (None, ""):
                    # 单格参考行（note 区新增，如 S16 编制依据等）：标签+值同格（note 式字段行）
                    set_cell_text(tcs[0], new_label_text + "\n" + str(val))
                else:
                    set_label_text(tcs[0], new_label_text)
        if len(tcs) >= 2:
            op_update_cell(tcs[1], it["value"], f"S{sec_no}[{it['label']}]",
                           empty_policy=empty_policy)
        final.append((tr, is_new, it))

    # 3.1) 显式删除的模板行物理移除（模板驱动下唯一删行通道）
    if del_trs:
        op_delete_rows(del_trs)

    # 4) 孤立受保护/结构行：模板有、写入项没有 → 保留（按原行号保持原位骨架）
    #    keep_structure=True 时，未匹配模板行也全部保留（模板为主）。
    #    missing_policy 控制这些"方案未给值"的模板字段行的值格处理：
    #    - preserve：保留模板原值（旧行为）；
    #    - no_data/none：值格非空（旧产品残留）→ 覆写为 missing_text，清除残留；
    #      受保护字段（preserve_labels）、纯值驱动节（_PURE_VALUE_SECTIONS，如 S11
    #      毒理省略即保留模板、不自行判数据来源）、以及多列子表行（S8 生物限值
    #      表头/数据行等，由容器重建管理）不受影响 —— 仅两列标准字段行参与填充。
    isolated = []
    for (ri, tr, seq, nl) in tmpl_rows:
        if id(tr) not in reused_ids and nl not in write_labels and tr not in del_trs:
            if keep_structure or nl in keep_labels:
                isolated.append(tr)
                if (keep_structure and missing_policy in ("no_data", "none")
                        and nl not in preserve_labels
                        and sec_no not in _PURE_VALUE_SECTIONS):
                    _tcs = tr.findall(qn("w:tc"))
                    if len(_tcs) == 2:
                        _old = _cell_w_text(tr, 1).strip()
                        if _old:
                            set_cell_text(_tcs[1], missing_text)
                            logs.append(("missing", sec_no, ri, nl, _old, missing_text))

    # 5) 删：非模板驱动（keep_structure=False）且写入项没有（不在保留清单）→ 物理删除
    #    （旧“写入项为准”行为，仅显式关闭模板驱动时生效）
    if not keep_structure:
        delete_trs = []
        for (ri, tr, seq, nl) in tmpl_rows:
            if id(tr) not in reused_ids and nl not in keep_labels and nl not in write_labels:
                delete_trs.append(tr)
                logs.append(("delete", sec_no, ri, nl, ""))
        op_delete_rows(delete_trs)

    # 6) 重建顺序（模板骨架驱动）：
    #    全部幸存模板行（复用字段 + 孤立受保护字段 + note 结构行）按原行号保持骨架；
    #    新增行按其写入顺序插到「下一个复用行」之前（无后续复用行则追加节尾）。
    #    注：旧实现只重建字段行，note 结构行（如 S8 的 8.1 暴露控制、手套材质行）
    #    会被排挤到节尾——现改为将 note 行一并纳入骨架重建，保持原位。
    base = []
    for ri in range(1, len(tbl.rows)):
        tr = tbl.rows[ri]._tr
        if tr.getparent() is None or id(tr) not in orig_idx:
            continue
        base.append((float(orig_idx[id(tr)]), tr))
    base.sort(key=lambda x: x[0])
    base_rows = [tr for (_k, tr) in base]
    base_pos = {id(tr): i for i, tr in enumerate(base_rows)}

    # 按写入顺序为每个新增行定插入目标：下一个复用行的 base 位；
    # 无后续复用行 → 追加节尾；本批无任何复用行（纯新增，如 S16 note 区）
    # → 插到节标题之后（base 最前），避免新增行挤到原 note 行之后。
    new_targets = []          # [(target_index, tr)]，同批保持写入顺序
    batch = []                # 当前批（遇到下一个复用行前的所有新增行）
    any_reused = False
    for tr, is_new, it in final:
        if is_new:
            batch.append(tr)
            continue
        any_reused = True
        bpos = base_pos.get(id(tr))
        if bpos is None:      # 防御：复用行不在 base（异常）→ 追加到末尾
            bpos = len(base_rows)
        for nt in batch:
            new_targets.append((bpos, nt))
        batch = []
    if batch:                 # 写入项末尾连续为新增行
        end = len(base_rows) if any_reused else 0
        for nt in batch:
            new_targets.append((end, nt))

    # 7) 按 target 分组插入（从后往前，避免低位插入挤动高位）；同 target 保持写入顺序。
    ordered = list(base_rows)
    from collections import defaultdict
    _groups = defaultdict(list)
    for tgt, nt in new_targets:
        _groups[tgt].append(nt)
    for tgt in sorted(_groups.keys(), reverse=True):
        for j, nt in enumerate(_groups[tgt]):
            ordered.insert(tgt + j, nt)

    # 物理重建（摘除后按序重插，避开原位错乱）
    anchor = tbl.rows[0]._tr
    prev = anchor
    for tr in ordered:
        parent = tr.getparent()
        if parent is not None:
            parent.remove(tr)
        prev.addnext(tr)
        prev = tr

    # 7.1) 新增行边框衔接修正（块首接上邻 bottom、块末接下邻 top、块内 dotted）
    new_tr_ids = {id(nt) for _t, nt in new_targets}
    fix_new_row_borders(tbl, ordered, new_tr_ids, tag=f"S{sec_no}")

    # 7.2) 首行衔接修正（删除引导行后首字段行 top 与标题底边断线）
    fix_first_row_border(tbl, ordered, tag=f"S{sec_no}")

    # 8) 日志
    for tr, is_new, it in final:
        src = it.get("src_label", it["label"])
        if is_new:
            logs.append(("insert", sec_no, src, it["value"]))
        else:
            logs.append(("update", sec_no, -1, src, "", it["value"]))
    # 隐患1防御：该节写入项大量靠“新增”插入（映射/通道未命中）→ 告警，提示核对映射
    n_write = len(mapped_items)
    n_ins = sum(1 for _t, is_new, _i in final if is_new)
    if n_write >= 3 and n_ins > n_write * 0.5:
        log(f"S{sec_no}: 写入项 {n_write} 个字段中 {n_ins} 个靠新增插入，"
            f"与模板字段差异大，请核对字段映射/标准字段清单", "WARN")
    return logs

# ------------------------- S0 页眉页脚覆写 -------------------------
def _read_s0_fields(template_path):
    """用结构读取读模板页眉页脚字段 → {norm_label: value}，用于子串替换定位。"""
    if read_msds is None:
        return {}
    try:
        r = read_msds(str(template_path))
        sd = r.sections.get(0)
        if sd is None:
            return {}
        return {norm_label(f.label): f.value for f in sd.fields}
    except Exception as e:
        log(f"S0: 读取模板页眉页脚失败 {e}", "WARN")
        return {}

def sync_section0(doc, items, template_path, empty_policy="warn"):
    """覆写页眉页脚字段（S0）—— 17 节自动化闭环的关键。
    策略（以模板格式为准）：在 header/footer 的段落与表格单元格中，
    用“模板字段原值”定位并替换为写入项值（子串替换，保留单元格前后缀
    如 'P' / '修订日期：' / '-MSDS'），格式继承原 run。
    - 页眉段落：Version
    - 页眉表格：产品名称
    - 页脚表格：公司名称 / 产品型号 / 修订日期
    - 页码为域自动生成，不覆写。
    """
    logs = []
    items = list(items or [])
    if not items:
        return logs
    tmpl = _read_s0_fields(template_path)
    wi = {norm_label(it["label"]): str(it.get("value", "")) for it in items}
    touched = []

    def do_replace(cell_obj, text, fld):
        """用模板原值定位并替换字段值；返回 (新文本, 是否变更)。"""
        old = tmpl.get(fld)
        new = wi.get(fld)
        if old is None or new is None:
            return text, False
        old_s = str(old).strip()
        if old_s and old_s in text:
            return text.replace(old_s, new), True
        return text, False

    for sec in doc.sections:
        # --- 页眉 ---
        for t in sec.header.tables:
            for row in t.rows:
                for c in row.cells:
                    text, ok = do_replace(c, c.text, "产品名称")
                    if ok:
                        set_cell_text(c, text)
                        touched.append(("页眉-产品名称", c.text, wi.get("产品名称", "")))
        for p in sec.header.paragraphs:
            if "Version" in p.text or "版本" in p.text:
                m = re.search(r"(?:Version|版本)[：:]?\s*([Vv]?\S+)", p.text)
                if m and "Version" in wi and m.group(1) in p.text:
                    old_v = m.group(1)
                    new_text = p.text.replace(old_v, wi["Version"])
                    set_paragraph_text(p._p, new_text)
                    touched.append(("页眉-Version", old_v, wi["Version"]))
        # --- 页脚 ---
        for t in sec.footer.tables:
            for row in t.rows:
                for c in row.cells:
                    text = c.text
                    for fld in ("公司名称", "产品型号", "修订日期"):
                        new_text, ok = do_replace(c, text, fld)
                        if ok:
                            set_cell_text(c, new_text)
                            touched.append((f"页脚-{fld}", text, wi.get(fld, "")))
                            text = new_text
    for desc, old, new in touched:
        logs.append(("update", 0, -1, desc, old, new))
    return logs

# ------------------------- prep 清理（模板残留，不依赖写入项） -------------------------
# 引导行 = 旧产品残留的单格说明行（如 S11/S12/S13 的"该产品无可用的…"），
# 覆写前物理删除；缺陷标签格（如 S8 手部防护标签格混入上一行值）覆写前清理为纯标签。
_GUIDE_DELETE_PATTERNS = (
    "该产品无可用的毒理学研究",
    "该产品无可用的生态毒理学研究",
    "以下为二乙二醇单丁醚",
    "必须遵守适用的国标、国家或当地法规进行废弃",
)


def prep_clean(doc):
    """覆写前清理模板残留：引导行物理删除 + 缺陷标签格清理。返回清理日志。
    引导行判定：单格通栏行且文本以已知残留特征开头（与 schema guide/prep_delete 对齐）。"""
    cleaned = []
    tables = section_tables(doc)
    for sec_no, tbl in tables.items():
        # 缺陷标签格清理（S8 手部防护标签格混入值）
        for ri in range(1, len(tbl.rows)):
            tcs = tbl.rows[ri]._tr.findall(qn("w:tc"))
            if len(tcs) < 2:
                continue
            txt = _cell_w_text(tbl.rows[ri]._tr, 0).strip()
            if sec_no == 8 and txt.startswith("手部防护") and "喷涂过程中" in txt:
                set_cell_text(tcs[0], "手部防护：")
                cleaned.append(f"S8 手部防护标签格清理（混入值移除）")
        # 引导行删除（单格通栏行）
        del_trs = []
        for ri in range(1, len(tbl.rows)):
            tcs = tbl.rows[ri]._tr.findall(qn("w:tc"))
            if len(tcs) != 1:
                continue
            txt = _cell_w_text(tbl.rows[ri]._tr, 0).strip()
            if any(txt.startswith(p) for p in _GUIDE_DELETE_PATTERNS):
                del_trs.append(tbl.rows[ri]._tr)
        if del_trs:
            op_delete_rows(del_trs)
            cleaned.append(f"S{sec_no} 引导行删除 x{len(del_trs)}")
    return cleaned


# ------------------------- 主流程 -------------------------
# 默认“模板驱动”的节：None = 全部节以模板为准（模板字段行全保留为骨架，
# 写入项只覆写值；模板多出字段保留；显式 "delete": true 才删模板行）。
# 可用写入项 JSON 顶层 "keep_structure": [2, ...]（仅指定节模板驱动）或
# CLI --keep-structure 覆盖。旧行为（写入项驱动：模板多出行默认删除）仅当
# 显式传入 keep_structure 为受限集合并配合 keep_structure=False 时生效。
DEFAULT_KEEP_STRUCTURE_SECTIONS = None   # None = 全部节默认模板驱动

def overwrite(template_path, write_items, out_path, preserve=None, sections=None,
              ref_pick=None, keep_structure=None, empty_policy=None, field_map=None,
              missing_policy=None, missing_text=None):
    """执行覆写：读模板 → 改/删/增 → 存盘。返回操作日志。
    ref_pick: 参考行策略。取值见 pick_ref_row。
      也支持在写入项 JSON 中按节覆盖：sections.<N>.ref = 'label:xxx' | 'last' | 'auto'
    keep_structure: 模板驱动节集合（None='all'=全部节：模板字段行全保留为骨架，
      写入项只覆写匹配值，模板多出字段保留，显式 "delete": true 才删模板行）。
      可用写入项 JSON 顶层 "keep_structure"（"all" 或节号列表）或参数覆盖。
    empty_policy: 空值覆写策略（overwrite/preserve/warn，见 EMPTY_POLICY_DEFAULT）。
      可用写入项 JSON 顶层 empty_policy 或参数覆盖。
    missing_policy: 模板字段行存在、写入项未给该字段时的值格处理
      （preserve 保留 / no_data 写"无数据" / none 写"无"，见 sync_section）。
      可用写入项 JSON 顶层 missing_policy 或参数覆盖；默认 preserve。
    missing_text: missing_policy 生效时的填充文本（默认"无数据"），
      可用写入项 JSON 顶层 missing_text 或参数覆盖。
    field_map: 外部字段映射配置文件路径；None 时用同目录 field_maps.json（若存在），
      否则用内置默认。隐患2：映射通道不锚定单个模板，模板更换只需换配置。
    """
    global FIELD_MAP
    if field_map:
        FIELD_MAP = _load_field_map(field_map)
    preserve = preserve or {}
    sec_objs = write_items.get("sections", {})
    # keep_structure 优先级：参数 > 写入项 JSON 顶层 > 默认（None=全部节模板驱动）
    if keep_structure is None:
        ks = write_items.get("keep_structure", DEFAULT_KEEP_STRUCTURE_SECTIONS)
        keep_structure = "all" if ks in (None, "all") else set(ks or [])
    if keep_structure == "all":
        # 全部节模板驱动：凡写入项覆盖的普通节都走模板骨架
        keep_structure = {int(s) for s in sec_objs if str(s) not in ("0",)}
    else:
        keep_structure = set(keep_structure or [])
    # empty_policy 优先级：参数 > 写入项 JSON 顶层 > 默认
    if empty_policy is None:
        empty_policy = write_items.get("empty_policy", EMPTY_POLICY_DEFAULT)
    empty_policy = str(empty_policy or EMPTY_POLICY_DEFAULT).strip().lower()
    if empty_policy not in ("overwrite", "preserve", "warn"):
        log(f"空值策略 {empty_policy!r} 无效，回退到默认 {EMPTY_POLICY_DEFAULT}", "WARN")
        empty_policy = EMPTY_POLICY_DEFAULT
    # missing_policy 优先级：参数 > 写入项 JSON 顶层 > 默认 preserve
    if missing_policy is None:
        missing_policy = write_items.get("missing_policy", "preserve")
    missing_policy = str(missing_policy or "preserve").strip().lower()
    if missing_policy not in ("preserve", "no_data", "none"):
        log(f"缺值策略 {missing_policy!r} 无效，回退到默认 preserve", "WARN")
        missing_policy = "preserve"
    if missing_text is None:
        missing_text = write_items.get("missing_text", "无数据")
    missing_text = str(missing_text or "无数据")
    if missing_policy != "preserve":
        log(f"缺值策略: {missing_policy}（模板字段方案未给 → 写 {missing_text!r}）")
    log(f"模板: {template_path}")
    doc = Document(template_path)
    tables = section_tables(doc)
    log(f"识别到 {len(tables)} 个 section 表格")

    # 用检索工具确认模板父子级 → 每节子标题标签集（结构行，不参与删除）
    sub_labels = section_sub_labels(template_path, sections=sections)

    # prep：清理模板残留（引导行删除 + 缺陷标签格清理），不依赖写入项
    prep_cleaned = prep_clean(doc)
    for pc in prep_cleaned:
        log(f"prep: {pc}")

    all_logs = []
    sec_objs = write_items.get("sections", {})

    for sec_no, payload in sec_objs.items():
        sec_no = int(sec_no)
        if sections is not None and sec_no not in sections:
            continue
        if sec_no == 0:
            # S0 页眉页脚：不在 body 表格里，走专用路径
            # 推导方案 dict（页眉_版本/页眉_产品名称/页脚_产品名称/页脚_修订日期）
            # → 引擎 S0 items（Version/产品名称/产品型号/修订日期）
            if isinstance(payload, dict):
                _s0map = {"页眉_版本": "Version", "页眉_产品名称": "产品名称",
                          "页脚_产品名称": "产品型号", "页脚_修订日期": "修订日期"}
                payload = [{"label": _s0map.get(k, k), "value": v}
                           for k, v in payload.items() if k in _s0map]
            logs = sync_section0(doc, payload, template_path, empty_policy=empty_policy)
            for lg in logs:
                kind = lg[0]
                if kind == "update":
                    log(f"S{lg[1]} 改 [{lg[3]}] {lg[4]!r} -> {lg[5]!r}")
            all_logs.extend(logs)
            continue
        if sec_no not in tables:
            log(f"Section{sec_no}：模板中不存在该节，跳过", "WARN")
            continue
        tbl = tables[sec_no]
        preserve_labels = set(preserve.get(sec_no, set())) | set(DEFAULT_PRESERVE.get(sec_no, set()))
        struct_labels = sub_labels.get(sec_no, set())
        # 参考行策略：节内 ref > 全局 ref_pick > 默认
        sec_pick = ref_pick
        if isinstance(payload, dict):
            sec_pick = payload.get("ref", sec_pick)
        elif isinstance(payload, list) and isinstance(payload, list):
            # 若节对象为 list，不支持内嵌 ref；沿用全局
            pass
        # 隐患3：keep_structure 节缺“必需字段”→ 该行保留模板默认值，需告警
        required = REQUIRED_FIELDS.get(sec_no)
        if required and sec_no in keep_structure and isinstance(payload, list):
            write_labels_sec = {norm_label(map_field_label(sec_no, it["label"]))
                                for it in payload}
            for rf in required:
                if rf not in write_labels_sec:
                    log(f"S{sec_no}: keep_structure 节缺必需字段 {rf}，"
                        f"该行保留模板默认值，请核实是否应为空值", "WARN")
        if sec_no == 3 and isinstance(payload, dict):
            logs = sync_section3(tbl, payload, preserve_labels, ref_pick=sec_pick,
                                 empty_policy=empty_policy)
        elif sec_no == 8 and isinstance(payload, dict) and \
                ("bio" in payload or "职业接触限值" in payload):
            logs = sync_section8(tbl, payload, preserve_labels, ref_pick=sec_pick,
                                 struct_labels=struct_labels, empty_policy=empty_policy,
                                 missing_policy=missing_policy, missing_text=missing_text)
        elif sec_no == 15 and isinstance(payload, dict) and "法规列表" in payload:
            logs = sync_section15(tbl, payload, preserve_labels, ref_pick=sec_pick,
                                  struct_labels=struct_labels, empty_policy=empty_policy,
                                  missing_policy=missing_policy, missing_text=missing_text)
        elif isinstance(payload, dict):
            # 推导方案 dict（{标准标签: 值}）→ items，模板驱动
            items = _dict_to_items(payload)
            logs = sync_section(tbl, sec_no, items, preserve_labels, ref_pick=sec_pick,
                                struct_labels=struct_labels,
                                keep_structure=sec_no in keep_structure,
                                empty_policy=empty_policy, missing_policy=missing_policy,
                                missing_text=missing_text)
        elif isinstance(payload, list):
            logs = sync_section(tbl, sec_no, payload, preserve_labels, ref_pick=sec_pick,
                                struct_labels=struct_labels,
                                keep_structure=sec_no in keep_structure,
                                empty_policy=empty_policy, missing_policy=missing_policy,
                                missing_text=missing_text)
        else:
            log(f"Section{sec_no}：未知写入项格式，跳过", "WARN")
            continue
        for lg in logs:
            kind = lg[0]
            if kind == "update":
                log(f"S{lg[1]} 改 row{lg[2]} [{lg[3]}] {lg[4]!r} -> {lg[5]!r}")
            elif kind == "delete":
                log(f"S{lg[1]} 删 row{lg[2]} [{lg[3]}]")
            elif kind == "insert":
                log(f"S{lg[1]} 增 [{lg[2]}] = {lg[3]!r}")
            elif kind == "insert-pending":
                log(f"S{lg[1]} 待增（后续执行）[{lg[2]}] = {lg[3]!r}")
            elif kind == "missing":
                log(f"S{lg[1]} 缺值[{lg[3]}] {lg[4]!r} -> {lg[5]!r}")
        all_logs.extend(logs)

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    log(f"已保存: {out_path}")
    return all_logs

# ------------------------- 闭环校验（调用结构读取 read_msds 读回） -------------------------
def verify_output(template_path, out_path, write_items, preserve=None, sections=None):
    """覆写后用 read_msds 读回输出文档，与写入项逐字段对比。
    sections: 若指定（如 {1,3,9}），只校验这些节；否则校验写入项中全部节。
    返回 (ok, 问题列表)。
    """
    preserve = preserve or {}
    if read_msds is None:
        log("结构读取工具不可用，跳过闭环校验", "WARN")
        return True, []
    problems = []
    out = read_msds(str(out_path))
    sec_objs = write_items.get("sections", {})
    for sec_no, payload in sec_objs.items():
        sec_no = int(sec_no)
        if sections is not None and sec_no not in sections:
            continue
        if sec_no not in out.sections:
            problems.append(f"S{sec_no}: 输出文档缺失该节")
            continue
        sd = out.sections[sec_no]
        if sec_no == 0:
            # S0 页眉页脚：读回 fields 即页眉页脚字段，逐字段对比写入项
            # 推导方案 dict（页眉_版本 等）→ items（Version 等），与 overwrite 对齐
            if isinstance(payload, dict):
                _s0map = {"页眉_版本": "Version", "页眉_产品名称": "产品名称",
                          "页脚_产品名称": "产品型号", "页脚_修订日期": "修订日期"}
                payload = [{"label": _s0map.get(k, k), "value": v}
                           for k, v in payload.items() if k in _s0map]
            found = {}
            for f in sd.fields:
                nl = norm_label(f.label)
                if nl:
                    found[nl] = f.value
            for it in payload:
                nl = norm_label(it["label"])
                if nl not in found:
                    problems.append(f"S0.{it['label']}: 输出缺失该字段")
                elif str(found[nl]).strip() != str(it.get("value", "")).strip():
                    problems.append(f"S0.{it['label']}: 期望{it['value']!r} 实际{found[nl]!r}")
            continue
        if sec_no == 3 and isinstance(payload, dict):
            # 产品类型
            found = {}
            for f in sd.fields:
                _s, label = norm_field(f.label)
                found[norm_label(label)] = f.value
            if "产品类型" in payload:
                exp = payload["产品类型"].strip()
                if found.get(norm_label("产品类型"), "").strip() != exp:
                    problems.append(f"S3.产品类型: 期望{exp!r} 实际{found.get(norm_label('产品类型'))!r}")
            # 成分（用结构读取的归一化函数对齐语义）
            exp_comps = payload.get("components", [])
            got_comps = [(c.name, c.cas, c.conc) for c in sd.components]
            for i, exp in enumerate(exp_comps):
                if i < len(got_comps):
                    got = got_comps[i]
                    exp_norm = (_norm_comp(exp.get("name", ""), normalize_component_name),
                                _norm_comp(exp.get("cas", ""), normalize_component_cas),
                                _norm_comp(exp.get("conc", ""), normalize_component_conc))
                    got_norm = (_norm_comp(got[0], normalize_component_name),
                                _norm_comp(got[1], normalize_component_cas),
                                _norm_comp(got[2], normalize_component_conc))
                    if exp_norm != got_norm:
                        problems.append(f"S3.成分{i}: 期望{exp} 实际{got}")
                else:
                    problems.append(f"S3.成分{i}: 期望{exp} 输出缺失")
            if len(got_comps) > len(exp_comps):
                problems.append(f"S3: 成分多出 {len(got_comps) - len(exp_comps)} 行")
        elif isinstance(payload, dict):
            # 推导方案 dict：容器键（S8 bio/职业接触限值、S15 法规列表）由专用
            # sync 处理，普通字段转 items 走下方 list 校验；新增字段（含 _after/_before）
            # 缺失不判 FAIL（新增为主动行为，结构读取归一化可能差异）。
            _skip = set()
            if sec_no == 8:
                _skip = {"bio", "职业接触限值"}
            elif sec_no == 15:
                _skip = {"法规列表"}
            items = []
            for _k, _v in payload.items():
                if _k in _skip:
                    continue
                # NOTE_ROW_KEYS 键（手套材质_*/生态毒性_延续/法规入口/其它的规定/免责声明）：
                # 清空语义（值空）→ 跳过校验（读回按 key 无法独立定位，清空为主动行为）
                if sec_no in NOTE_ROW_KEYS and _k in NOTE_ROW_KEYS[sec_no]:
                    _val = _v.get("value") if isinstance(_v, dict) else _v
                    if _val in (None, ""):
                        continue
                _it = {"label": _k,
                       "value": _v.get("value") if isinstance(_v, dict) else _v}
                if isinstance(_v, dict):
                    if "after" in _v:
                        _it["_after"] = _v["after"]
                    if "before" in _v:
                        _it["_before"] = _v["before"]
                items.append(_it)
            payload = items
        if isinstance(payload, list):
            found = {}
            for f in sd.fields:
                _s, label = norm_field(f.label)
                nl = norm_label(label)
                if nl:
                    # setdefault：同一标签出现多行时（如 S12 生态毒性 + 12.1 延续行，
                    # 延续行被读回为同标签第二个 field），保留 Word 顺序首个真值格
                    found.setdefault(nl, f.value)
                # 数字开头的化学品名（如 '2-丁氧基乙醇'）会被 norm_field 误拆序号，
                # 完整标签也登记，避免写入项匹配失败
                full = norm_label(f.label)
                if full:
                    found.setdefault(full, f.value)
            # 通栏/子标题行（note）读回后落在 sd.lines，也纳入匹配。
            # 单格 '标签：\n多行值' 整条在 lines 里，取首行做标签、其余做值。
            for ln in sd.lines:
                ln_s = str(ln)
                first = ln_s.split("\n")[0]
                _s, label = norm_field(first)
                nl = norm_label(label)
                if nl and nl not in found:
                    found[nl] = "\n".join(ln_s.split("\n")[1:])
            for it in payload:
                exp_label = norm_label(map_field_label(sec_no, it["label"]))
                exp_val = it["value"]
                if ("_after" in it or "_before" in it) and exp_label not in found:
                    # 新增字段：结构读取归一化差异不致命，日志提示
                    log(f"verify: S{sec_no} 新增字段 {it['label']!r} 未在读回中定位（跳过校验）", "WARN")
                    continue
                if exp_label not in found:
                    problems.append(f"S{sec_no}.{it['label']}: 输出缺失该字段")
                elif found[exp_label].strip() != exp_val.strip():
                    problems.append(f"S{sec_no}.{it['label']}: 期望{exp_val!r} 实际{found[exp_label]!r}")
            # 保留字段应仍存在
            preserve_labels = set(preserve.get(sec_no, set())) | set(DEFAULT_PRESERVE.get(sec_no, set()))
            for pl in preserve_labels:
                if norm_label(pl) not in found:
                    problems.append(f"S{sec_no}: 保留字段缺失 {pl}")
    if problems:
        return False, problems
    return True, []

# ------------------------- CLI -------------------------
def main():
    ap = argparse.ArgumentParser(description="MSDS 覆写引擎（增删查改，以写入项为准）")
    ap.add_argument("--template", required=True, help="模板 docx")
    ap.add_argument("--write-items", required=True, help="写入项 JSON（推断引擎输出契约）")
    ap.add_argument("--out", required=True, help="输出 docx 路径")
    ap.add_argument("--sections", help="仅处理指定节，如 1,3,9")
    ap.add_argument("--ref-pick", default=None,
                    help="新增行参考行策略: auto(默认)/last/label:标签名/row:行号")
    ap.add_argument("--keep-structure", default=None,
                    help="保留模板结构的节（不删除模板字段），如 2；默认含2")
    ap.add_argument("--empty-policy", default=None,
                    help="空值覆写策略: warn(默认,覆写+告警)/overwrite/preserve")
    ap.add_argument("--field-map", default=None,
                    help="字段映射配置文件 json（默认同目录 field_maps.json）")
    args = ap.parse_args()

    write_items = json.loads(Path(args.write_items).read_text(encoding="utf-8"))
    sections = None
    if args.sections:
        sections = {int(x) for x in args.sections.split(",") if x.strip()}
    keep_structure = None
    if args.keep_structure:
        keep_structure = {int(x) for x in args.keep_structure.split(",") if x.strip()}

    overwrite(args.template, write_items, args.out, sections=sections,
              ref_pick=args.ref_pick, keep_structure=keep_structure,
              empty_policy=args.empty_policy, field_map=args.field_map)
    ok, problems = verify_output(args.template, args.out, write_items, sections=sections)
    if ok:
        log("闭环校验通过 OK：覆写结果与写入项一致。")
    else:
        log("闭环校验失败：", "FAIL")
        for p in problems:
            log("  - " + p, "FAIL")
        sys.exit(1)

# ------------------------- S3 同步（产品类型 + 成分增删） -------------------------
def _clean_comp_name(name):
    """成分名称清洗：去所有空白（中文名称里的空格/半角空格会导致换行断裂）。"""
    return re.sub(r"\s+", "", name or "")

def _tc_pPr(tc):
    """返回单元格首个段落的 pPr（无则创建）。"""
    p = tc.find(qn("w:p"))
    if p is None:
        p = tc.makeelement(qn("w:p"), {})
        tc.append(p)
    pPr = p.find(qn("w:pPr"))
    if pPr is None:
        pPr = p.makeelement(qn("w:pPr"), {})
        p.insert(0, pPr)
    return pPr

def _format_component_name_cell(cell):
    """成分名称列格式：去首行缩进/右缩进 + 居中（避免长名称折行断裂）。
    覆写时只改 pPr（缩进/对齐），不破坏 rPr 格式继承。
    """
    if isinstance(cell, _Cell):
        tc = cell._tc
    else:
        tc = cell
    pPr = _tc_pPr(tc)
    # 去缩进
    ind = pPr.find(qn("w:ind"))
    if ind is not None:
        for attr in ("w:left", "w:right", "w:firstLine",
                     "w:leftChars", "w:rightChars", "w:firstLineChars"):
            ind.attrib.pop(qn(attr), None)
        if len(ind.attrib) == 0:
            pPr.remove(ind)
    # 居中
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        jc = pPr.makeelement(qn("w:jc"), {})
        pPr.append(jc)
    jc.set(qn("w:val"), "center")

def sync_section3(tbl, sec_data, preserve_labels, ref_pick=None, empty_policy="warn"):
    """S3：产品类型改 + 成分行增/删。sec_data: {'产品类型':..., 'components':[...]}
    ref_pick: 参考行策略，仅对成分行增补时使用。
    empty_policy: 空值覆写策略（见 op_update_cell）。
    成分名称列：覆写时去空白 + 居中格式（消空格/缩进防折行断裂）。
    """
    logs = []
    prod_row, comp_header, comp_start = find_component_rows(tbl)
    if prod_row is not None:
        if "产品类型" in sec_data:
            cell = tbl.rows[prod_row].cells[1]
            old = cell.text
            op_update_cell(cell, sec_data["产品类型"], "S3.产品类型", empty_policy=empty_policy)
            logs.append(("update", 3, prod_row, "产品类型", old, sec_data["产品类型"]))
    components = sec_data.get("components", [])
    if comp_start is None:
        log("S3 未定位到成分表头，跳过成分同步", "WARN")
        return logs
    n_tmpl = len(tbl.rows) - comp_start
    n_want = len(components)
    # 改：公共部分逐格覆写
    for i in range(min(n_tmpl, n_want)):
        comp = components[i]
        row = tbl.rows[comp_start + i]
        old_name = row.cells[0].text
        for ci, key in ((0, "name"), (1, "cas"), (2, "conc")):
            cell = row.cells[ci]
            if ci == 0:
                val = _clean_comp_name(comp.get(key, ""))
            else:
                val = comp.get(key, "")
            op_update_cell(cell, val, f"S3成分[{i}].{key}", empty_policy=empty_policy)
        _format_component_name_cell(row.cells[0])
        new_name = row.cells[0].text
        logs.append(("update", 3, comp_start + i, f"成分{i}", old_name, new_name))
    # 删：模板成分行多于写入项
    if n_tmpl > n_want:
        delete_trs = [tbl.rows[comp_start + i]._tr for i in range(n_want, n_tmpl)]
        op_delete_rows(delete_trs)
        logs.append(("delete", 3, -1, f"成分行x{n_tmpl - n_want}", ""))
    # 增：写入项成分行多于模板
    if n_want > n_tmpl:
        comp_rows_seq = [(None, tbl.rows[comp_start + i]._tr) for i in range(n_tmpl)]
        if comp_rows_seq:
            ref_tr = pick_ref_row(tbl, comp_rows_seq, None, ref_pick)
        else:
            ref_tr = tbl.rows[comp_header]._tr
        for i in range(n_tmpl, n_want):
            comp = components[i]
            new_tr = copy.deepcopy(ref_tr)
            tcs = new_tr.findall(qn("w:tc"))
            for ci, key in ((0, "name"), (1, "cas"), (2, "conc")):
                if ci < len(tcs):
                    val = _clean_comp_name(comp.get(key, "")) if ci == 0 else comp.get(key, "")
                    ref_text = _cell_w_text(new_tr, ci)
                    set_cell_text(tcs[ci], replace_value_keep_spacing(ref_text, val))
            # 新增行名称列也居中
            if tcs:
                _format_component_name_cell(tcs[0])
            ref_tr.addnext(new_tr)
            ref_tr = new_tr
            logs.append(("insert", 3, f"成分{i}", comp.get("name", "")))
    return logs

# ------------------------- 推导方案 dict 适配 + 专用容器节 -------------------------
def _dict_to_items(payload):
    """推导方案 dict（{标准标签: 值}）→ items list（引擎标准格式）。
    值可为：
      - str: 覆写（'' 清空）
      - {"value":.., "after":锚点|"before":锚点}: 新增字段
    返回顺序 = dict 插入顺序（新增字段已按锚点排列，引擎按写入顺序插入到
    「下一个复用行」之前，从而满足 after/before 锚点语义）。
    """
    items = []
    for k, v in payload.items():
        if isinstance(v, dict):
            it = {"label": k, "value": v.get("value")}
            if "after" in v:
                it["_after"] = v["after"]
            if "before" in v:
                it["_before"] = v["before"]
            items.append(it)
        else:
            items.append({"label": k, "value": v})
    return items


def sync_section8(tbl, payload, preserve_labels, ref_pick=None, empty_policy="warn",
                  struct_labels=None, missing_policy="preserve", missing_text="无数据"):
    """S8 专用同步：普通防护字段（走 sync_section）+ bio 生物限值子表重建
    + 职业接触限值子表新增。
    payload: {'呼吸系统防护':..., '手部防护':..., ..., 'bio':[占位行],
              '职业接触限值':{subtitle,headers,spans,rows,before|after}}
    """
    logs = []
    containers = ("bio", "职业接触限值")
    items = []
    for k, v in payload.items():
        if k in containers:
            continue
        if isinstance(v, dict):
            items.append({"label": k, "value": v.get("value")})
        else:
            items.append({"label": k, "value": v})
    logs += sync_section(tbl, 8, items, preserve_labels, ref_pick=ref_pick,
                         struct_labels=struct_labels or set(), keep_structure=True,
                         empty_policy=empty_policy, missing_policy=missing_policy,
                         missing_text=missing_text)

    new_tr_ids = set()

    # ---- bio 生物限值子表重建 ----
    if "bio" in payload:
        bio = payload["bio"] or []
        head_idx = None
        for ri in range(1, len(tbl.rows)):
            _s, label = norm_field(tbl.rows[ri].cells[0].text)
            if norm_label(label) == "组分名称":
                head_idx = ri
                break
        if head_idx is None:
            log("S8: 未定位到生物限值表头（组分名称），跳过 bio 重建", "WARN")
        else:
            del_trs = []
            for ri in range(head_idx + 1, len(tbl.rows)):
                tcs = tbl.rows[ri]._tr.findall(qn("w:tc"))
                if len(tcs) < 2:
                    break
                _s, label = norm_field(tbl.rows[ri].cells[0].text)
                nl = norm_label(label)
                if nl in ("工程控制", "职业接触限值", "组分名称") or re.match(r"^\d", _s or ""):
                    break
                del_trs.append(tbl.rows[ri]._tr)
            op_delete_rows(del_trs)
            if del_trs:
                logs.append(("delete", 8, -1, f"生物限值数据行x{len(del_trs)}", ""))
            head_tr = tbl.rows[head_idx]._tr
            prev = head_tr
            keys = ("组分名称", "标准来源", "生物监测指标", "生物限值", "采样时间")
            for i, row in enumerate(bio):
                new_tr = copy.deepcopy(head_tr)
                tcs = new_tr.findall(qn("w:tc"))
                for ci, key in enumerate(keys):
                    if ci < len(tcs):
                        val = row.get(key, "")
                        ref = _cell_w_text(new_tr, ci)
                        set_cell_text(tcs[ci], replace_value_keep_spacing(ref, val))
                prev.addnext(new_tr)
                prev = new_tr
                new_tr_ids.add(id(new_tr))
                logs.append(("insert", 8, f"生物限值[{i}]", row.get("组分名称", "")))

    # ---- 职业接触限值子表新增 ----
    if "职业接触限值" in payload:
        sub = payload["职业接触限值"]
        anchor_label = sub.get("before") or sub.get("after")
        use_before = "before" in sub
        if not anchor_label:
            log("S8: 职业接触限值子表缺少 before/after 锚点", "WARN")
        else:
            anchor_tr = None
            sub_ref = None
            head_ref = None
            for ri in range(1, len(tbl.rows)):
                _s, label = norm_field(tbl.rows[ri].cells[0].text)
                nl = norm_label(label)
                if anchor_tr is None and nl == norm_label(anchor_label):
                    anchor_tr = tbl.rows[ri]._tr
                if sub_ref is None and nl == "暴露控制":
                    sub_ref = tbl.rows[ri]._tr
                if head_ref is None and nl == "组分名称":
                    head_ref = tbl.rows[ri]._tr
            if anchor_tr is None:
                log(f"S8: 职业接触限值子表锚点 {anchor_label!r} 未找到", "WARN")
            else:
                parts = []
                # subtitle 行（克隆 8.1 暴露控制 单格通栏行）
                if sub_ref is not None:
                    sub_tr = copy.deepcopy(sub_ref)
                    tcs = sub_tr.findall(qn("w:tc"))
                    if tcs:
                        set_cell_text(tcs[0], sub.get("subtitle", "职业接触限值："))
                    parts.append(sub_tr)
                # 表头行 + 数据行（克隆生物限值表头 5 列行）
                ref5 = head_ref if head_ref is not None else anchor_tr
                head_tr = copy.deepcopy(ref5)
                hd_tcs = head_tr.findall(qn("w:tc"))
                headers = sub.get("headers", [])
                for ci, h in enumerate(headers):
                    if ci < len(hd_tcs):
                        set_cell_text(hd_tcs[ci], replace_value_keep_spacing(
                            _cell_w_text(head_tr, ci), h))
                parts.append(head_tr)
                for i, row in enumerate(sub.get("rows", [])):
                    d_tr = copy.deepcopy(head_tr)
                    d_tcs = d_tr.findall(qn("w:tc"))
                    for ci, h in enumerate(headers):
                        if ci < len(d_tcs):
                            val = row.get(h, "")
                            set_cell_text(d_tcs[ci], replace_value_keep_spacing(
                                _cell_w_text(d_tr, ci), val))
                    parts.append(d_tr)
                # 插入到锚点前/后（保持 parts 顺序）
                if use_before:
                    cur = anchor_tr.getprevious()
                    for pt in parts:
                        if cur is None:
                            anchor_tr.addprevious(pt)
                            cur = pt
                        else:
                            cur.addnext(pt)
                            cur = pt
                else:
                    cur = anchor_tr
                    for pt in parts:
                        cur.addnext(pt)
                        cur = pt
                for pt in parts:
                    new_tr_ids.add(id(pt))
                    if pt is parts[0]:
                        logs.append(("insert", 8, "职业接触限值子标题", sub.get("subtitle", "")))
                    elif pt is parts[1]:
                        logs.append(("insert", 8, "职业接触限值表头", ",".join(headers)))
                    else:
                        logs.append(("insert", 8, "职业接触限值行", ""))

    if new_tr_ids:
        ordered = [tbl.rows[0]._tr] + [tbl.rows[i]._tr for i in range(1, len(tbl.rows))]
        fix_new_row_borders(tbl, ordered, new_tr_ids, tag="S8")
    return logs


def sync_section15(tbl, payload, preserve_labels, ref_pick=None, empty_policy="warn",
                   struct_labels=None, missing_policy="preserve", missing_text="无数据"):
    """S15 法规信息：普通 note 覆写（走 sync_section）+ 法规列表数组重建。
    payload: {'其它的规定':..., '法规列表':[...]}
    法规列表重建：保留入口行'符合下列法规要求：'，删除其后原法规单格行，
    按列表长度克隆入口行样式逐条写入。
    """
    logs = []
    items = []
    for k, v in payload.items():
        if k == "法规列表":
            continue
        if isinstance(v, dict):
            items.append({"label": k, "value": v.get("value")})
        else:
            items.append({"label": k, "value": v})
    logs += sync_section(tbl, 15, items, preserve_labels, ref_pick=ref_pick,
                         struct_labels=struct_labels or set(), keep_structure=True,
                         empty_policy=empty_policy, missing_policy=missing_policy,
                         missing_text=missing_text)
    if "法规列表" not in payload:
        return logs
    laws = payload["法规列表"] or []
    entry = None
    for ri in range(1, len(tbl.rows)):
        if "符合下列法规要求" in _cell_w_text(tbl.rows[ri]._tr, 0):
            entry = tbl.rows[ri]._tr
            break
    if entry is None:
        log("S15: 未定位到'符合下列法规要求'入口行，跳过法规重建", "WARN")
        return logs
    n = _delete_note_value_rows(entry, tbl)
    if n:
        logs.append(("delete", 15, -1, f"法规原条目行x{n}", ""))
    prev = entry
    new_tr_ids = []
    for law in laws:
        new_tr = copy.deepcopy(entry)
        tcs = new_tr.findall(qn("w:tc"))
        if tcs:
            set_cell_text(tcs[0], law)
        prev.addnext(new_tr)
        prev = new_tr
        new_tr_ids.append(new_tr)
        logs.append(("insert", 15, "法规条目", law))
    if new_tr_ids:
        ordered = [tbl.rows[0]._tr] + [tbl.rows[i]._tr for i in range(1, len(tbl.rows))]
        fix_new_row_borders(tbl, ordered, {id(t) for t in new_tr_ids}, tag="S15")
    return logs


if __name__ == "__main__":
    main()

