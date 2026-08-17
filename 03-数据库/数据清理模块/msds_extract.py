#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
批量提取文件夹内 MSDS/SDS 的成分信息并汇总到 Excel。

默认支持：
- PDF：可复制文字的 PDF；同时尝试提取表格
- DOCX：正文与表格
- XLSX/XLSM：工作表表格

默认输出字段：
产品、成份、标准、含量、来源文件、页码/工作表、提取方式、需复核、复核原因、原始行

注意：
1. “标准”默认解释为 CAS号 / EC号 / 标准号 / 标识号。
2. 扫描版 PDF 需要先 OCR，或在本程序前增加 OCR 模块。
3. 不同供应商的 MSDS 模板差异很大，建议先用 20~50 份样本调试表头别名。
"""

from __future__ import annotations

import argparse
import logging
import re
import sys
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Iterable, Sequence

try:
    import pymupdf
except ImportError as exc:
    raise SystemExit("缺少 PyMuPDF：pip install pymupdf") from exc

try:
    from docx import Document
except ImportError as exc:
    raise SystemExit("缺少 python-docx：pip install python-docx") from exc

try:
    from openpyxl import Workbook, load_workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter
except ImportError as exc:
    raise SystemExit("缺少 openpyxl：pip install openpyxl") from exc


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".xlsm"}

SECTION3_START_PATTERNS = [
    re.compile(
        r"(?:第\s*3\s*(?:部分|节)|第三部分)\s*[:：.\-—]?\s*"
        r"(?:成分|组成|组分|成份)[^\r\n]{0,50}",
        re.IGNORECASE,
    ),
    re.compile(
        r"\bSECTION\s*3\b\s*[:：.\-—]?\s*"
        r"(?:COMPOSITION|INGREDIENTS?|COMPOSITION\s*/\s*INFORMATION)",
        re.IGNORECASE,
    ),
]

SECTION4_PATTERN = re.compile(
    r"(?:\r?\n)\s*(?:第\s*4\s*(?:部分|节)|第四部分|\bSECTION\s*4\b)",
    re.IGNORECASE,
)

PRODUCT_PATTERNS = [
    re.compile(
        r"(?:产品名称|化学品名称|商品名称|物料名称|品名)\s*[:：]?\s*"
        r"([^\r\n|]{2,120})",
        re.IGNORECASE,
    ),
    re.compile(
        r"(?:Product\s*(?:name|identifier)|Trade\s*name|Material\s*name)"
        r"\s*[:：]?\s*([^\r\n|]{2,120})",
        re.IGNORECASE,
    ),
]

INGREDIENT_HEADERS = (
    "化学品名称",
    "化学名称",
    "成分名称",
    "组分名称",
    "成份名称",
    "成份",
    "成分",
    "组分",
    "物质名称",
    "ingredient",
    "chemicalname",
    "component",
    "substance",
)

STANDARD_HEADERS = (
    "casno",
    "cas号",
    "cas编号",
    "casnumber",
    "化学文摘社登记号",
    "ecno",
    "ec号",
    "einecs",
    "标准号",
    "标准",
    "标识号",
    "识别号",
    "编号",
    "identifier",
)

CONTENT_HEADERS = (
    "含量",
    "浓度范围",
    "浓度",
    "百分比",
    "质量分数",
    "质量百分比",
    "比例",
    "content",
    "concentration",
    "weightpercent",
    "weight%",
    "w/w",
    "range",
)

CAS_RE = re.compile(r"\b\d{2,7}-\d{2}-\d\b")
EC_RE = re.compile(r"\b\d{3}-\d{3}-\d\b")
PERCENT_RE = re.compile(
    r"(?P<value>"
    r"(?:[<>≤≥]\s*)?\d+(?:\.\d+)?"
    r"(?:\s*(?:-|–|—|~|～|至|to)\s*(?:[<>≤≥]\s*)?\d+(?:\.\d+)?)?"
    r"\s*(?:%|wt\.?\s*%|w/w|质量\s*%)"
    r")",
    re.IGNORECASE,
)
TAIL_CONTENT_RE = re.compile(
    r"(?P<value>"
    r"(?:[<>≤≥]\s*)?\d+(?:\.\d+)?"
    r"(?:\s*(?:-|–|—|~|～|至|to)\s*(?:[<>≤≥]\s*)?\d+(?:\.\d+)?)?"
    r"\s*%?"
    r")\s*$",
    re.IGNORECASE,
)


@dataclass
class TableBlock:
    location: str
    rows: list[list[str]]
    context: str = ""


@dataclass
class ExtractedRow:
    product: str
    ingredient: str
    standard: str
    content: str
    source_file: str
    location: str
    method: str
    needs_review: str
    review_reason: str
    raw_row: str


@dataclass
class ExceptionRow:
    source_file: str
    status: str
    detail: str


def clean_cell(value: object) -> str:
    if value is None:
        return ""
    text = str(value).replace("\u00a0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def compact_header(value: str) -> str:
    return re.sub(r"[\s:：/\\()（）\[\]【】._\-—]+", "", value).lower()


def normalize_content(value: str) -> str:
    value = clean_cell(value)
    value = value.replace("％", "%")
    value = re.sub(r"\s+", " ", value)
    value = re.sub(r"\s*([~～\-–—])\s*", r"\1", value)
    return value.strip(" ;；,，")


def normalize_standard(value: str) -> str:
    value = clean_cell(value)
    value = value.replace("CAS No.", "CAS").replace("CAS NO.", "CAS")
    value = re.sub(r"\s+", " ", value)
    return value.strip(" ;；,，")


def normalize_ingredient(value: str) -> str:
    value = clean_cell(value)
    value = re.sub(r"^\s*(?:\d+[.)、]|[-•·])\s*", "", value)
    return value.strip(" ;；,，|")


def extract_product_name(full_text: str, path: Path) -> str:
    search_text = full_text[:12000]
    for pattern in PRODUCT_PATTERNS:
        match = pattern.search(search_text)
        if match:
            candidate = clean_cell(match.group(1))
            candidate = re.split(
                r"\s{2,}|(?:建议用途|推荐用途|用途|Product\s*code|产品代码)",
                candidate,
                maxsplit=1,
                flags=re.IGNORECASE,
            )[0]
            if 1 < len(candidate) <= 120:
                return candidate
    return path.stem


def extract_section3(full_text: str) -> str:
    if not full_text:
        return ""
    starts = []
    for pattern in SECTION3_START_PATTERNS:
        match = pattern.search(full_text)
        if match:
            starts.append(match.start())
    if not starts:
        return ""
    start = min(starts)
    tail = full_text[start:]
    end_match = SECTION4_PATTERN.search(tail)
    return tail[: end_match.start()] if end_match else tail[:15000]


def header_field(cell: str) -> str | None:
    compact = compact_header(cell)
    if any(alias in compact for alias in INGREDIENT_HEADERS):
        return "ingredient"
    if any(alias in compact for alias in CONTENT_HEADERS):
        return "content"
    if any(alias in compact for alias in STANDARD_HEADERS):
        return "standard"
    return None


def find_header_mapping(rows: Sequence[Sequence[str]]) -> tuple[int, dict[str, int]] | None:
    """
    在前 12 行寻找表头。至少需要识别“成份”和“含量”两列。
    “标准”列可选；缺失时会从原始行中识别 CAS/EC 号。
    """
    for row_index, row in enumerate(rows[:12]):
        mapping: dict[str, int] = {}
        for col_index, cell in enumerate(row):
            field = header_field(clean_cell(cell))
            if field and field not in mapping:
                mapping[field] = col_index

        if "ingredient" in mapping and "content" in mapping:
            return row_index, mapping

        # 某些表头分成上下两行，尝试拼接当前行和下一行。
        if row_index + 1 < min(len(rows), 12):
            next_row = rows[row_index + 1]
            merged_mapping: dict[str, int] = {}
            max_cols = max(len(row), len(next_row))
            for col_index in range(max_cols):
                upper = clean_cell(row[col_index]) if col_index < len(row) else ""
                lower = clean_cell(next_row[col_index]) if col_index < len(next_row) else ""
                field = header_field(f"{upper} {lower}")
                if field and field not in merged_mapping:
                    merged_mapping[field] = col_index
            if "ingredient" in merged_mapping and "content" in merged_mapping:
                return row_index + 1, merged_mapping
    return None


def get_cell(row: Sequence[str], index: int | None) -> str:
    if index is None or index >= len(row):
        return ""
    return clean_cell(row[index])


def extract_identifiers(text: str) -> str:
    identifiers: list[str] = []
    for value in CAS_RE.findall(text):
        if value not in identifiers:
            identifiers.append(value)
    for value in EC_RE.findall(text):
        if value not in identifiers:
            identifiers.append(value)
    return "; ".join(identifiers)


def valid_cas_checksum(cas: str) -> bool:
    """
    CAS 校验：从校验位左侧开始，数字由右向左分别乘 1、2、3...，
    求和后对 10 取模，应等于最后一位。
    """
    if not CAS_RE.fullmatch(cas):
        return False
    digits = cas.replace("-", "")
    body, check_digit = digits[:-1], int(digits[-1])
    total = sum(int(digit) * multiplier for multiplier, digit in enumerate(reversed(body), 1))
    return total % 10 == check_digit


def review_flags(ingredient: str, standard: str, content: str, method: str) -> tuple[str, str]:
    reasons: list[str] = []
    if not ingredient:
        reasons.append("成份为空")
    if not content:
        reasons.append("含量为空")
    if not standard:
        reasons.append("标准/编号为空")

    cas_numbers = CAS_RE.findall(standard)
    invalid_cas = [cas for cas in cas_numbers if not valid_cas_checksum(cas)]
    if invalid_cas:
        reasons.append("CAS校验失败：" + ",".join(invalid_cas))

    if method == "文本正则":
        reasons.append("文本正则提取，建议抽查")

    return ("是", "；".join(reasons)) if reasons else ("否", "")


def row_is_section_heading(text: str) -> bool:
    return bool(
        re.search(
            r"(?:第\s*[4-9]\s*(?:部分|节)|第[四五六七八九十]+部分|\bSECTION\s*[4-9]\b)",
            text,
            re.IGNORECASE,
        )
    )


def parse_table_block(
    block: TableBlock,
    product: str,
    source_file: Path,
) -> list[ExtractedRow]:
    mapping_result = find_header_mapping(block.rows)
    if not mapping_result:
        return []

    header_index, mapping = mapping_result
    results: list[ExtractedRow] = []

    for row in block.rows[header_index + 1 :]:
        cleaned = [clean_cell(cell) for cell in row]
        raw = " | ".join(cleaned).strip(" |")
        if not raw:
            continue
        if row_is_section_heading(raw):
            break

        ingredient = normalize_ingredient(get_cell(cleaned, mapping.get("ingredient")))
        content = normalize_content(get_cell(cleaned, mapping.get("content")))
        standard = normalize_standard(get_cell(cleaned, mapping.get("standard")))

        if not standard:
            standard = extract_identifiers(raw)

        # 防止把“合计/备注/说明”当作成份。
        if compact_header(ingredient) in {"合计", "总计", "备注", "说明", "note", "total"}:
            continue

        # 若成份、标准、含量全空，则不是有效数据行。
        if not any((ingredient, standard, content)):
            continue

        needs_review, reason = review_flags(ingredient, standard, content, "表格")
        results.append(
            ExtractedRow(
                product=product,
                ingredient=ingredient,
                standard=standard,
                content=content,
                source_file=str(source_file),
                location=block.location,
                method="表格",
                needs_review=needs_review,
                review_reason=reason,
                raw_row=raw,
            )
        )
    return results


def parse_section_text(
    section_text: str,
    product: str,
    source_file: Path,
) -> list[ExtractedRow]:
    """
    表格提取失败后的兜底方案。
    对形如“乙醇 64-17-5 10-20%”的行效果较好。
    """
    results: list[ExtractedRow] = []
    for line_number, raw_line in enumerate(section_text.splitlines(), 1):
        line = clean_cell(raw_line)
        if len(line) < 4 or row_is_section_heading(line):
            continue

        compact = compact_header(line)
        if any(
            keyword in compact
            for keyword in (
                "成分组成信息",
                "成份组成信息",
                "compositioninformation",
                "化学品名称cas含量",
                "组分名称cas含量",
            )
        ):
            continue

        identifiers = extract_identifiers(line)
        percent_match = PERCENT_RE.search(line)
        content = normalize_content(percent_match.group("value")) if percent_match else ""

        # 有 CAS/EC 但没有百分号时，尝试提取行尾数值或范围。
        working = CAS_RE.sub(" ", line)
        working = EC_RE.sub(" ", working)
        if not content and identifiers:
            tail_match = TAIL_CONTENT_RE.search(working)
            if tail_match:
                content = normalize_content(tail_match.group("value"))

        # 文本行至少需要“编号或百分比”之一，避免误抓普通说明。
        if not identifiers and not content:
            continue

        ingredient_text = line
        ingredient_text = CAS_RE.sub(" ", ingredient_text)
        ingredient_text = EC_RE.sub(" ", ingredient_text)
        if percent_match:
            ingredient_text = (
                ingredient_text[: percent_match.start()]
                + " "
                + ingredient_text[percent_match.end() :]
            )
        elif content:
            ingredient_text = re.sub(re.escape(content) + r"\s*$", " ", ingredient_text)

        ingredient = normalize_ingredient(ingredient_text)
        ingredient = re.sub(
            r"\b(?:CAS|EC|EINECS|含量|浓度|content|concentration)\b\s*[:：]?",
            " ",
            ingredient,
            flags=re.IGNORECASE,
        )
        ingredient = normalize_ingredient(ingredient)

        if len(ingredient) < 2 or len(ingredient) > 180:
            continue

        needs_review, reason = review_flags(ingredient, identifiers, content, "文本正则")
        results.append(
            ExtractedRow(
                product=product,
                ingredient=ingredient,
                standard=identifiers,
                content=content,
                source_file=str(source_file),
                location=f"第3节文本第{line_number}行",
                method="文本正则",
                needs_review=needs_review,
                review_reason=reason,
                raw_row=line,
            )
        )
    return results


def read_pdf(path: Path) -> tuple[str, list[TableBlock]]:
    text_parts: list[str] = []
    tables: list[TableBlock] = []
    with pymupdf.open(path) as document:
        for page_index, page in enumerate(document, 1):
            page_text = page.get_text("text", sort=True) or ""
            text_parts.append(page_text)

            try:
                finder = page.find_tables()
                for table_index, table in enumerate(finder.tables, 1):
                    matrix = table.extract()
                    rows = [[clean_cell(cell) for cell in row] for row in matrix]
                    tables.append(
                        TableBlock(
                            location=f"第{page_index}页-表{table_index}",
                            rows=rows,
                            context=page_text,
                        )
                    )
            except Exception as exc:  # 表格识别失败不影响正文提取
                logging.debug("PDF 表格识别失败 %s 第%s页：%s", path.name, page_index, exc)

    return "\n".join(text_parts), tables


def read_docx(path: Path) -> tuple[str, list[TableBlock]]:
    document = Document(path)
    text_parts = [clean_cell(paragraph.text) for paragraph in document.paragraphs]
    tables: list[TableBlock] = []

    for table_index, table in enumerate(document.tables, 1):
        rows = [
            [clean_cell(cell.text) for cell in row.cells]
            for row in table.rows
        ]
        tables.append(
            TableBlock(
                location=f"Word表{table_index}",
                rows=rows,
                context="",
            )
        )
    return "\n".join(text_parts), tables


def read_excel(path: Path) -> tuple[str, list[TableBlock]]:
    workbook = load_workbook(path, read_only=True, data_only=True)
    text_parts: list[str] = []
    tables: list[TableBlock] = []

    for worksheet in workbook.worksheets:
        rows: list[list[str]] = []
        for values in worksheet.iter_rows(values_only=True):
            cleaned = [clean_cell(value) for value in values]
            if any(cleaned):
                rows.append(cleaned)
                text_parts.append(" | ".join(cleaned))
        if rows:
            tables.append(
                TableBlock(
                    location=f"工作表:{worksheet.title}",
                    rows=rows,
                    context="",
                )
            )
    workbook.close()
    return "\n".join(text_parts), tables


def read_document(path: Path) -> tuple[str, list[TableBlock]]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return read_pdf(path)
    if suffix == ".docx":
        return read_docx(path)
    if suffix in {".xlsx", ".xlsm"}:
        return read_excel(path)
    raise ValueError(f"不支持的文件类型：{suffix}")


def deduplicate(rows: Iterable[ExtractedRow]) -> list[ExtractedRow]:
    seen: set[tuple[str, ...]] = set()
    result: list[ExtractedRow] = []
    for row in rows:
        key = (
            compact_header(row.product),
            compact_header(row.ingredient),
            compact_header(row.standard),
            compact_header(row.content),
            row.source_file,
            row.location,
        )
        if key not in seen:
            seen.add(key)
            result.append(row)
    return result


def collect_files(input_folder: Path, output_file: Path) -> list[Path]:
    files = [
        path
        for path in input_folder.rglob("*")
        if path.is_file()
        and path.suffix.lower() in SUPPORTED_EXTENSIONS
        and path.resolve() != output_file.resolve()
        and not path.name.startswith("~$")
    ]
    return sorted(files)


def process_folder(
    input_folder: Path,
    output_file: Path,
) -> tuple[list[ExtractedRow], list[ExceptionRow]]:
    all_rows: list[ExtractedRow] = []
    exceptions: list[ExceptionRow] = []
    files = collect_files(input_folder, output_file)

    if not files:
        raise FileNotFoundError(f"未找到支持的文件：{input_folder}")

    for index, path in enumerate(files, 1):
        logging.info("[%s/%s] 正在处理：%s", index, len(files), path.name)
        try:
            full_text, table_blocks = read_document(path)
            product = extract_product_name(full_text, path)

            table_rows: list[ExtractedRow] = []
            for block in table_blocks:
                table_rows.extend(parse_table_block(block, product, path))

            section3 = extract_section3(full_text)
            text_rows = parse_section_text(section3, product, path) if section3 else []

            rows = deduplicate([*table_rows, *text_rows])
            all_rows.extend(rows)

            if not rows:
                if path.suffix.lower() == ".pdf" and len(re.sub(r"\s+", "", full_text)) < 80:
                    status = "需要OCR"
                    detail = "PDF几乎没有可提取文字，可能是扫描件"
                elif not section3 and not table_rows:
                    status = "未找到第3节"
                    detail = "未识别到成分/组成信息标题或成分表头"
                else:
                    status = "未提取到成分"
                    detail = "已读取文档，但没有形成有效成分行"
                exceptions.append(ExceptionRow(str(path), status, detail))

        except Exception as exc:
            logging.exception("处理失败：%s", path)
            exceptions.append(ExceptionRow(str(path), "处理失败", str(exc)))

    return deduplicate(all_rows), exceptions


def autosize_columns(worksheet, widths: dict[int, int]) -> None:
    for column_index, width in widths.items():
        worksheet.column_dimensions[get_column_letter(column_index)].width = width


def write_output(
    rows: Sequence[ExtractedRow],
    exceptions: Sequence[ExceptionRow],
    output_file: Path,
) -> None:
    workbook = Workbook()
    summary = workbook.active
    summary.title = "MSDS汇总"

    headers = [
        "产品",
        "成份",
        "标准",
        "含量",
        "来源文件",
        "页码/工作表",
        "提取方式",
        "需复核",
        "复核原因",
        "原始行",
    ]
    summary.append(headers)

    for item in rows:
        data = asdict(item)
        summary.append(
            [
                data["product"],
                data["ingredient"],
                data["standard"],
                data["content"],
                data["source_file"],
                data["location"],
                data["method"],
                data["needs_review"],
                data["review_reason"],
                data["raw_row"],
            ]
        )

    header_fill = PatternFill("solid", fgColor="1F4E78")
    review_fill = PatternFill("solid", fgColor="FFF2CC")
    error_fill = PatternFill("solid", fgColor="FCE4D6")
    header_font = Font(color="FFFFFF", bold=True)

    for cell in summary[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")

    summary.freeze_panes = "A2"
    summary.auto_filter.ref = f"A1:J{max(summary.max_row, 1)}"
    summary.sheet_view.showGridLines = False

    for row_index in range(2, summary.max_row + 1):
        summary.cell(row_index, 2).alignment = Alignment(wrap_text=True, vertical="top")
        summary.cell(row_index, 5).alignment = Alignment(wrap_text=True, vertical="top")
        summary.cell(row_index, 9).alignment = Alignment(wrap_text=True, vertical="top")
        summary.cell(row_index, 10).alignment = Alignment(wrap_text=True, vertical="top")
        if summary.cell(row_index, 8).value == "是":
            for column_index in range(1, 11):
                summary.cell(row_index, column_index).fill = review_fill

    autosize_columns(
        summary,
        {
            1: 24,
            2: 30,
            3: 22,
            4: 16,
            5: 45,
            6: 20,
            7: 12,
            8: 10,
            9: 36,
            10: 60,
        },
    )

    exception_sheet = workbook.create_sheet("异常与待处理")
    exception_sheet.append(["来源文件", "状态", "说明"])
    for item in exceptions:
        exception_sheet.append([item.source_file, item.status, item.detail])

    for cell in exception_sheet[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")

    exception_sheet.freeze_panes = "A2"
    exception_sheet.auto_filter.ref = f"A1:C{max(exception_sheet.max_row, 1)}"
    exception_sheet.sheet_view.showGridLines = False
    autosize_columns(exception_sheet, {1: 55, 2: 18, 3: 60})

    for row_index in range(2, exception_sheet.max_row + 1):
        for column_index in range(1, 4):
            exception_sheet.cell(row_index, column_index).fill = error_fill
            exception_sheet.cell(row_index, column_index).alignment = Alignment(
                wrap_text=True, vertical="top"
            )

    output_file.parent.mkdir(parents=True, exist_ok=True)
    workbook.save(output_file)


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="批量提取 MSDS/SDS 成分并汇总到 Excel"
    )
    parser.add_argument("input_folder", type=Path, help="MSDS 文件夹路径")
    parser.add_argument(
        "-o",
        "--output",
        type=Path,
        default=Path("MSDS成分汇总.xlsx"),
        help="输出 Excel 文件，默认：MSDS成分汇总.xlsx",
    )
    parser.add_argument(
        "--log-level",
        choices=("DEBUG", "INFO", "WARNING", "ERROR"),
        default="INFO",
        help="日志级别，默认 INFO",
    )
    return parser


def main() -> int:
    args = build_parser().parse_args()
    logging.basicConfig(
        level=getattr(logging, args.log_level),
        format="%(levelname)s %(message)s",
    )

    input_folder = args.input_folder.expanduser().resolve()
    output_file = args.output.expanduser().resolve()

    if not input_folder.is_dir():
        logging.error("输入文件夹不存在：%s", input_folder)
        return 2

    try:
        rows, exceptions = process_folder(input_folder, output_file)
        write_output(rows, exceptions, output_file)
    except Exception as exc:
        logging.exception("程序执行失败：%s", exc)
        return 1

    logging.info(
        "完成：提取 %s 条成分记录，%s 个异常文件。输出：%s",
        len(rows),
        len(exceptions),
        output_file,
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
