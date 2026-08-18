# -*- coding: utf-8 -*-
"""Section 3 穿透式底层深度检索脚本.

直接扫描 255 份 docx 原文档的所有段落、表格、合并单元格、页眉页脚及附注文本，
查找 Section 3 中的每一个化学成分、CAS 号、游离单体、添加剂及特殊表述，
与现有数据库和飞书文档做 100% 交叉比对。
"""

import docx
import re
import sys
from pathlib import Path
from collections import defaultdict, Counter

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from core.docx_reader import read_msds
from core.msds_db import open_db

DB_PATH = r"F:\正式项目与模块化内容\冠志\MSDS\Word 覆写模块\数据库\正式库\Data Base\msds_standard.db"
DST_DIR = Path(r"F:\正式项目与模块化内容\冠志\MSDS\Word 覆写模块\数据库\正式库\入库word  第一批")

CAS_REGEX = re.compile(r"\b(\d{2,7}-\d{2}-\d)\b")

def deep_scan_word_s3():
    files = sorted([f for f in DST_DIR.glob("*.docx") if not f.name.startswith("~$")])
    print(f"=== 开始对 {len(files)} 份 Word 文档执行底层穿透式检索 ===")

    all_raw_cas_found = defaultdict(lambda: {"models": set(), "contexts": []})
    all_raw_components = defaultdict(lambda: {"models": set(), "cas_list": set(), "contexts": []})
    text_outside_tables = []
    anomalous_s3_files = []

    for p in files:
        m_name = p.stem.split(" ")[0]
        try:
            doc = docx.Document(str(p))
        except Exception as e:
            print(f"Error opening {p.name}: {e}")
            continue

        in_s3 = False
        s3_tables = []
        s3_paras = []

        # 1. 扫描段落
        for p_idx, para in enumerate(doc.paragraphs):
            t = para.text.strip()
            if not t:
                continue
            if any(k in t for k in ["3. 成分/组成资料", "3.成分/组成资料", "3. 成分/组成信息", "3.成分/组成信息", "3. 组成/成分资料", "3.组成/成分资料"]):
                in_s3 = True
                continue
            if in_s3 and any(f"{n}." in t or f"{n}、" in t for n in [4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16]):
                in_s3 = False
                break
            if in_s3:
                s3_paras.append(t)
                # 检查段落中的 CAS
                for cas in CAS_REGEX.findall(t):
                    all_raw_cas_found[cas]["models"].add(m_name)
                    all_raw_cas_found[cas]["contexts"].append(f"段落: {t}")

        # 2. 扫描表格
        in_s3_tbl = False
        for t_idx, table in enumerate(doc.tables):
            # 检查表头或首行
            tbl_text = " ".join(c.text.strip() for c in table.rows[0].cells if c.text.strip())
            if any(k in tbl_text for k in ["3. 成分/组成资料", "3.成分/组成资料", "3. 成分/组成信息", "3.成分/组成信息", "3. 组成/成分资料", "3.组成/成分资料"]):
                in_s3_tbl = True
                s3_tables.append((t_idx, table))
                continue
            if in_s3_tbl and any(f"{n}." in tbl_text or f"{n}、" in tbl_text for n in [4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16]):
                in_s3_tbl = False
                continue
            if in_s3_tbl:
                s3_tables.append((t_idx, table))

        # 分析提取 S3 表格中的每一行
        for t_idx, table in s3_tables:
            for r_idx, row in enumerate(table.rows):
                cell_texts = [c.text.strip() for c in row.cells]
                row_str = " | ".join(cell_texts)
                # 跳过大标题行
                if "成分/组成资料" in row_str or "产品类型" in row_str and len(cell_texts) == 1:
                    continue
                # 提取行内 CAS
                cas_list = CAS_REGEX.findall(row_str)
                for cas in cas_list:
                    all_raw_cas_found[cas]["models"].add(m_name)
                    all_raw_cas_found[cas]["contexts"].append(f"表格行: {row_str}")

                # 提取成分名 (通常在第1列或第2列)
                if len(cell_texts) >= 2:
                    # 识别是否是成分数据行 (非表头行如 "成分 | CAS | 含量")
                    if any(h in cell_texts[0] for h in ["成分", "化学名", "物质名称", "标识", "Component", "序号", "No"]):
                        continue
                    comp_cand = cell_texts[0]
                    # 如果第1列是 3.2 或序号，则取第2列
                    if re.match(r"^(?:3\.\d+|\d+)$", comp_cand) and len(cell_texts) >= 3:
                        comp_cand = cell_texts[1]
                    
                    comp_cand_clean = comp_cand.strip().replace("\n", " ")
                    if comp_cand_clean and comp_cand_clean not in ["混合物", "纯物质", "3.2成分", "3.2 成分", "3.1产品类型", "3.1 产品类型"]:
                        all_raw_components[comp_cand_clean]["models"].add(m_name)
                        for cas in cas_list:
                            all_raw_components[comp_cand_clean]["cas_list"].add(cas)
                        all_raw_components[comp_cand_clean]["contexts"].append(row_str)

    print(f"\n[穿透扫描结果 1] 底层 Word 全文共发现独立 CAS 号: {len(all_raw_cas_found)} 个")
    print(f"[穿透扫描结果 2] 底层 Word 全文共发现独立成分候选表述: {len(all_raw_components)} 个")

    return all_raw_cas_found, all_raw_components

if __name__ == "__main__":
    raw_cas_dict, raw_comp_dict = deep_scan_word_s3()

    conn = open_db(DB_PATH)
    cur = conn.cursor()
    cur.execute("SELECT f.value FROM msds_field f WHERE f.section = 3 AND f.kind = 'component'")
    db_cas_set = set()
    for (val,) in cur.fetchall():
        cas_m = re.search(r"CAS[:：\s]*([^\s|]+)", val)
        if cas_m:
            c = re.sub(r"[^\d\-]", "", cas_m.group(1))
            if re.match(r"^\d{2,7}-\d{2}-\d$", c):
                db_cas_set.add(c)

    raw_cas_set = set(raw_cas_dict.keys())
    diff_cas = raw_cas_set - db_cas_set
    print(f"\n=== Word 原文存在但 DB 未录入的 CAS 号 ({len(diff_cas)} 个) ===")
    for cas in diff_cas:
        print(f"CAS: {cas}")
        print(f"  涉及型号: {raw_cas_dict[cas]['models']}")
        print(f"  上下文: {raw_cas_dict[cas]['contexts'][:3]}")

    print(f"\n=== 全量 51 个 CAS 号完整列表 ===")
    for idx, (cas, info) in enumerate(sorted(raw_cas_dict.items(), key=lambda x: -len(x[1]['models'])), 1):
        in_db = "已在DB" if cas in db_cas_set else "新增/遗漏"
        print(f"{idx:<2}. CAS: {cas:<15} [{in_db}] 覆盖型号数: {len(info['models']):<3}")
