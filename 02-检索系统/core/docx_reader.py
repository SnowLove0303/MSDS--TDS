# -*- coding: utf-8 -*-
"""MSDS docx 核心读取器.

以 MSDS_CN 国彩 模板为基准 (16 表格 = 16 节):
  - 每个表格 R0 为节标题行 (整行合并, 加粗)
  - 字段行为 标签|值 两列
  - S3 为三列成分表 (化学品名称|CAS编号|含量%(w/w))
  - S11/S12 有通栏说明行 (单列)
  - S14/S15 是单列表格

读取器只读不改, 输出 ParseResult (含异常报告).
"""
from __future__ import annotations

import hashlib
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from .detectors import (
    dedupe_merged,
    hint_section,
    is_component_header_row,
    is_field_row,
    is_section_title,
    is_sub_heading,
    is_value_shape,
    looks_like_field_label,
    normalize_label,
    split_classification_line,
)
from .structure import (
    Anomaly,
    ComponentData,
    FieldData,
    ParseResult,
    SectionData,
    SubTable,
    normalize_component_cas,
    normalize_component_conc,
    normalize_component_name,
    parse_vert_component_table,
    split_component_cells,
    split_flat_component_text,
)

# S3 三列表头 (成分表)
S3_HEADERS = ("化学品名称", "CAS编号", "含量%（w/w）")

# 需要整块保留为 lines 的节 (单列表格 / 法规清单)
LINES_ONLY_SECTIONS = {14, 15}


class _NumberingResolver:
    """Word 自动编号 (<w:numPr>) 解析器.

    大量 MSDS 的字段序号 (如 "1.1 产品名称" / "9.1 外观") 不是文本, 而是
    Word 自动编号列表生成 — python-docx 的 paragraph.text 不含这些数字,
    导致序号丢失 (全库 644 文件 30021 段受影响). 本解析器从 numbering part
    读取编号定义 (numId → abstractNum → lvlText/start), 按文档顺序计算
    每段的实际编号并拼接到文本前.

    编号规则 (与 Word 一致):
      - 每个 numId 独立计数; lvl0 起始值 = 其 start (常为节号或 1)
      - lvl0 段落: 编号 = lvlText 模板 (如 "9.%1" → "9.1"), 本级递增
      - 更深级段落 (lvl1+): 本级递增, 父级引用 lvl0 当前值
        (numId=1 lvl1 "%1.%2" → "1.1", "1.2"...)
      - 更深级计数器在升回浅级时重置
    """

    def __init__(self, doc):
        self._numdefs: dict[str, dict[int, tuple[int, str]]] = {}
        np = None
        for rel in doc.part.rels.values():
            if "numbering" in rel.reltype:
                np = rel.target_part
                break
        if np is not None:
            root = np.element
            abs_map = {}
            for an in root.findall(qn("w:abstractNum")):
                abs_map[an.get(qn("w:abstractNumId"))] = an
            num_map: dict[str, str] = {}
            for num in root.findall(qn("w:num")):
                numId = num.get(qn("w:numId"))
                a = num.find(qn("w:abstractNumId"))
                if a is not None:
                    num_map[numId] = a.get(qn("w:val"))
            for numId, absId in num_map.items():
                an = abs_map.get(absId)
                if an is None:
                    continue
                levels: dict[int, tuple[int, str]] = {}
                for lvl in an.findall(qn("w:lvl")):
                    ilvl = int(lvl.get(qn("w:ilvl")))
                    lvlText_el = lvl.find(qn("w:lvlText"))
                    start_el = lvl.find(qn("w:start"))
                    lvlText = (lvlText_el.get(qn("w:val"))
                               if lvlText_el is not None else f"%{ilvl + 1}.")
                    start = (int(start_el.get(qn("w:val")))
                             if start_el is not None else 1)
                    levels[ilvl] = (start, lvlText)
                if levels:
                    self._numdefs[numId] = levels
        self._counters: dict[str, dict[int, int]] = {}

    def resolve(self, numId: str, ilvl: int) -> str:
        """返回该段自动编号 (如 "1.1" / "9.16"), 并更新计数器; 无法解析返回 ''."""
        levels = self._numdefs.get(numId)
        if not levels or ilvl not in levels:
            return ""
        c = self._counters.setdefault(numId, {})
        if 0 not in c:
            c[0] = levels[0][0]            # lvl0 起始 = start (节号或 1)
        start, lvl_text = levels[ilvl]
        # 本级当前值: ilvl0 取 lvl0 计数器, 更深级从 start 起 (若未初始化)
        if ilvl == 0:
            val = c[0]
        else:
            val = c.get(ilvl, start)
        # 用递增前的计数器快照解析 lvlText 模板 (%n 指各浅级当前值),
        # 整体替换保留字面量前缀/后缀: "9.%1" → "9.1", "%1.%2" → "1.2",
        # "%2)" → "2)", "6.%1." → "6.1."
        # 注意: 必须在递增之前取值, 否则 "9.%1" 会把 9.1 算成 2 (先 +1 后取)
        result = re.sub(r"%(\d+)",
                        lambda m: str(c.get(int(m.group(1)) - 1, 1)),
                        lvl_text)
        if not result or "%" in result:
            result = str(val)
        # 递增本级计数器 (更深级计数器在升回浅级时重置)
        c[ilvl] = val + 1
        for d in list(c):
            if d > ilvl:
                del c[d]
        return result


def _cell_lines(cell, resolver: _NumberingResolver | None = None) -> list[tuple[str, bool]]:
    """单元格内每段 (文本, 是否加粗), 自动编号已拼接到文本前.

    - 自动编号: 段落含 <w:numPr> → 计算编号前缀 (如 "1.1 产品名称：")
    - 加粗: 段落任一 run 加粗 → bold=True (用于标签列识别)
    """
    out: list[tuple[str, bool]] = []
    for para in cell.paragraphs:
        t = para.text or ""
        # GHS 象形图: Word 以 <w:drawing> 图片嵌入段落, para.text 为空
        # (reader 之前完全丢失). 检测 drawing 数 → 输出占位标记, 供
        # S2 象形图 / S14 运输标签等图片信息保留.
        n_img = len(para._p.findall('.//' + qn('w:drawing')))
        prefix = ""
        pPr = para._p.pPr
        if resolver is not None and pPr is not None and pPr.numPr is not None:
            # python-docx 的 CT_DecimalNumber.val 是 int, 而 numbering part 的
            # numId 键是 str → 必须显式转 str, 否则 _numdefs.get(int) 查不到
            numId = (str(pPr.numPr.numId.val)
                     if pPr.numPr.numId is not None else None)
            ilvl = pPr.numPr.ilvl.val if pPr.numPr.ilvl is not None else 0
            if numId is not None:
                num = resolver.resolve(numId, ilvl)
                if num:
                    prefix = num + " "
        bold = any(run.bold and run.text.strip() for run in para.runs)
        # 图片段: 文本为空 → 占位; 有文本 → 文本 + 图片标记.
        # 多图合并计数 '[象形图×2]', 供 GUI/规范化识别"此处应放 N 张象形图".
        # 源文件图片段可能残留纯符号文本 (OS-8802 英文版 '\\') → 丢弃, 只留标记.
        if n_img and t and re.fullmatch(r"[\\/|·•\-_ ]+", t):
            t = ""
        img_mark = ("[象形图×%d]" % n_img) if n_img > 1 else "[象形图]" if n_img else ""
        raw_t = (prefix + t).strip()
        # S2 危险分类行 "皮肤腐蚀/刺激健康危害  1B类" 依赖多空格做分隔符
        # (split_text_block 据此拆 分类名=类别 字段) → 保留原始空格.
        # 其余段落折叠连续空格 (显示更干净).
        if split_classification_line(raw_t) is not None:
            t = raw_t
        else:
            t = re.sub(r" {2,}", " ", raw_t)
        if img_mark:
            t = (t + (" " if t else "") + img_mark).strip()
        if t:
            out.append((t, bold))
    return out


def _cell_text(cell, resolver: _NumberingResolver | None = None) -> str:
    """单元格文本: 保留换行(多段落)与制表符(跨列分隔), 仅压缩连续空格.

    python-docx 的 paragraph.text 已把 <w:br/> 转为 \\n、<w:tab/> 转为 \\t,
    此处不再折叠制表符, 否则会破坏 S8 等"标签格内含跨列内容"的行结构.
    """
    return "\n".join(t for t, _ in _cell_lines(cell, resolver))


def _cell_images(cell, doc) -> list[ImageData]:
    """提取一个单元格内所有嵌入图片 (drawing → blip → image part).

    GHS 象形图/运输标签以 <w:drawing> 嵌入段落, 文本读取 (para.text) 为空.
    返回 ImageData 列表 (blob 原始字节 + 扩展名 + 绘制尺寸 EMU), 供 GUI 显示
    原图 / 导出文件 / 存库. 找不到对应 image part 时跳过.
    """
    from core.structure import ImageData
    out: list[ImageData] = []
    for d in cell._tc.findall('.//' + qn('w:drawing')):
        # 绘制尺寸 (EMU)
        w = h = 0
        ext = d.findall('.//' + qn('wp:extent'))
        if ext:
            try:
                w = int(ext[0].get("cx") or 0)
                h = int(ext[0].get("cy") or 0)
            except (TypeError, ValueError):
                w = h = 0
        for b in d.findall('.//' + qn('a:blip')):
            rid = b.get(qn('r:embed'))
            if not rid:
                continue
            try:
                part = doc.part.related_parts[rid]
            except (KeyError, ValueError):
                continue
            if not getattr(part, "blob", None):
                continue
            out.append(ImageData(blob=part.blob,
                                 ext=part.partname.ext.lower(),
                                 width=w, height=h))
    return out


def _dedupe_row(row, resolver: _NumberingResolver | None = None
                ) -> tuple[list[str], dict[int, set[int]]]:
    """行内单元格去重 (识别 Word 表格合并结构: gridSpan 横向跨列 + 同名合并).

    识别 Word 表格结构:
      - 显式 <w:gridSpan val=N> 跨列单元格 → 只保留一次 (合并 N 个逻辑列)
      - python-docx 对合并单元格在同一行返回重复 _tc → 按 id(_tc) 去重
      - 兜底: 相邻文本相同 (早期模板用相同文本伪合并) → 去重

    返回 (cells, bold_cells): cells[i] 为合并后文本 (含 \\n 多段),
    bold_cells[i] 为该 cell 内加粗段的索引集合 (供标签列归类).
    """
    seen_tc: list[int] = []
    items: list[tuple[str, set[int], int]] = []   # (text, bset, grid_span)
    for c in row.cells:
        lines = _cell_lines(c, resolver)
        text = "\n".join(t for t, _ in lines)
        bset = {i for i, (_, b) in enumerate(lines) if b}
        # 识别显式横向跨列 gridSpan
        gspan = 1
        tcPr = c._tc.find(qn('w:tcPr'))
        gs = tcPr.find(qn('w:gridSpan')) if tcPr is not None else None
        if gs is not None:
            try:
                gspan = int(gs.get(qn('w:val')) or 1)
            except (TypeError, ValueError):
                gspan = 1
        tc_id = id(c._tc)
        # 合并单元格 (同 _tc) → 只保留一次; 显式 gridSpan>1 也合并
        if tc_id in seen_tc:
            continue
        seen_tc.append(tc_id)
        if gspan > 1 and items and items[-1][0] == text:
            # 跨列单元格文本与前一逻辑列相同 → 属同一合并块的延续, 合并
            continue
        items.append((text, bset, gspan))
    cells: list[str] = []
    bold_cells: dict[int, set[int]] = {}
    for i, (text, bset, _g) in enumerate(items):
        text = (text or "").strip()
        if not cells or text != cells[-1]:
            cells.append(text)
            bold_cells[len(cells) - 1] = bset
    return cells, bold_cells


def is_bio_limit_header(cells: list[str]) -> bool:
    """判断是否 S8.2 生物限值表头行.

    生物限值表 (GB 8.2) 列: 组分名称 | 标准来源 | 生物监测指标 | 生物限值 | 采样时间.
    PEA-4139 模板为 6 列 (标准来源 列重复合并), 去重后 5 列有效. 识别条件:
      - 首格含 "组分名称" (或 组分/名称), 且
      - 全行含 ≥4 个关键列名 (标准来源/生物监测指标/生物限值/采样时间/限值)
    """
    if not cells or not cells[0].strip():
        return False
    if any("\n" in c or len(c) > 40 for c in cells):
        return False
    first = cells[0].strip()
    if not re.search(r"组分|成分|物质名称", first):
        return False
    joined = " ".join(cells)
    kws = ["标准来源", "生物监测", "生物限值", "限值", "采样时间", "时间"]
    hits = sum(1 for k in kws if k in joined)
    return hits >= 4


def _collect_story(container) -> str:
    """收集页眉/页脚完整内容: 段落 + 表格单元格 (模板页眉含产品名表格)."""
    parts = []
    for p in container.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    for tb in container.tables:
        for row in tb.rows:
            seen = []
            for cell in row.cells:
                if id(cell._tc) in seen:
                    continue
                seen.append(id(cell._tc))
                t = cell.text.strip()
                if t:
                    parts.append(t)
    return "\n".join(parts)


# ---- 页眉页脚 → Section 0 字段化 (父子级) ----

_PAGE_FIXED_KEYWORDS = ("版本", "Version", "修订", "页码", "公司", "电话", "传真",
                        "地址", "物料", "安全技术说明书", "MSDS")
# 页码: '5 / 5' / '3 / 5' → 页码字段
_PAGE_NUM_RE = re.compile(r"^\s*\d+\s*/\s*\d+\s*$")
# 页脚复合内容结尾锚: '-MSDS' / '_MSDS' / '-SDS' (忽略大小写)
_PAGE_COMPOUND_TAIL_RE = re.compile(r"[-_](?:MSDS|SDS)$", re.IGNORECASE)


def _page_compound(text: str) -> tuple[str, str] | None:
    """拆分页脚复合内容 → (公司名称, 产品型号); 非复合返回 None.

    '广州冠志化工有限公司 2-苯氧基乙醇-MSDS'       → (公司, '2-苯氧基乙醇')
    'GUANGZHOU GUANZHI NEW MSTAR TECHNOLOGY CO.,LTD. BEK-500L-MSDS'
        → (英文公司, 'BEK-500L')
    以 '-MSDS/-SDS' 结尾为锚 (MSDS 文档页脚特征), 型号可数字/中文/英文开头;
    公司名含空格也支持 (从右往左在最后一个空格处切分).
    """
    text = (text or "").strip()
    if not _PAGE_COMPOUND_TAIL_RE.search(text):
        return None
    head = text[: _PAGE_COMPOUND_TAIL_RE.search(text).start()].rstrip()
    sp = head.rfind(" ")
    if sp <= 0:
        return None
    cn, en = head[:sp].strip(), head[sp + 1:].strip()
    if not cn or not en:
        return None
    # 公司名必须非空且型号不含空格 (型号段为最后一个空格后的连续串)
    if " " in en:
        return None
    return cn, en


def _page_fields(text: str) -> list[FieldData]:
    """把一行页眉/页脚文本拆成 0~N 个字段 (父子级下字段名不含 页眉/页脚 前缀).

    - 页码 '3 / 5'        → (页码, '3 / 5')
    - 冒号 '修订日期：...'  → (修订日期, ...)   (剥离 P/F/D 标记前缀)
    - 复合 '公司 型号-MSDS' → (公司名称, ...) + (产品型号, ...)
    - 其他 '物料安全数据表'  → (原文, '') 标题行
    """
    text = (text or "").strip()
    if not text:
        return []
    if _PAGE_NUM_RE.match(text):
        return [FieldData(label="页码", value=text, editable=False)]
    # 冒号拆分 (不含 '-', 避免误拆 PEA-4139 等产品编号)
    label = value = ""
    for sep in ("：", ":"):
        if sep in text:
            a, _, b = text.partition(sep)
            if a.strip() and b.strip():
                label, value = a.strip(), b.strip()
                break
    if not label:
        # 无冒号 → 页脚复合 (公司名 + 型号) 或 页眉标题/页脚独立字段
        comp = _page_compound(text)
        if comp:
            return [FieldData(label="公司名称", value=comp[0], editable=False),
                    FieldData(label="产品型号", value=comp[1], editable=False)]
        label = text
    # 去掉字段标记前缀 (如 "P修订日期" → "修订日期");
    # 仅当下一字符为中文才剥离, 避免误伤 "PEA-4139" 这类英文产品编号
    if (label[:1] in ("P", "F", "D") and len(label) > 1
            and label[1].isalpha() and not label[1].isascii()):
        label = label[1:]
    editable = _page_editable(label, value)
    return [FieldData(label=label, value=value, editable=editable)]


def _page_editable(label: str, value: str) -> bool:
    """页眉/页脚字段默认编辑状态: 明确固定项不可编辑, 其余默认可编辑."""
    if _PAGE_NUM_RE.match(label or ""):
        return False   # 页码
    if any(k in label for k in _PAGE_FIXED_KEYWORDS):
        return False   # 版本/公司/电话/修订日期/物料标题等固定信息
    return True


def _collect_page_fields(story, prefix: str, out: list[FieldData]) -> None:
    """收集页眉或页脚的全部字段 (段落 + 表格单元格)."""
    for p in story.paragraphs:
        t = p.text.strip()
        if t:
            out.extend(_page_fields(t))
    for tb in story.tables:
        for row in tb.rows:
            seen = []
            for cell in row.cells:
                if id(cell._tc) in seen:
                    continue
                seen.append(id(cell._tc))
                t = cell.text.strip()
                if not t:
                    continue
                # 页眉表格 = 产品名称/型号表 (模板中存 "PEA-4139")
                if prefix == "页眉":
                    out.append(FieldData(label="产品名称", value=t, editable=True))
                else:
                    out.extend(_page_fields(t))


def _parse_page_section(hdr_story, ftr_story) -> SectionData | None:
    """解析页眉/页脚为 section 0, 父子级: 0.页眉页脚 → 页眉/页脚 → 字段.

    '页眉'/'页脚' 作为 sub 子标题 (order 交替 + lines 触发 split_text_block),
    其下挂各自字段 (产品名称/版本/公司名称/修订日期/页码 等). 字段名不含
    '页眉-'/'页脚-' 前缀 (父子级已表达归属), 供 GUI/总库/Excel 层级呈现.
    """
    sec = SectionData(number=0, title="页眉页脚", full_title="0.页眉页脚")
    has_any = False
    for sub_label, story, prefix in (("0.1 页眉", hdr_story, "页眉"),
                                     ("0.2 页脚", ftr_story, "页脚")):
        fields: list[FieldData] = []
        _collect_page_fields(story, prefix, fields)
        if not fields:
            continue
        sec.lines.append(sub_label)      # → split_text_block → sub 行 (页眉/页脚)
        sec.order.append("line")
        for f in fields:
            sec.fields.append(f)
            sec.order.append("field")
        has_any = True
    if not has_any:
        return None
    return sec


def _subsection_section(text: str) -> int | None:
    """从 '4.1 Description of first aid measures' 提取归属节号 4.

    文本框内容块首段带小节号 (N.N) → 整块归属节 N (XML 流顺序中
    内容块常出现在本节标题之前, 逐段流式会错位归入上一节).
    """
    m = re.match(r"^\s*(\d{1,2})\.\d+\s", text or "")
    return int(m.group(1)) if m else None


# 文本框归位创建目标节时的标准节标题兜底 (英文模板; 中文节标题块随后
# 到达时 _ensure_section 已存在则沿用, title 保持一致)
_DEFAULT_SECTION_TITLES = {
    1: "Identification of the substance/mixture",
    2: "Hazards identification",
    3: "Composition/information on ingredients",
    4: "First aid measures",
    5: "Firefighting measures",
    6: "Accidental release measures",
    7: "Handling and storage",
    8: "Exposure controls/personal protection",
    9: "Physical and chemical properties",
    10: "Stability and reactivity",
    11: "Toxicological information",
    12: "Ecological information",
    13: "Disposal considerations",
    14: "Transport information",
    15: "Regulatory information",
    16: "Other information",
}


# 无小节号文本框块的首行特征归属 (国彩模板: 整节内容块先于节标题块到达,
# 首段无 'N.N' 前缀 → 用特征词定位目标节). 与行级 SECTION_HINT 隔离,
# 仅在文本框块级触发, 避免污染表格行兜底.
_TEXTBOX_HINT = {
    "reactivity": 10,
    "chemical stability": 10,
    "thermal decomposition": 10,
    "hazardous decomposition": 10,
    "additional information about design": 8,
    "relevant safety": 15,
    "this information is based": 16,
}


def _textbox_hint(text: str) -> int | None:
    t = (text or "").strip().lower()
    for key, num in _TEXTBOX_HINT.items():
        if key in t:
            return num
    return None


def _iter_doc_content(doc):
    """按 body XML 顺序产出内容事件: 表格行 + VML 文本框块.

    修复: VML 文本框 (w:pict > v:textbox > w:txbxContent) 承载的正文
    python-docx 顶层 API (doc.tables/doc.paragraphs) 完全不可见, 主循环
    原先整块丢失 (OS-12020 EN Guocai 12 节全在文本框, 只解析出表格的
    3 节). 每个文本框 = 一节内容块 (首段常为 'N.N 标题'), 整组 yield
    供归位处理, 与表格行按 body 顺序交错.

    事件: ("table", ti, None) 对应 doc.tables[ti] 的后续行;
          ("textbox", None, [lines]) 文本框内全部非空段落.
    """
    body = doc.element.body
    ti = 0
    for child in body.iterchildren():
        if child.tag == qn("w:tbl"):
            yield ("table", ti, None)
            ti += 1
        elif child.tag in (qn("w:pict"), qn("w:p")):
            # 每个文本框独立成块 (同一 w:p 下可多个 w:pict 浮动框,
            # 若共享 lines 会吞并后一个块 → 节标题丢失)
            for tbx in child.findall(".//" + qn("w:txbxContent")):
                lines: list[str] = []
                for p in tbx.findall(qn("w:p")):
                    text = "".join(t.text or "" for t in
                                   p.findall(".//" + qn("w:t")))
                    if text.strip():
                        lines.append(text.strip())
                if lines:
                    yield ("textbox", None, lines)


def read_msds(path: str | Path) -> ParseResult:
    """读取一个 MSDS docx 为 ParseResult.

    流式行解析: 遍历所有表格的所有行, 维护"当前节"指针,
    遇到节标题行 (含内嵌节标题) 即切换当前节.
    兼容三种结构:
      - 一表一节 (标准模板)
      - 一表多节 (S3 表内嵌 S4, 如 PU-1034)
      - 少表多节 (S4 无独立表格)
    """
    path = Path(path)
    result = ParseResult(file_path=str(path), file_name=path.name)
    try:
        result.sha256 = hashlib.sha256(path.read_bytes()).hexdigest()
    except OSError:
        pass

    doc = Document(str(path))
    result.tables_count = len(doc.tables)
    result.paragraphs_count = len(doc.paragraphs)
    if doc.sections:
        sec = doc.sections[0]
        result.header = _collect_story(sec.header)
        result.footer = _collect_story(sec.footer)
        # 页眉/页脚字段化 → section 0 (可编辑/不可编辑)
        s0 = _parse_page_section(sec.header, sec.footer)
        if s0 is not None:
            result.sections[0] = s0

    current: SectionData | None = None
    norm_stats = {"split": 0, "fw": 0}
    _bio_active: SubTable | None = None   # 正在收集的 S8.2 生物限值子表 (None = 未在收集)
    mix_pending = False   # S3 'Mixtures' 标题行的下一行 = 主要成分 (国彩版跨行)
    mix_stash: dict[int, list[tuple[str, str]]] = {}  # 节号 → Mixtures note 成分 (延迟提升)
    resolver = _NumberingResolver(doc)   # 自动编号 → 文本序号 (修复普遍序号丢失)

    # 事件流: 表格行 + VML 文本框段落 按 body 顺序交错 (文本框修复,
    # 见 _iter_doc_content). 表格行走原逻辑; 文本框行作为单列段落进入
    # 节切换 + 说明行 (信息保留, 结构以 lines 呈现).
    for _kind, _ti, _text in _iter_doc_content(doc):
        if _kind == "table":
            tb = doc.tables[_ti]
            for _ri, row in enumerate(tb.rows):
                # 收集本行嵌入图片 (GHS 象形图等): 记录所在节/表/行, 合并单元格去重
                _seen_tc: set[int] = set()
                for _ci, _cell in enumerate(row.cells):
                    if id(_cell._tc) in _seen_tc:
                        continue
                    _seen_tc.add(id(_cell._tc))
                    for _img in _cell_images(_cell, doc):
                        _img.section = current.number if current is not None else 0
                        _img.table, _img.row, _img.cell = _ti, _ri, _ci
                        result.images.append(_img)
                cells, bold_cells = _dedupe_row(row, resolver)
                if not cells or not cells[0]:
                    # 首格空: 仅当普通节中后续列有内容 → 无标题行 (内容放文本列,
                    # BL-8085 S5 特殊危害后的补充行); 全空行 / 成分表行跳过.
                    if not (len(cells) > 1 and any(c.strip() for c in cells[1:])
                            and current is not None and not current.is_component_table):
                        continue
                first = cells[0]

                # 节标题 (含内嵌): 切换当前节
                num, title = is_section_title(first)
                if num is not None:
                    current = _ensure_section(result, num, title, first)
                    continue

                # 未进入任何节: 尝试用特征词兜底
                if current is None:
                    hint = hint_section(first)
                    if hint is not None:
                        current = _ensure_section(result, hint, first, first)
                    else:
                        continue

                # 成分表列标题行 (如 ['成分',''] / ['成分 /','']) 先于表头出现
                # → 预标记成分表 (成分表头紧随其后, 否则该行会误走字段逻辑)
                if not current.is_component_table and len(cells) <= 2 and (
                        cells[0].rstrip(" /／ ") in
                        ("成分", "组分", "成分/组成", "危险成分", "名称", "化学品名称")):
                    current.is_component_table = True
                    continue

                # 单列表格内嵌成分表: 整格文本用制表符/空格分隔列、换行分隔行
                # (PA-4757/HPU-7651/PA-4408/OS-8030 等英文+中文单列 S3)
                if len(cells) == 1 and not current.is_component_table:
                    flat = split_flat_component_text(first)
                    if flat is not None and flat["rows"]:
                        current.is_component_table = True
                        current.component_header = " | ".join(flat["header"])
                        _apply_flat_component(flat, current, norm_stats)
                        continue

                # 列内多行成分表: 每列含多段(表头+数据), 表头含 name/cas/conc 语义
                # (PU-3210/BEK-500L/PU-6088 等英文模板 3 列表格, 每格多段)
                if not current.is_component_table and len(cells) >= 2:
                    vert = parse_vert_component_table(cells)
                    if vert is not None and vert["rows"]:
                        current.is_component_table = True
                        current.component_header = " | ".join(vert["header"])
                        for nm, cs, cc in vert["rows"]:
                            comps = split_component_cells(nm, cs, cc)
                            if norm_stats is not None and len(comps) > 1:
                                norm_stats["split"] += len(comps) - 1
                            current.components.extend(comps)
                        continue

                # S8.2 生物限值子表 (GB 8.2): 表头行 + 后续数据行 → 结构化子表.
                # 不进 fields/lines (避免表头被当字段值、数据行被当 note), 保持
                # Word 表格列结构, 供 GUI 以子表形式呈现 (参照 PEA-4139 模板).
                if current is not None and current.number == 8 and not _bio_active:
                    if is_bio_limit_header(cells):
                        _bio_active = SubTable(title="生物限值", seq="8.2",
                                               header=[c.strip() for c in cells
                                                       if c.strip()])
                        current.sub_tables.append(_bio_active)
                        continue
                if _bio_active is not None:
                    # 收集数据行: 多列或单列有效值, 且首格无冒号、非编号标题、非节标题
                    if not re.search(r"：|:", cells[0]):
                        if (not is_sub_heading(cells)
                                and not is_section_title(cells[0])[0]):
                            _bio_active.rows.append(
                                [c.strip() for c in cells if c.strip()])
                            continue
                    _bio_active = None   # 遇到新标题/字段行 → 结束收集

                # 成分表头 → 当前节标记为成分表 (记录实际表头文本, 供检索/显示)
                if is_component_header_row(cells):
                    current.is_component_table = True
                    current.component_header = " | ".join(c.strip() for c in cells if c.strip())
                    continue

                # S3 单列 'Mixtures' 标题 + 主要成分 note (OS/PTF 英文模板)
                # 冠志版: 单格 '3.1 Mixtures\n· polyether urethane'
                # 国彩版: 两行 '3.1 Mixtures' + 'polyether urethane 20±1%（w/w） /'
                # 注意: 暂存不立即 append — 仅当节最终无真成分表时才提升
                # (PA-4757 'Description: Polymer' 是总类描述, 其后有 3.2 真成分表,
                #  不该作成分; OS 系列 'polyether urethane' 是真主要成分才提升)
                if current.number == 3 and not current.is_component_table:
                    if _is_mix_header(first):
                        # 保留 'Mixtures' 声明本身 (对应中文 '产品类型：混合物'),
                        # 作为 sub 标题 (iter_rows 会拆序号). 单格
                        # '3.1 Mixtures\n· polyether urethane' → 首行作标题,
                        # 其后行为成分 note (延迟提升到 components).
                        head = first.split("\n", 1)[0].strip()
                        if head and not any(f.label == head for f in current.fields):
                            current.fields.append(FieldData(label=head,
                                                            value="", editable=False))
                            current.order.append("field")
                        if "\n" in first:
                            for seg in first.split("\n")[1:]:
                                comp = _parse_mix_row(seg)
                                if comp is not None:
                                    mix_stash.setdefault(current.number, []).append(comp)
                            mix_pending = False
                            continue
                        mix_pending = True
                        continue
                    if mix_pending:
                        mix_pending = False
                        comp = _parse_mix_row(first)
                        if comp is not None:
                            mix_stash.setdefault(current.number, []).append(comp)
                        continue

                if current.is_component_table:
                    _parse_component_row(cells, current, norm_stats)
                else:
                    _parse_field_row(cells, current, bold_cells)
        else:
            # ---- VML 文本框块 (整节内容, 整组归位) ----
            lines = _text  # list[str]
            first = lines[0] if lines else ""
            # 节标题块 → 切换当前节 (节已由归位预建则更新 title/full_title)
            num, title = is_section_title(first)
            if num is not None:
                if num in result.sections:
                    result.sections[num].title = title
                    result.sections[num].full_title = first
                    current = result.sections[num]
                else:
                    current = _ensure_section(result, num, title, first)
                for ln in lines[1:]:
                    _parse_field_row([ln], current, None)
                continue
            if current is None:
                hint = hint_section(first)
                if hint is not None:
                    current = _ensure_section(result, hint, first, first)
                else:
                    # 文档开头无节标题的前置内容 → 视为 S1 前置块
                    # (信息保留优先; 后随 SECTION 1 标题块会更新 title)
                    current = _ensure_section(
                        result, 1, _DEFAULT_SECTION_TITLES[1], first)
            # 首段带小节号 (N.N) → 整块归属对应节 (目标节不存在则预建,
            # 修正"内容块先于标题块"的流顺序错位)
            tnum = _subsection_section(first)
            if tnum is not None:
                target = result.sections.get(tnum)
                if target is None:
                    target = _ensure_section(
                        result, tnum, _DEFAULT_SECTION_TITLES.get(tnum, first),
                        first)
                for ln in lines:
                    _parse_field_row([ln], target, None)
                continue
            # 无前缀文本框块: 首行特征词兜底归属 (国彩模板整节内容块无
            # 小节号且先于标题块到达 → 需直接归对应节而非 current)
            hn = _textbox_hint(first)
            if hn is not None:
                target = result.sections.get(hn)
                if target is None:
                    target = _ensure_section(
                        result, hn, _DEFAULT_SECTION_TITLES.get(hn, first),
                        first)
                for ln in lines:
                    _parse_field_row([ln], target, None)
                continue
            for ln in lines:
                _parse_field_row([ln], current, None)

    result.sections_count = len(result.sections)
    # Mixtures note 延迟提升: 仅当节无真成分表 (components 为空) 时才采用
    for _n, _stash in mix_stash.items():
        _sec = result.sections.get(_n)
        if _sec is not None and not _sec.components and _stash:
            for _nm, _cc in _stash:
                _sec.components.append(ComponentData(name=_nm, cas="", conc=_cc,
                                                     raw_name=_nm, raw_conc=_cc))
    # 兼容性归一化摘要 (改动可追溯)
    if norm_stats["split"]:
        result.anomalies.append(Anomaly(
            "warn", 3, f"S3 成分拆分 {norm_stats['split']} 行",
            "一行多成分单元格按换行符拆分 (名称/CAS/含量三列对齐)",
        ))
    if norm_stats["fw"]:
        result.anomalies.append(Anomaly(
            "warn", 3, f"S3 全角符号归一 {norm_stats['fw']} 行",
            "全角 ＞＜～％－（）等转半角",
        ))
    _post_check(result)
    return result


def _ensure_section(result: ParseResult, num: int, title: str, full_title: str) -> SectionData:
    if num not in result.sections:
        result.sections[num] = SectionData(number=num, title=title, full_title=full_title)
    return result.sections[num]


# 全角符号标记 (用于统计归一化改动; 与 structure._FW2HW 对应)
_FW_MARKS = ("＞", "＜", "～", "％", "－", "（", "）", "，", "；", "：")


def _apply_flat_component(flat: dict, sec: SectionData,
                          stats: dict | None = None) -> None:
    """把单列表格成分表解析结果 (split_flat_component_text) 写入节.

    - pre_lines:  表头前说明 → tab 两列 (如 '产品类型：\t混合物') 作字段,
                  其余作通栏说明行
    - rows:       成分行 → 经 split_component_cells 归一化后入 components
    - post_lines: 表头后非成分说明 (尾注 / S4 溢出通栏) → 通栏说明行
    """
    for ln in flat["pre_lines"]:
        if "\t" in ln:
            a, _, b = ln.partition("\t")
            if a.strip() and b.strip():
                if re.fullmatch(r"\d+(?:\.\d+)*\.?", a.strip()):
                    # 纯序号 + tab + 标题 ('3.2\tComposition', PA-4851) →
                    # 序号+标题 子标题 (无值, tab 后是标题文字而非内容)
                    sec.fields.append(FieldData(
                        label=normalize_label(a.strip() + " " + b.strip()),
                        value="", editable=False))
                else:
                    sec.fields.append(FieldData(label=normalize_label(a.strip()),
                                                value=b.strip()))
                sec.order.append("field")
                continue
        sec.lines.append(ln)
        sec.order.append("line")
    for nm, cs, cc in flat["rows"]:
        comps = split_component_cells(nm, cs, cc)
        if stats is not None and len(comps) > 1:
            stats["split"] += len(comps) - 1
        sec.components.extend(comps)
    for ln in flat["post_lines"]:
        sec.lines.append(ln)
        sec.order.append("line")


# S3 单列 'Mixtures' 标题正则 (自动编号恢复后形如 '3.1 Mixtures')
_MIX_HEAD_RE = re.compile(r"^(?:\d+(?:\.\d+)?\s*)?mixtures\b", re.I)


def _is_mix_header(text: str) -> bool:
    """是否 S3 'Mixtures' 标题行 (可带 '3.1 ' 前缀, 后接成分 note)."""
    first = (text or "").split("\n", 1)[0].strip()
    return bool(_MIX_HEAD_RE.match(first))


def _parse_mix_row(text: str) -> tuple[str, str] | None:
    """从 S3 'Mixtures' 成分行解析 (name, conc). None = 非成分行.

    OS/PTF 英文模板: 单格 'Mixtures\\n· polyether urethane' 或
    跨行 'Mixtures' 标题 + 成分行 'polyether urethane 20±1%（w/w） /'.
    """
    v = (text or "").replace("\xa0", " ")
    v = re.sub(r"^[\s·•\-]+", "", v)          # 去项目符号/缩进
    v = re.sub(r"^Description\s*:\s*", "", v, flags=re.I)
    v = re.sub(r"\s*/\s*$", "", v).strip()
    if not v:
        return None
    low = v.lower()
    if (low in ("void",) or low.startswith("dangerous components")
            or "no hazardous components" in low or low.startswith("composition")):
        return None
    # 尾部含量 (如 'polyether urethane 20±1%（w/w）')
    m = re.search(
        r"\s+(\d+(?:[±＋\-]\d+)?\s*(?:%\s*(?:（w/w）)?|（w/w）|w/w|wt%|W/W))\s*$", v)
    conc = ""
    if m:
        conc = m.group(1)
        v = v[:m.start()].rstrip()
    return v, conc


def _parse_component_row(cells: list[str], sec: SectionData,
                         stats: dict | None = None) -> None:
    """解析成分表的一行 (成分表模式).

    兼容性原则: 数据行按 \\n 拆分成多个成分 (一行多成分), 并对
    名称/CAS/含量做全角→半角归一化 (详见 structure.split_component_cells).
    stats 累计 'split' (拆分出的额外成分行数) 与 'fw' (全角归一化行数),
    由 read_msds 汇总为 Anomaly, 保证改动可追溯.
    """
    first = cells[0]
    # 产品类型
    if first.startswith("产品类型"):
        sec.fields.append(FieldData(label="产品类型", value=" / ".join(cells[1:]) or ""))
        sec.order.append("field")
        return
    # "成分" 列标题行 (合并行, 如 ['成分',''] / ['成分 /',''] ) → 跳过
    if len(cells) <= 2 and first.rstrip(" /／ ") in (
            "成分", "组分", "成分/组成", "危险成分", "名称", "化学品名称"):
        return
    # 单列: 说明行 / 表头残留
    if len(cells) == 1:
        sec.lines.append(first)
        sec.order.append("line")
        return
    # 数据行: 名称 | CAS | 含量 (兼容一行多成分 + 全角符号)
    if stats is not None:
        for c in cells[:3]:
            if any(m in c for m in _FW_MARKS):
                stats["fw"] += 1
                break
    comps = split_component_cells(cells[0], cells[1],
                                  cells[2] if len(cells) >= 3 else "")
    if stats is not None and len(comps) > 1:
        stats["split"] += len(comps) - 1
    sec.components.extend(comps)


_BULLET_RE = re.compile(r"^[·•\-]")
_BULLET_NUM_RE = re.compile(r"^[·•\-]\s*\d+(?:\.\d+)?\s*\S")


def _split_stacked_fields(label: str, value: str) -> list[tuple[str, str]] | None:
    """一行内多字段堆叠 (左格多标签 + 右格多值) → 按行对齐拆分.

    源文件把多个 14.x 子项挤在同一行表格 (换行分隔) 是格式规范问题,
    reader 做兼容, 把堆叠字段拆成多个可编辑字段, 供 GUI/覆写规范化使用.
    兼容 S14 三类变体:
      - OS-1310 中文: '运输危险级别：\\n包装类型：\\n环境危险：\\n特殊防范措施\\n附加信息：\\n\\n按《MARPOL73/78》...'
        ↔ '9\\nIII\\n否\\n参见6-8节\\n处理前...\\n\\n不适用'
      - HPU-7651 中文: '海上运输：\\n空运\\n用户特殊注意事项：' ↔ 3 值 (无编号)
      - BEK-500L 英文: '14.1 UN-Number\\n· ADR...\\n· 14.2 ...' ↔ 'Void\\nVoid\\n...'
    判定 (避免误拆 S9 复合标签 'Combustion value:\\nSaturated vapor pressure:'):
      - 独立标签 ≥3 行 (2 行复合标签不触发)
      - 值行 ≥2
      - 标签数与值行数基本对齐 (±1), 防止行数错位误拆
      - 项目符号子行 ('· ADR...' / '· Class' 无编号) 并入前一标签值;
        '· 14.x' 编号子行仍为标签 (去 '·' 前缀)
    返回 [(label_i, value_i)]; None = 非堆叠 (保持单字段).
    """
    raw_labels = [ln.strip() for ln in (label or "").split("\n") if ln.strip()]
    raw_values = [ln.strip() for ln in (value or "").split("\n") if ln.strip()]
    if len(raw_labels) < 3 or len(raw_values) < 2:
        return None
    tags: list[str] = []
    subs: list[list[str]] = []
    for ln in raw_labels:
        if _BULLET_RE.match(ln) and not _BULLET_NUM_RE.match(ln):
            if subs:                       # '· ADR...' → 并入前一标签
                subs[-1].append(ln)
            continue
        if len(ln.strip()) <= 1:           # 单字残行 (断行标签的尾字) 跳过
            continue
        tags.append(ln)
        subs.append([])
    if len(tags) < 3 or abs(len(tags) - len(raw_values)) > 1:
        return None
    out: list[tuple[str, str]] = []
    for i, t in enumerate(tags):
        t = re.sub(r"^[·•\-]\s*", "", t).strip()
        v = raw_values[i] if i < len(raw_values) else ""
        extra = "\n".join(subs[i]).strip()
        if extra:
            v = (v + "\n" + extra).strip()
        out.append((t, v))
    if len(raw_values) > len(tags):
        # 多余值行 (值行比标签多) 并入最后一个标签
        out[-1] = (out[-1][0],
                   (out[-1][1] + "\n" + "\n".join(raw_values[len(tags):])).strip())
    return out


def _parse_field_row(cells: list[str], sec: SectionData,
                     bold_cells: dict[int, set[int]] | None = None) -> None:
    first = cells[0]
    bold0 = set((bold_cells or {}).get(0, ()))   # cells[0] 内加粗段索引
    # 无标题行: 首格空、次列有内容 (['', '内容']) → 内容在字段列.
    # 父子级归属判别: 内容在字段列 = 前一个带标签字段的延续内容 (如
    # S5 '物质或混合物的特殊危害' 后 ['', '在着火或爆炸情况下，不要吸进烟尘。'],
    # 与原文档一致, 该句与 '燃烧时释放一氧化碳...' 同级同属 5.3),
    # 因此并入前一个字段的 value (换行分隔), 而非独立成通栏跨列行.
    # 只有内容出现在首格/单列 (['内容', '']) 才判为独立通栏总结句 (见下).
    if not first.strip():
        content = " / ".join(c.strip() for c in cells[1:] if c.strip()).strip()
        if content:
            if sec.fields:
                # 属于父子级延续: 并入最近字段 (无论其是否带标签), 保留原换行
                if sec.fields[-1].value:
                    sec.fields[-1].value += "\n"
                sec.fields[-1].value += content
            else:
                # 兜底: 节首个字段即无标题 → 独立字段 (全库未出现)
                sec.fields.append(FieldData(label="", value=content))
                sec.order.append("field")
        return
    # 单列: 通栏说明行 / 子标题
    if len(cells) == 1:
        sec.lines.append(first)
        sec.order.append("line")
        if bold0:                       # 记录加粗段 → 标签列归类放宽
            sec.line_bold[len(sec.lines) - 1] = bold0
        return
    # 子标题 (如 "8.1 暴露控制") 无值
    if is_sub_heading(cells):
        sec.lines.append(first)
        sec.order.append("line")
        return
    if is_field_row(cells) or (bold0 and len(cells) >= 2
                               and looks_like_field_label(first)):
        # 加粗补充: 无冒号英文标签 (如 "Waste Disposal Method" 左列加粗、
        # 右列有值) → 归为字段行, 标签列 (加粗) 受保护不可覆写
        # 守卫: 加粗但明显是值形态 (42±2% / 乳白色液体 / 乙醇) 的不判字段,
        # 避免值混进标签列 (looks_like_field_label 拦截值形态).
        label = first
        value = " / ".join(cells[1:]).strip()
        # 值格为空且标签无编号前缀 → 通栏引导词/说明行 (与 "该产品无可用的毒理学研究。"
        # 同级), 而非字段标签. 例: CX-470 S11 '类似产品的风险评估数据：'
        # (Word 中 cell1 为空, 与 '该产品无可用的毒理学研究。' 结构相同, 只是带冒号结尾,
        # 被 is_field_row 的冒号判定误归为字段). 有编号前缀的空值行
        # (如 '11.6 致癌性' 模板占位未填) 仍视为字段.
        if not value and not re.match(r"^\s*\d+\.\d+\s*\S", label):
            sec.lines.append(first)
            sec.order.append("line")
            if bold0:
                sec.line_bold[len(sec.lines) - 1] = bold0
            return
        # 标签格内含制表符 = 一行跨多列 (S8 R3: '手部防护：\t喷涂过程中要求有呼吸防护设备。')
        # tab 前为标签, tab 后为内容 → 并入值
        if "\t" in first:
            label_part, _, rest = first.partition("\t")
            if rest.strip():
                label = label_part
                value = (rest.strip() + " / " + value).strip(" /")
        # 一行多字段: 标签格内多行各为独立编号字段 (英文文件把
        # "9.19 Dynamic viscosity:\n9.20 Explosion characteristics\n9.21 Dust explosion level"
        # 合并进一格, 值格多行对齐) → 按行拆分, 每行一个字段.
        # 兼容行内多编号: "9.19 最低成膜温度MFFT/℃9.20 玻璃化温度Tg/℃：" 两个编号挤在一行
        # (无换行, RA-15000) → 先在后续编号前断行, 再按行对齐.
        # 条件: 标签多行且每行均以 "数字.数字" 开头 (避免误拆
        # "Combustion value:\nSaturated vapor pressure:" 这类复合标签).
        label_lines: list[str] = []
        for ln in label.split("\n"):
            # 行内多编号展开: 在非首个 "9.x " 前断行. 兼容编号后无空格变体:
            #   "9.19 最低成膜温度...9.20 玻璃化温度Tg/℃：" (空格)
            #   "4.4.接触皮肤：4.5.吸入：" (点+汉字, RA-15000 无空格)
            #   "1.2产品使用建议" (汉字紧跟编号)
            # 排除 "2.1.1 分类" 子编号 (点后是数字, 不触发; 且 multi_num
            # 的 all-check 会因 '.1 分类' 非编号开头而回退).
            label_lines.extend(p for p in re.split(
                r"(?=\d{1,2}\.\d{1,2}(?:\s+|[.、．：:][一-鿿A-Za-z]|[一-鿿]))", ln)
                if p.strip())
        multi_num = (len(label_lines) > 1
                     and all(re.match(r"^\s*\d+\.\d+\s*\S", ln) for ln in label_lines))
        if multi_num:
            value_lines = value.split("\n") if "\n" in value else [value]
            for i, ln in enumerate(label_lines):
                ln = ln.strip()
                val = value_lines[i].strip() if i < len(value_lines) else ""
                sec.fields.append(FieldData(label=normalize_label(ln), value=val))
                sec.order.append("field")
            return
        # 非编号多标签堆叠 (S14 运输信息: 源文件把 14.3-14.7 挤在同一行表格,
        # 左格 '运输危险级别：\n包装类型：...' 右格 '9\nIII\n否...' 按行对齐).
        # multi_num 只处理全编号标签 (S9), 此处兜底冒号/短标签/英文·编号混合.
        stacked = _split_stacked_fields(label, value)
        if stacked:
            for ln, val in stacked:
                sec.fields.append(FieldData(label=normalize_label(ln), value=val))
                sec.order.append("field")
            return
        label = normalize_label(label)
        # 单字残行丢弃: 源文件标签跨行断开产生的残行 (如 '制：'/'法：'/'害：'
        # 断成单独一行, normalize 后只剩单字 '制'/'法'/'害'). 单字不成字段,
        # 有值并入前一字段的 value, 无值丢弃. (RA-15000 '制', PA-4408 '法' 等)
        if len(label) <= 1 and not is_value_shape(label):
            if value and sec.fields:
                prev = sec.fields[-1]
                prev.value = (prev.value + "\n" + value).strip()
            elif value:
                sec.lines.append(value)
                sec.order.append("line")
            return
        sec.fields.append(FieldData(label=label, value=value))
        sec.order.append("field")
    else:
        # 无法归类 → 说明行 (过滤空格, 避免 '文本 / ' 尾部斜杠: 源表 2 列空值行)
        sec.lines.append(" / ".join(c for c in cells if c.strip()))
        sec.order.append("line")


def _renumber_s9(result: ParseResult) -> None:
    """S9 编号重排修复 (禁止跳号).

    检测到 S9 字段编号不连续 / 重复 / 不从 9.1 开始时, 按文档出现顺序
    重新连续编号 9.1, 9.2, ..., 9.N, 并同步更新 FieldData.label 中的编号,
    使下游 (检索/表格/数据库/覆写) 获得连续编号.

    覆盖"一行多编号"情况 (英文文件把 "9.19\\n9.20\\n9.21" 合并进一格):
    label 内所有 ^9.\\d+ 编号按出现顺序一并重排. 原始编号映射记入
    anomalies (info, 供审计追溯). 无编号字段 (如 S9 气味) 保持无编号.
    """
    sec9 = result.sections.get(9)
    if not sec9:
        return
    # 收集每个 field 的编号列表 (按文档顺序)
    field_nums: list[tuple[FieldData, list[int]]] = []
    for f in sec9.fields:
        nums = [int(m) for m in re.findall(r"(?m)^9\.(\d+)", f.label)]
        if nums:
            field_nums.append((f, nums))
    if not field_nums:
        return
    all_nums = [n for _, ns in field_nums for n in ns]
    if not all_nums:
        return
    # 已连续 (从 9.1 起, 无重复, 无缺号) → 无需修复
    if (min(all_nums) == 1
            and len(all_nums) == len(set(all_nums))
            and max(all_nums) - min(all_nums) + 1 == len(all_nums)):
        return
    # 重编号: 按出现顺序连续 9.1..9.N
    new_seq = 1
    mapping: list[str] = []
    for f, old_nums in field_nums:
        new_nums = list(range(new_seq, new_seq + len(old_nums)))
        new_seq += len(old_nums)
        if new_nums == old_nums:
            continue
        it = iter(new_nums)
        new_label = re.sub(r"(?m)^9\.\d+", lambda m: f"9.{next(it)}", f.label)
        if new_label != f.label:
            mapping.append(f"9.{old_nums[0]}→9.{new_nums[0]}")
            f.label = new_label
    if mapping:
        result.anomalies.append(Anomaly(
            "info", 9, "S9 编号重排(禁跳号)",
            "按文档顺序连续重编号: " + ", ".join(mapping),
        ))


def _post_check(result: ParseResult) -> None:
    """读取后置修复+检查: S9 编号重排(禁跳号) / 缺失节.

    注意: 不再报告"字段值为空"告警 —— 模板字段留空待填充是常态
    (全库 646 文件 590 个均有空值字段, 如 供应商信息:/产品名称: 子标题行),
    空值在 GUI 三列表中直观呈现即可, 不属于解析异常.
    """
    # S9 编号: 先重排修复 (禁止跳号), 若仍有残留缺失再告警 (用 findall
    # 识别一行多编号, 英文文件把 "9.19\n9.20\n9.21" 合并进一格)
    _renumber_s9(result)
    sec9 = result.sections.get(9)
    if sec9:
        nums = []
        for f in sec9.fields:
            nums.extend(int(m) for m in re.findall(r"(?m)^9\.(\d+)", f.label))
        if nums:
            expected = set(range(min(nums), max(nums) + 1))
            missing = sorted(expected - set(nums))
            if missing:
                result.anomalies.append(Anomaly(
                    "warn", 9, "S9 编号不连续", f"缺失: {', '.join(f'9.{n}' for n in missing)}",
                ))
    # 16 节完整性
    for n in range(1, 17):
        if n not in result.sections:
            result.anomalies.append(Anomaly("error", 0, f"缺失第{n}节"))


# ============================================================
# 内化默认模板 (用户确认 2026-08-17: 定稿模板 PEA-4139 MSDS_CN 冠志 为唯一源/参照)
#   定稿模板优先, 内化副本 (templates/MSDS_CN 国彩 模板.docx) 作回退 (离线/迁移).
# ============================================================
_TEMPLATE_DINGGAO = Path(
    r"F:\正式项目与模块化内容\冠志\MSDS\MSDS 数据清理模块\标准模板\标准模板\定稿模板\PEA-4139 MSDS_CN 冠志 模板.docx")
_TEMPLATE_RESOURCE = (Path(__file__).resolve().parent.parent
                      / "templates" / "MSDS_CN 国彩 模板.docx")
_TEMPLATE_EXTERNAL = Path(
    r"F:\正式项目与模块化内容\MSDS 数据清理模块\标准模板\标准模板\定稿模板\模板草稿\MSDS_CN 国彩 模板.docx")


def _resolve_template() -> Path:
    """解析默认模板: 优先定稿模板 (冠志版), 回退内化副本/外部源路径."""
    for cand in (_TEMPLATE_DINGGAO, _TEMPLATE_RESOURCE, _TEMPLATE_EXTERNAL):
        if cand.exists():
            return cand
    return _TEMPLATE_RESOURCE  # 均不存在 → 返回内化路径 (报错信息友好)


TEMPLATE_PATH = _resolve_template()
