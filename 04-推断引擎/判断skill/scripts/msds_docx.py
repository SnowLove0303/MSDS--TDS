# -*- coding: utf-8 -*-
"""
msds_docx.py — 把 content.json（16 节推导内容）渲染成 Word 文档。

## 两种输出（每次使用都建议产出）
1. **方案 Word**（默认）：含封面/诊断汇总表/推导依据/16节（【原文问题】+【规范应写内容】）。
   - 可写入 MSDS 的内容：正常显示；
   - 仅说明的内容（推导理由/编辑指令/字段提醒）：灰色楷体 + 【说明】前缀。
2. **纯 MSDS 正文 Word**（--pure）：只抽取"可写入 MSDS"的内容，**按标准表格结构**输出，
   对标冠志 16 表模板（templates/sds_structure.json），并符合 GB/T 17519-2013 封面/页眉/页脚格式。

## content.json 结构（见 templates/content.example.json）
{
  "title": "封面标题（方案）",
  "sds": {"product": "产品名", "company": "供应商", "date": "修订日期",
          "first_date": "最初编制日期", "sds_no": "SDS编号", "version": "版本", "basis": "编制依据"},
  "meta": ["...", "..."],
  "overview": ["...", "..."],
  "diagnosis_header": ["列1", "列2", "列3"],
  "diagnosis_rows": [["节", "问题", "程度"], ...],
  "basis_blocks": [ {"title": "2.1 标准依据", "items": ["..."]}, ... ],
  "sections": [ {"num": "第1部分", "title": "化学品及企业标识",
                  "diagnosis": "原文问题", "content": [...] }, ... ],
  "references": ["..."]
}

## content[] 条目类型（可写入 MSDS vs 仅说明）
- **字符串** → 视为"可写入 MSDS"正文。若形如"字段名：值"（冒号≤28字），纯 MSDS 渲染为表格行 [字段名|值]；
  否则为整行（跨列）。
- **dict**：
  - {"t":"msds","text":"..."}    可写入 MSDS（同字符串）
  - {"t":"note","text":"..."}    仅说明（不进纯 MSDS）
  - {"t":"pictogram","code":"GHS05"}   嵌象形图（S2 用 GHS01~09；S14 用 TDG 系列）
  - {"t":"table","label":"成分表","header":["化学品名称","CAS编号","含量%"],"rows":[[...],...]}
      结构化多列表（纯 MSDS 渲染为独立子表；label 对应 sds_structure.json 的 __table__: 占位）

用法:
    python msds_docx.py "<content.json>" [--out "<方案>.docx"]            # 方案 Word
    python msds_docx.py "<content.json>" --pure --out "<纯MSDS>.docx"      # 纯 MSDS（表格化）
"""
import argparse, io, json, os, re, sys

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("缺少 python-docx，请先安装: pip install python-docx")

SKILL_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
PICTO_DIR = os.path.join(SKILL_DIR, "templates", "pictograms")
STRUCTURE_PATH = os.path.join(SKILL_DIR, "templates", "sds_structure.json")

BLUE = (0x1F, 0x4E, 0x79)
WHITE = (0xFF, 0xFF, 0xFF)


# ---------------------------------------------------------------- 基础工具
def _norm_label(s):
    """归一化字段名：去空格、全半角括号/冒号统一，用于匹配。"""
    s = str(s).replace(' ', '').replace('　', '')
    return s.replace('（', '(').replace('）', ')').replace('：', ':')


def split_label_value(text):
    """把『字段名：值』拆成 (label, value)；拆不出则返回 (None, 原文)。

    规则：行首 ≤28 字的『xxx：』视为字段名；label 不能以句子标点（。，；、,.）或冒号结尾。
    括号（如『产品名称（商品名）』『GHS危险性类别（依据…）』）是合法字段名，不禁用。
    """
    text = (text or '').strip()
    m = re.match(r'^\s*(.{1,28}?)\s*[:：]\s*(.*)$', text, re.S)
    if m:
        label = m.group(1).strip()
        if label and not label[-1] in '。，；、,.·—:：':
            return label, m.group(2).strip()
    return None, text


def set_font(run, name_cn='宋体', name_en='Times New Roman', size=10.5, bold=False, color=None):
    run.font.name = name_en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name_cn)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn('w:shd'), {qn('w:val'): 'clear', qn('w:fill'): fill})
    tcPr.append(shd)


def cell_text(cell, text, size=9.5, bold=False, cn='宋体', color=None, align=None, space_after=0):
    """写入单元格文本；text 中的 '\\n' 转为单元格内换行（软换行），
    保持多句内容在同一单元格内分行（对齐模板的换行风格）。"""
    cell.text = ''
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    lines = str(text).split('\n')
    for i, ln in enumerate(lines):
        if i:
            p.add_run().add_break()   # WD_BREAK.LINE 软换行
        r = p.add_run(ln)
        set_font(r, name_cn=cn, size=size, bold=bold, color=color)
    return p


def add_page_number(par):
    """在段落中插入页码域。"""
    for tpl in ('begin', 'instr', 'end'):
        r = par.add_run()
        if tpl == 'begin':
            f = OxmlElement('w:fldChar'); f.set(qn('w:fldCharType'), 'begin'); r._r.append(f)
        elif tpl == 'instr':
            it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = ' PAGE '
            r._r.append(it)
        else:
            f = OxmlElement('w:fldChar'); f.set(qn('w:fldCharType'), 'end'); r._r.append(f)


def pictogram_path(code):
    p = os.path.join(PICTO_DIR, code + ".png")
    return p if os.path.exists(p) else None


def add_pictogram_para(doc, code, size_cm=2.0, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    img = pictogram_path(code)
    if img:
        p.add_run().add_picture(img, width=Cm(size_cm))
    else:
        r = p.add_run("【象形图】" + code)
        set_font(r, name_cn='宋体', size=10.5)
    return p


def item_parts(item):
    """把 content[] 条目解析成 (text, is_note)。"""
    if isinstance(item, dict):
        return item.get("text", ""), item.get("t") == "note"
    return item, False


def load_structure():
    if os.path.exists(STRUCTURE_PATH):
        with open(STRUCTURE_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    return None


# ================================================================ 方案 Word
def build(content, out_path):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.2); s.bottom_margin = Cm(2.2)
        s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

    def para(text, size=10.5, bold=False, cn='宋体', align=None, space_after=6, color=None):
        p = doc.add_paragraph()
        r = p.add_run(text)
        set_font(r, name_cn=cn, size=size, bold=bold, color=color)
        if align:
            p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        return p

    def heading(text, level=1):
        sizes = {0: 16, 1: 14, 2: 12, 3: 11}
        p = doc.add_paragraph()
        r = p.add_run(text)
        set_font(r, name_cn='黑体', size=sizes.get(level, 11), bold=True, color=BLUE)
        p.paragraph_format.space_before = Pt(12 if level > 0 else 6)
        p.paragraph_format.space_after = Pt(6)
        return p

    def render_table_item(payload):
        tbl = doc.add_table(rows=1, cols=len(payload.get("header", [])))
        tbl.style = 'Table Grid'
        for j, h in enumerate(payload.get("header", [])):
            c = tbl.rows[0].cells[j]
            cell_text(c, str(h), size=9.5, bold=True, cn='黑体')
            cell_shading(c, 'DCE6F1')
        for row in payload.get("rows", []):
            cells = tbl.add_row().cells
            for j, v in enumerate(row):
                cell_text(cells[j], str(v), size=9.5)
        para("", size=4)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(content.get("title", "化学品安全技术说明书规范化推导方案"))
    set_font(r, name_cn='黑体', size=18, bold=True, color=BLUE)
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("——依据第1、3、9部分推导其余各部分的编写方案")
    set_font(r2, name_cn='楷体', size=13, color=(0x44, 0x44, 0x44))
    doc.add_paragraph()

    for line in content.get("meta", []):
        para(line, size=10, space_after=3)

    heading("一、任务概述与原文诊断", 0)
    for t in content.get("overview", []):
        para(t)
    if content.get("diagnosis_rows"):
        heading("1.1 原文主要规范性问题汇总", 2)
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = 'Table Grid'
        for i, h in enumerate(content.get("diagnosis_header", ["部分", "原文存在问题", "严重程度"])):
            c = tbl.rows[0].cells[i]
            c.text = ''
            r = c.paragraphs[0].add_run(h); set_font(r, name_cn='黑体', size=10, bold=True)
            cell_shading(c, 'DCE6F1')
        for row in content["diagnosis_rows"]:
            cells = tbl.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = ''
                r = cells[i].paragraphs[0].add_run(str(v)); set_font(r, size=9.5)
        para("", size=4)

    heading("二、推导依据", 0)
    for blk in content.get("basis_blocks", []):
        heading(blk.get("title", ""), 2)
        for t in blk.get("items", []):
            para("· " + t, size=10)

    heading("三、各部分的规范化推导（具体应写内容）", 0)
    para("以下按 GB/T 16483-2008 的 16 个部分逐一给出：原文存在问题 → 推导后的规范应写内容。"
         "其中第1、3、9部分为已知事实（据原文档整理并规范化），其余部分由前三部分推导得出。"
         "『字段名：值』条目可直接映射到冠志模板表格行。", size=10.5)
    for s in content.get("sections", []):
        heading(f"{s.get('num','')}  {s.get('title','')}", 1)
        para("【原文问题】" + s.get("diagnosis", ""), size=10, color=(0x9C, 0x27, 0x00))
        para("【推导后的规范应写内容】", size=10.5, bold=True, cn='黑体')
        for line in s.get("content", []):
            if isinstance(line, dict) and line.get("t") == "pictogram":
                add_pictogram_para(doc, line.get("code", ""))
                continue
            if isinstance(line, dict) and line.get("t") == "table":
                render_table_item(line)
                continue
            text, is_note = item_parts(line)
            if is_note:
                para("【说明】" + text, size=10, cn='楷体', color=(0x80, 0x80, 0x80), space_after=2)
            else:
                para(text, size=10.5, space_after=2)

    heading("四、参考资料", 0)
    for t in content.get("references", []):
        para("· " + t, size=10)

    os.makedirs(os.path.dirname(os.path.abspath(out_path)) or ".", exist_ok=True)
    doc.save(out_path)
    print("[OK] 方案 Word 已生成:", out_path)


# ================================================================ 纯 MSDS（表格化）
def setup_header_footer(doc, sds):
    """GB/T 17519-2013 §4.2：首页上部封面；首页后各页眉=产品名称/修订日期/SDS编号；页脚页码。"""
    product = sds.get("product") or ""
    sds_no = sds.get("sds_no") or ""
    date = sds.get("date") or ""
    sec = doc.sections[0]
    sec.different_first_page_header_footer = True

    for hf in (sec.footer, sec.first_page_footer):
        p = hf.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("第 "); set_font(r, size=9)
        add_page_number(p)
        r = p.add_run(" 页  共 "); set_font(r, size=9)
        r2 = p.add_run()
        f = OxmlElement('w:fldChar'); f.set(qn('w:fldCharType'), 'begin'); r2._r.append(f)
        it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = ' NUMPAGES '
        r2._r.append(it)
        f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end'); r2._r.append(f2)
        set_font(r2, size=9)
        r = p.add_run(" 页"); set_font(r, size=9)

    ph = sec.header.paragraphs[0]
    ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ph.add_run(f"{product}    SDS编号：{sds_no}    修订日期：{date}")
    set_font(r, name_cn='宋体', size=9)


def render_pure_cover(doc, sds):
    """GB/T 17519-2013 §4.2.1 首页上部。"""
    product = sds.get("product") or ""
    date = sds.get("date") or ""
    first_date = sds.get("first_date") or ""
    sds_no = sds.get("sds_no") or ""
    version = sds.get("version") or ""
    basis = sds.get("basis") or "按照 GB/T 16483、GB/T 17519 编制"

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("化学品安全技术说明书")
    set_font(r, name_cn='黑体', size=20, bold=True, color=BLUE)
    p.paragraph_format.space_after = Pt(18)

    tbl = doc.add_table(rows=3, cols=4)
    tbl.style = 'Table Grid'
    rows = [
        ("产品名称：", product, "编制依据：", basis),
        ("修订日期：", date, "SDS编号：", sds_no),
        ("最初编制日期：", first_date, "版本：", version),
    ]
    for i, (la, va, lb, vb) in enumerate(rows):
        cell_text(tbl.rows[i].cells[0], la, size=10.5, bold=True)
        cell_text(tbl.rows[i].cells[1], va, size=10.5)
        cell_text(tbl.rows[i].cells[2], lb, size=10.5, bold=True)
        cell_text(tbl.rows[i].cells[3], vb, size=10.5)
    doc.add_paragraph()
    return tbl


def _collect_msds_items(section):
    """从节的 content 提取可写入 MSDS 的条目（剔除 note），返回统一列表。

    每项为 (kind, label, payload)：
      ('kv',   label, (label, value))
      ('full', None,  text)
      ('picto', None, {'code':...})
      ('table', table_label, {'label','header','rows'})
    """
    out = []
    for x in section.get("content", []):
        if isinstance(x, dict) and x.get("t") == "pictogram":
            out.append(('picto', None, x))
        elif isinstance(x, dict) and x.get("t") == "table":
            out.append(('table', x.get("label", ""), x))
        else:
            text, is_note = item_parts(x)
            if is_note or not text.strip():
                continue
            label, value = split_label_value(text)
            if label is not None:
                out.append(('kv', label, (label, value)))
            else:
                out.append(('full', None, text))
    return out


def render_section_table(doc, sec_def, section, empty_policy='skip'):
    """渲染一节为标准表格。

    sec_def: sds_structure.json 的节定义（title/fields）
    empty_policy: 'skip' 缺字段跳过；'ndata' 缺字段填『无数据』；'na' 填『不适用』
    """
    fields = sec_def.get("fields", [])
    num = sec_def.get("num", "")
    title = sec_def.get("title", "")
    items = _collect_msds_items(section)

    # 字段别名映射：fields 每项可为 "显示名" 或 ["显示名","别名1",...]
    # field_meta: [(display, [alias_norm,...]), ...] 保持 fields 顺序
    field_meta = []
    kv_slots = {}      # norm(alias) -> list[(label, value)]
    table_slots = {}   # norm(table_label) -> list[payload]
    list_slots = {}    # norm(list_label) -> list[full_text]
    def _fname(f):
        return str(f[0]) if isinstance(f, (list, tuple)) else str(f)

    for f in fields:
        fn = _fname(f)
        if fn.startswith('__table__:'):
            key = _norm_label(fn.split(':', 1)[1])
            table_slots.setdefault(key, [])
        elif fn.startswith('__list__:'):
            key = _norm_label(fn.split(':', 1)[1])
            list_slots.setdefault(key, [])
        else:
            if isinstance(f, (list, tuple)):
                display = str(f[0]); aliases = [_norm_label(x) for x in f]
            else:
                display = str(f); aliases = [_norm_label(f)]
            # 保序去重：同字段别名重复会导致同一内容被重复渲染
            seen, uniq = set(), []
            for a in aliases:
                if a not in seen:
                    seen.add(a); uniq.append(a)
            field_meta.append((display, uniq))
            for a in uniq:
                kv_slots.setdefault(a, [])

    def _to_list(line):
        first_key = next(iter(list_slots))
        list_slots[first_key].append(line)

    extras = []
    for kind, label, payload in items:
        if kind == 'kv':
            key = _norm_label(label)
            if key in kv_slots:
                kv_slots[key].append(payload)
            elif list_slots:          # 法规/其他信息节：未匹配 kv 整行统一归入列表
                _to_list(f"{label}：{payload[1]}")
            else:
                extras.append(payload)
        elif kind == 'table':
            key = _norm_label(label)
            if key in table_slots:
                table_slots[key].append(payload)
            else:
                extras.append(payload)
        elif kind == 'picto':
            extras.append(('PICTO', payload))
        else:                          # full 整行
            if list_slots:
                _to_list(payload)
            else:
                extras.append(('FULL', payload))

    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hr = tbl.add_row()
    merged = hr.cells[0].merge(hr.cells[1])
    cell_text(merged, f"{num}.{title}" if not str(num).startswith('第') else f"{num} {title}",
              size=10.5, bold=True, cn='黑体', color=WHITE)
    cell_shading(merged, '1F4E79')

    # 分离 FULL 概述行（无字段名整行，如 S11/S12 首行）→ 渲染在表头之后、字段之前
    full_items = [it for it in extras if isinstance(it, tuple) and it and it[0] == 'FULL']
    extras = [it for it in extras if not (isinstance(it, tuple) and it and it[0] == 'FULL')]
    for _, payload in full_items:
        row = tbl.add_row()
        merged = row.cells[0].merge(row.cells[1])
        cell_text(merged, payload, size=9.5, bold=True)

    # 分离象形图条目：优先嵌入"象形图"字段值列；无该字段的节（如S14）才作合并行
    picto_items = [it for it in extras if isinstance(it, tuple) and it and it[0] == 'PICTO']
    extras = [it for it in extras if not (isinstance(it, tuple) and it and it[0] == 'PICTO')]

    for display, aliases in field_meta:
        collected = []
        for a in aliases:
            collected.extend(kv_slots[a])
        if '象形图' in display:
            # 象形图字段：左列字段名，右列 = 文本值（若有）+ 图片（多个）
            row = tbl.add_row()
            cell_text(row.cells[0], display, size=9.5, bold=True, cn='宋体')
            vc = row.cells[1]
            vc.text = ''
            first = True
            if collected and collected[0][1]:
                p = vc.paragraphs[0]
                r = p.add_run(collected[0][1])
                set_font(r, name_cn='宋体', size=9.5)
                first = False
            for _, payload in picto_items:
                img = pictogram_path(payload.get("code", ""))
                p = vc.paragraphs[0] if first else vc.add_paragraph()
                first = False
                if img:
                    p.add_run().add_picture(img, width=Cm(1.5))
                else:
                    r = p.add_run("【象形图】" + payload.get("code", ""))
                    set_font(r, name_cn='宋体', size=9.5)
            picto_items = []
            continue
        for _, value in collected:
            row = tbl.add_row()
            cell_text(row.cells[0], display, size=9.5, bold=True, cn='宋体')
            cell_text(row.cells[1], value, size=9.5)
        if not collected and empty_policy in ('ndata', 'na'):
            row = tbl.add_row()
            cell_text(row.cells[0], display, size=9.5, bold=True)
            cell_text(row.cells[1], '无数据' if empty_policy == 'ndata' else '不适用', size=9.5)

    # 未消费的象形图（无"象形图"字段的节，如 S14 运输标志）：回补为合并行
    for _, payload in picto_items:
        row = tbl.add_row()
        merged = row.cells[0].merge(row.cells[1])
        img = pictogram_path(payload.get("code", ""))
        if img:
            merged.paragraphs[0].add_run().add_picture(img, width=Cm(1.6))
        else:
            cell_text(merged, "【象形图】" + payload.get("code", ""))

    # 纯列表节（__list__ 且无普通字段，如 S15 法规信息 / S16 其他信息）：
    # 按模板单列表格渲染（表头 + 每条内容一行）
    if list_slots and not field_meta:
        for lines in list_slots.values():
            for line in lines:
                row = tbl.add_row()
                merged = row.cells[0].merge(row.cells[1])
                cell_text(merged, line, size=9.5)

    for item in extras:
        label, value = item
        row = tbl.add_row()
        cell_text(row.cells[0], label, size=9.5, bold=True)
        cell_text(row.cells[1], value, size=9.5)

    for f in fields:
        fn = _fname(f)
        if fn.startswith('__table__:'):
            label = _norm_label(fn.split(':', 1)[1])
            for payload in table_slots[label]:
                doc.add_paragraph()
                ph = doc.add_paragraph()
                r = ph.add_run(payload.get("label", "成分表"))
                set_font(r, name_cn='黑体', size=10, bold=True)
                sub = doc.add_table(rows=1, cols=len(payload.get("header", [])))
                sub.style = 'Table Grid'
                for j, h in enumerate(payload.get("header", [])):
                    c = sub.rows[0].cells[j]
                    cell_text(c, str(h), size=9, bold=True, cn='黑体')
                    cell_shading(c, 'DCE6F1')
                for row in payload.get("rows", []):
                    cells = sub.add_row().cells
                    for j, v in enumerate(row):
                        cell_text(cells[j], str(v), size=9)
        elif fn.startswith('__list__:'):
            label = _norm_label(fn.split(':', 1)[1])
            # 纯列表节已在主表单列渲染；仅当同时存在普通字段时补充附加列表
            if list_slots[label] and field_meta:
                doc.add_paragraph()
                for text in list_slots[label]:
                    p = doc.add_paragraph()
                    r = p.add_run("· " + text)
                    set_font(r, name_cn='宋体', size=9.5)
                    p.paragraph_format.space_after = Pt(2)


def build_pure(content, out_path):
    """纯 MSDS 正文：GB/T 17519 封面/页眉/页脚 + 冠志 16 表表格结构。"""
    doc = Document()
    for s in doc.sections:
        s.page_width = Cm(21.0); s.page_height = Cm(29.7)
        s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.0); s.right_margin = Cm(2.0)

    sds = content.get("sds", {})
    setup_header_footer(doc, sds)
    render_pure_cover(doc, sds)

    structure = load_structure()
    sec_defs = {d.get("num"): d for d in structure.get("sections", [])} if structure else {}
    if not sec_defs:
        nums = [str(i) for i in range(1, 17)]
        sec_defs = {n: {"num": n, "title": s.get("title", ""), "fields": []}
                    for n, s in zip(nums, content.get("sections", []))}

    order = ["1", "2", "3", "4", "5", "6", "7", "8", "9", "10",
             "11", "12", "13", "14", "15", "16"]
    content_by_num = {}
    for s in content.get("sections", []):
        m = re.search(r'(\d+)$', str(s.get("num", "")).replace("第", "").replace("部分", ""))
        key = m.group(1) if m else None
        if key:
            content_by_num[key] = s
        else:
            content_by_num.setdefault(len(content_by_num) + 1, s)

    for n in order:
        sec_def = sec_defs.get(n)
        sec_content = content_by_num.get(n)
        if not sec_def or not sec_content:
            continue
        policy = 'ndata' if n == '9' else ('na' if n == '14' else 'skip')
        render_section_table(doc, sec_def, sec_content, empty_policy=policy)

    os.makedirs(os.path.dirname(os.path.abspath(out_path)) or ".", exist_ok=True)
    doc.save(out_path)
    print("[OK] 纯 MSDS 正文 Word 已生成（表格化）:", out_path)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("content_json")
    ap.add_argument("--pure", action="store_true", help="只输出纯 MSDS 正文（表格化，剔除说明）")
    ap.add_argument("--out", help="输出 .docx 路径（默认与 content.json 同名替换为 .docx）")
    args = ap.parse_args()
    with open(args.content_json, "r", encoding="utf-8") as f:
        content = json.load(f)
    out = args.out or os.path.splitext(args.content_json)[0] + ".docx"
    if args.pure:
        build_pure(content, os.path.abspath(out))
    else:
        build(content, os.path.abspath(out))


if __name__ == "__main__":
    main()
