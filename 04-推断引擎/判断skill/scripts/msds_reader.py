# -*- coding: utf-8 -*-
"""
msds_reader.py — 读取任意 MSDS .docx，抽出 16 节（重点 S1/S3/S9），输出 facts.json 与终端摘要。

用法:
    python msds_reader.py "<源MSDS.docx>" [--out "<输出目录>/facts.json"] [--all]

- 兼容段落式与表格式（含合并单元格）模板。
- --all 时把全部 16 节原文都写入 facts.json（默认只存 S1/S3/S9 与节索引）。
"""
import argparse, io, json, os, re, sys

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

try:
    import docx
except ImportError:
    sys.exit("缺少 python-docx，请先安装: pip install python-docx")

# GB/T 16483 16 节标题关键词（用于把表格/段落归到节号）
SECTION_PATTERNS = [
    (1,  r"1\s*[.．、]\s*(物料及供应商标识|化学品及企业标识|化学品和企业标识)"),
    (2,  r"2\s*[.．、]\s*(危险性概述)"),
    (3,  r"3\s*[.．、]\s*(成分[／/]组成信息|成分/组成资料|成分和组成信息)"),
    (4,  r"4\s*[.．、]\s*(急救措施|急救)"),
    (5,  r"5\s*[.．、]\s*(消防措施|火灾爆炸措施|灭火措施)"),
    (6,  r"6\s*[.．、]\s*(泄漏应急处理|意外泄漏措施|泄漏的应急措施)"),
    (7,  r"7\s*[.．、]\s*(操作处置与储存|操作和储存)"),
    (8,  r"8\s*[.．、]\s*(接触控制和个体防护|接触控制[／/]个人防护|暴露控制和个人防护)"),
    (9,  r"9\s*[.．、]\s*(理化特性|物理和化学性质|物理和化学特性)"),
    (10, r"10\s*[.．、]\s*(稳定性和反应性|稳定性和反应)"),
    (11, r"11\s*[.．、]\s*(毒理学信息|毒理学资料|毒性资料|毒理性)"),
    (12, r"12\s*[.．、]\s*(生态学信息|生态信息|生态资料|生态毒理)"),
    (13, r"13\s*[.．、]\s*(废弃处置|废弃注意事项|处理注意事项|废弃处理)"),
    (14, r"14\s*[.．、]\s*(运输信息|运输资料)"),
    (15, r"15\s*[.．、]\s*(法规信息|法规资料|法规及标准)"),
    (16, r"16\s*[.．、]\s*(其他信息|其它信息|其他资料)"),
]

def classify_section(text: str):
    """返回 (节号, 匹配标题) 或 (None, None)。"""
    if not text:
        return None, None
    for num, pat in SECTION_PATTERNS:
        m = re.search(pat, text.strip())
        if m:
            return num, m.group(0).strip()
    return None, None

def cell_text(cell):
    t = cell.text or ""
    return "\n".join(x.strip() for x in t.split("\n") if x.strip()).strip()

def read_docx(path):
    d = docx.Document(path)
    raw = []  # 按出现顺序收集文本块：[(来源, 节号或None, 文本)]
    # 段落
    for p in d.paragraphs:
        t = p.text.strip()
        if t:
            num, _ = classify_section(t)
            raw.append(("para", num, t))
    # 表格：按行（横向合并单元格行内去重；纵向合并按 python-docx 返回的原始 cell 文本保留）
    for ti, table in enumerate(d.tables):
        for row in table.rows:
            seen_in_row = set()
            cells = []
            for c in row.cells:
                if id(c._tc) in seen_in_row:
                    continue
                seen_in_row.add(id(c._tc))
                cells.append(cell_text(c))
            line = " | ".join(x for x in cells if x)
            if line:
                num, _ = classify_section(line)
                raw.append((f"table{ti}", num, line))
    return d, raw

def group_sections(raw):
    """按节号归组，节内保留顺序。"""
    sections = {}
    order = []
    cur = None
    for src, num, text in raw:
        if num is not None:
            cur = num
            if num not in sections:
                sections[num] = []
                order.append(num)
        if cur is not None and num is None:
            sections[cur].append(text)
        elif num is not None and text != sections.get(num, []):
            # 标题行本身也作为一个锚点信息记录（放在该节首）
            pass
    return sections, order

def extract_key_facts(sections):
    """从 S1/S3/S9 提取关键事实（启发式，供 AI 快速核对）。"""
    facts = {"sections_found": {}, "s1": {}, "s3": [], "s9": {}}

    s3 = sections.get(3, [])
    if s3:
        # 尝试解析"名称 | CAS | 含量"或"名称 / CAS / 含量"
        for line in s3:
            for sep in [" | ", " / ", "  ", "，"]:
                parts = [p.strip() for p in line.split(sep) if p.strip()]
                if len(parts) >= 2:
                    name = parts[0]
                    cas = ""
                    content = ""
                    for p in parts[1:]:
                        if re.fullmatch(r"\d{2,7}-\d{2}-\d{1}", p) or "CAS" in p.upper():
                            cas = p
                        elif re.search(r"%|\d\s*[—\-~]\s*\d", p) or "交易" in p or "秘密" in p:
                            content = p
                    if name and (cas or content):
                        facts["s3"].append({"name": name, "cas": cas, "content": content})
                        break
    return facts

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("docx")
    ap.add_argument("--out", help="facts.json 输出路径")
    ap.add_argument("--all", action="store_true", help="保存全部 16 节原文")
    args = ap.parse_args()

    path = os.path.abspath(args.docx)
    if not os.path.exists(path):
        sys.exit(f"文件不存在: {path}")
    if os.path.basename(path).startswith("~$"):
        sys.exit("这是 Word 锁文件，请提供正式文件")

    doc, raw = read_docx(path)
    sections, order = group_sections(raw)

    facts = {
        "source": path,
        "sections_found": {n: f"S{n}" for n in order},
        "missing_sections": [n for n in range(1, 17) if n not in sections],
        "s3_raw": sections.get(3, []),
        "s9_raw": sections.get(9, []),
        "s1_raw": sections.get(1, []),
    }
    if args.all:
        facts["all_raw"] = {str(n): sections[n] for n in order}

    # 终端摘要
    print(f"文件: {os.path.basename(path)}")
    print(f"检测到节: {', '.join('S%d'%n for n in order) or '无'}  |  缺失节: "
          f"{', '.join('S%d'%n for n in facts['missing_sections']) or '无'}")
    print("\n===== S1 物料及供应商标识 (摘要) =====")
    for line in sections.get(1, [])[:14]:
        print("  " + line[:120])
    print("\n===== S3 成分/组成信息 (原文) =====")
    for line in sections.get(3, []):
        print("  " + line[:160])
    print("\n===== S9 物理和化学特性 (摘要) =====")
    for line in sections.get(9, [])[:30]:
        print("  " + line[:120])

    if args.out:
        os.makedirs(os.path.dirname(os.path.abspath(args.out)), exist_ok=True)
        with open(args.out, "w", encoding="utf-8") as f:
            json.dump(facts, f, ensure_ascii=False, indent=2)
        print(f"\n[OK] facts.json 已写入: {args.out}")

if __name__ == "__main__":
    main()
